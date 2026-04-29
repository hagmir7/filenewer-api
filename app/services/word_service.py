"""
Word / DOCX service functions.
"""

import io
import re
import logging

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .pdf_service import is_arabic, process_arabic_text, register_arabic_fonts

logger = logging.getLogger(__name__)


def word_to_pdf(source, filename: str = "document.docx") -> bytes:
    """
    Convert Word (.docx) → PDF with full Arabic / RTL support.
    """
    from docx import Document as DocxDocument
    from docx.text.paragraph import Paragraph as DocxPara
    from docx.table import Table as DocxTable
    from docx.oxml.ns import qn
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
        PageBreak,
    )

    # ── Read DOCX ─────────────────────────────────
    if hasattr(source, "read"):
        file_bytes = source.read()
    else:
        file_bytes = source

    docx = DocxDocument(io.BytesIO(file_bytes))

    # ── Register Arabic fonts ──────────────────────
    registered_fonts = register_arabic_fonts()
    has_arabic_font = bool(registered_fonts)

    arabic_font = "Amiri" if "Amiri" in registered_fonts else "Helvetica"
    arabic_font_bold = (
        "Amiri-Bold" if "Amiri-Bold" in registered_fonts else "Helvetica-Bold"
    )
    latin_font = "Helvetica"
    latin_font_bold = "Helvetica-Bold"

    # ── Helper: pick font based on content ─────────
    def pick_font(text: str, bold: bool = False) -> str:
        if is_arabic(text):
            return arabic_font_bold if bold else arabic_font
        return latin_font_bold if bold else latin_font

    # ── Helper: process text (Arabic or Latin) ─────
    def prepare_text(text: str) -> str:
        if is_arabic(text):
            return process_arabic_text(text)
        return text

    # ── Helper: escape XML special chars ───────────
    def escape(text: str) -> str:
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    # ── Custom styles ──────────────────────────────
    def make_style(name, font, font_bold, size, color, before, after, align=TA_LEFT):
        return ParagraphStyle(
            name,
            fontName=font_bold,
            fontSize=size,
            textColor=color,
            spaceBefore=before,
            spaceAfter=after,
            leading=size * 1.4,
            alignment=align,
            wordWrap="RTL" if font == arabic_font else "LTR",
        )

    # ── Process paragraph runs → markup ────────────
    def runs_to_markup(para) -> tuple[str, bool]:
        """
        Returns (markup_string, has_arabic).
        Uses correct font per run.
        """
        parts = []
        has_arabic = False

        for run in para.runs:
            raw = run.text
            if not raw:
                continue

            arabic = is_arabic(raw)
            if arabic:
                has_arabic = True
                raw = process_arabic_text(raw)

            text = escape(raw)
            font = pick_font(run.text, bold=run.bold)

            # Wrap in font tag
            text = f'<font name="{font}">{text}</font>'

            if run.bold:
                text = f"<b>{text}</b>"
            if run.italic:
                text = f"<i>{text}</i>"
            if run.underline:
                text = f"<u>{text}</u>"

            parts.append(text)

        return "".join(parts), has_arabic

    # ── Pick paragraph style ───────────────────────
    def pick_style(para, has_arabic: bool) -> ParagraphStyle:
        name = (para.style.name or "Normal").lower()
        align = TA_RIGHT if has_arabic else TA_JUSTIFY
        font = arabic_font if has_arabic else latin_font
        bold = arabic_font_bold if has_arabic else latin_font_bold
        wrap = "RTL" if has_arabic else "LTR"

        base = {
            "fontName": font,
            "leading": 20,
            "wordWrap": wrap,
            "alignment": align,
        }

        if "heading 1" in name:
            return ParagraphStyle(
                "H1",
                **base,
                fontSize=22,
                textColor=colors.HexColor("#2E75B6"),
                fontName=bold,
                spaceBefore=18,
                spaceAfter=14,
            )

        if "heading 2" in name:
            return ParagraphStyle(
                "H2",
                **base,
                fontSize=16,
                textColor=colors.HexColor("#2E75B6"),
                fontName=bold,
                spaceBefore=14,
                spaceAfter=10,
            )

        if "heading 3" in name:
            return ParagraphStyle(
                "H3",
                **base,
                fontSize=13,
                textColor=colors.HexColor("#1F4E79"),
                fontName=bold,
                spaceBefore=10,
                spaceAfter=8,
            )

        if "list bullet" in name:
            return ParagraphStyle(
                "Bullet", **base, fontSize=11, leftIndent=20, spaceAfter=4
            )

        if "list number" in name:
            return ParagraphStyle(
                "Number", **base, fontSize=11, leftIndent=20, spaceAfter=4
            )

        return ParagraphStyle("Body", **base, fontSize=11, spaceAfter=6)

    # ── Process table ──────────────────────────────
    def process_table(tbl) -> Table:
        data = []
        col_cnt = max(len(row.cells) for row in tbl.rows)
        col_w = (A4[0] - 2 * inch) / col_cnt

        for r_idx, row in enumerate(tbl.rows):
            row_data = []
            for cell in row.cells:
                cell_text = "\n".join(p.text for p in cell.paragraphs)
                arabic = is_arabic(cell_text)

                if arabic:
                    cell_text = process_arabic_text(cell_text)

                font = (
                    arabic_font_bold
                    if (r_idx == 0 and arabic)
                    else (
                        arabic_font
                        if arabic
                        else (latin_font_bold if r_idx == 0 else latin_font)
                    )
                )

                align = TA_RIGHT if arabic else TA_LEFT
                wrap = "RTL" if arabic else "LTR"

                style = ParagraphStyle(
                    f"cell_{r_idx}",
                    fontName=font,
                    fontSize=10,
                    textColor=colors.white if r_idx == 0 else colors.black,
                    alignment=align,
                    wordWrap=wrap,
                    leading=14,
                )
                row_data.append(Paragraph(escape(cell_text), style))
            data.append(row_data)

        pdf_tbl = Table(data, colWidths=[col_w] * col_cnt, repeatRows=1)
        pdf_tbl.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2E75B6")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    (
                        "ROWBACKGROUNDS",
                        (0, 1),
                        (-1, -1),
                        [colors.HexColor("#DCE6F1"), colors.white],
                    ),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#AAAAAA")),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("PADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        return pdf_tbl

    # ── Build story ────────────────────────────────
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
        title=filename.replace(".docx", "").replace(".doc", ""),
    )

    story = []

    for element in docx.element.body:
        tag = element.tag.split("}")[-1]

        # ── Paragraph ─────────────────────────────
        if tag == "p":
            para = DocxPara(element, docx)

            # Page break
            if any(
                br.get(qn("w:type")) == "page"
                for br in element.findall(".//" + qn("w:br"))
            ):
                story.append(PageBreak())
                continue

            # Empty
            if not para.text.strip():
                story.append(Spacer(1, 8))
                continue

            markup, has_arabic = runs_to_markup(para)
            style = pick_style(para, has_arabic)

            # Bullet
            name = (para.style.name or "").lower()
            if "list bullet" in name:
                prefix = "• " if not has_arabic else " •"
                try:
                    story.append(Paragraph(prefix + markup, style))
                except Exception:
                    story.append(
                        Paragraph(prefix + escape(prepare_text(para.text)), style)
                    )
                continue

            # Numbered
            if "list number" in name:
                try:
                    story.append(Paragraph("1. " + markup, style))
                except Exception:
                    story.append(
                        Paragraph("1. " + escape(prepare_text(para.text)), style)
                    )
                continue

            try:
                story.append(Paragraph(markup, style))
            except Exception:
                story.append(Paragraph(escape(prepare_text(para.text)), style))

        # ── Table ──────────────────────────────────
        elif tag == "tbl":
            try:
                tbl = DocxTable(element, docx)
                story.append(Spacer(1, 8))
                story.append(process_table(tbl))
                story.append(Spacer(1, 12))
            except Exception as e:
                story.append(
                    Paragraph(
                        f"[Table error: {e}]",
                        ParagraphStyle("err", fontName="Helvetica", fontSize=10),
                    )
                )

        elif tag == "sectPr":
            pass

    doc.build(story)
    buffer.seek(0)
    return buffer.read()


def word_to_jpg(
    source,
    filename: str = "document.docx",
    dpi: int = 200,
    quality: int = 85,
    pages: list = None,
) -> list[dict]:
    """
    Convert Word (.docx) → JPG images (one per page).

    Strategy:
        1. Convert Word → PDF (using word_to_pdf service)
        2. Convert PDF pages → JPG images (using pymupdf)

    Args:
        source   : uploaded file object OR raw bytes
        filename : original filename
        dpi      : image resolution 72-600    (default: 200)
        quality  : JPG quality 1-95           (default: 85)
        pages    : list of page numbers (1-based) None = all pages

    Returns:
        list of {
            'page'    : int,
            'bytes'   : bytes,
            'width'   : int,
            'height'  : int,
            'filename': str,
            'size_kb' : float,
        }
    """
    import fitz
    from PIL import Image

    # ── Read file bytes ────────────────────────────
    if hasattr(source, "read"):
        file_bytes = source.read()
    else:
        file_bytes = source

    if not file_bytes:
        raise ValueError("Empty file.")

    # ── Step 1: Word → PDF ─────────────────────────
    pdf_bytes = word_to_pdf(
        io.BytesIO(file_bytes),
        filename=filename,
    )

    if not pdf_bytes:
        raise RuntimeError("Failed to convert Word to PDF.")

    # ── Step 2: PDF → JPG ─────────────────────────
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    total_pages = len(doc)

    # ── Validate pages ────────────────────────────
    if pages is not None:
        invalid = [p for p in pages if not (1 <= p <= total_pages)]
        if invalid:
            raise ValueError(
                f"Invalid page numbers: {invalid}. "
                f"Document has {total_pages} pages (1-{total_pages})."
            )
        pages_to_convert = pages
    else:
        pages_to_convert = list(range(1, total_pages + 1))

    # ── Convert each page → JPG ───────────────────
    zoom = dpi / 72
    matrix = fitz.Matrix(zoom, zoom)
    results = []

    base_name = (
        filename.replace(".docx", "")
        .replace(".doc", "")
        .replace(".DOCX", "")
        .replace(".DOC", "")
    )

    for page_num in pages_to_convert:
        page = doc[page_num - 1]
        pixmap = page.get_pixmap(matrix=matrix, alpha=False)

        image = Image.frombytes(
            "RGB",
            [pixmap.width, pixmap.height],
            pixmap.samples,
        )

        buffer = io.BytesIO()
        image.save(
            buffer,
            format="JPEG",
            quality=quality,
            optimize=True,
        )
        buffer.seek(0)
        jpg_bytes = buffer.read()

        results.append(
            {
                "page": page_num,
                "bytes": jpg_bytes,
                "width": pixmap.width,
                "height": pixmap.height,
                "filename": f"{base_name}_page_{page_num}.jpg",
                "size_kb": round(len(jpg_bytes) / 1024, 2),
            }
        )

    doc.close()
    return results


def word_to_txt(
    source,
    filename: str = "document.docx",
    include_headers: bool = True,
    include_tables: bool = True,
    include_comments: bool = False,
    preserve_spacing: bool = True,
    page_separator: str = "",
    encoding: str = "utf-8",
) -> dict:
    """
    Extract text from Word (.docx) → plain text.

    Args:
        source           : uploaded file object OR raw bytes
        filename         : original filename
        include_headers  : include heading markers    (default: True)
        include_tables   : include table content      (default: True)
        include_comments : include comments           (default: False)
        preserve_spacing : preserve paragraph spacing (default: True)
        page_separator   : separator between sections (default: '')
        encoding         : output encoding            (default: utf-8)

    Returns:
        {
            'text'         : str,
            'paragraphs'   : int,
            'words'        : int,
            'chars'        : int,
            'tables'       : int,
            'headings'     : int,
            'size_original': int,
            'size_text'    : int,
        }
    """
    from docx import Document as DocxDocument
    from docx.text.paragraph import Paragraph as DocxPara
    from docx.table import Table as DocxTable
    from docx.oxml.ns import qn

    # ── Read file bytes ────────────────────────────
    if hasattr(source, "read"):
        file_bytes = source.read()
    else:
        file_bytes = source

    if not file_bytes:
        raise ValueError("Empty file.")

    original_size = len(file_bytes)

    # ── Open document ─────────────────────────────
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"Cannot open Word document: {e}")

    # ── Extract content ───────────────────────────
    lines = []
    para_count = 0
    table_count = 0
    heading_count = 0

    # ── Heading markers ────────────────────────────
    heading_markers = {
        "Heading 1": "=",
        "Heading 2": "-",
        "Heading 3": "~",
        "Heading 4": "+",
        "Title": "#",
        "Subtitle": "*",
    }

    def process_paragraph(para) -> str:
        """Extract text from paragraph with style info."""
        nonlocal para_count, heading_count

        text = para.text.strip()
        if not text:
            return ""

        style_name = para.style.name if para.style else "Normal"

        # ── Heading detection ──────────────────────
        if include_headers and any(h in style_name for h in heading_markers):
            heading_count += 1
            marker = next(
                (v for k, v in heading_markers.items() if k in style_name), ""
            )
            if marker:
                underline = marker * len(text)
                return f"{text}\n{underline}"

        para_count += 1
        return text

    def process_table(table) -> str:
        """Extract text from table as formatted text."""
        nonlocal table_count
        table_count += 1

        rows_text = []
        col_widths = []

        # ── Collect all cell texts ─────────────────
        rows_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = " ".join(
                    p.text.strip() for p in cell.paragraphs if p.text.strip()
                )
                row_data.append(cell_text)
            rows_data.append(row_data)

        if not rows_data:
            return ""

        # ── Calculate column widths ────────────────
        num_cols = max(len(row) for row in rows_data)
        col_widths = [0] * num_cols

        for row in rows_data:
            for i, cell in enumerate(row):
                if i < num_cols:
                    col_widths[i] = max(col_widths[i], len(cell))

        # ── Build table text ───────────────────────
        def format_row(row_data):
            cells = []
            for i in range(num_cols):
                cell = row_data[i] if i < len(row_data) else ""
                cells.append(cell.ljust(col_widths[i]))
            return "| " + " | ".join(cells) + " |"

        separator = "+" + "+".join("-" * (w + 2) for w in col_widths) + "+"

        table_lines = [separator]
        for r_idx, row in enumerate(rows_data):
            table_lines.append(format_row(row))
            if r_idx == 0:
                table_lines.append(separator)  # header separator
        table_lines.append(separator)

        return "\n".join(table_lines)

    # ── Process document elements ──────────────────
    for element in doc.element.body:
        tag = element.tag.split("}")[-1]

        if tag == "p":
            para = DocxPara(element, doc)
            text = process_paragraph(para)
            if text:
                lines.append(text)
                if preserve_spacing:
                    lines.append("")

        elif tag == "tbl":
            if include_tables:
                table = DocxTable(element, doc)
                table_text = process_table(table)
                if table_text:
                    if preserve_spacing:
                        lines.append("")
                    lines.append(table_text)
                    if preserve_spacing:
                        lines.append("")

        elif tag == "sectPr":
            if page_separator:
                lines.append(page_separator)

    # ── Extract comments if requested ─────────────
    if include_comments:
        try:
            comments_part = doc.part.comments_part
            if comments_part:
                lines.append("\n\n--- COMMENTS ---")
                for comment in comments_part.element.findall(".//" + qn("w:comment")):
                    author = comment.get(qn("w:author"), "Unknown")
                    date = comment.get(qn("w:date"), "")[:10]
                    comment_text = "".join(
                        t.text or "" for t in comment.findall(".//" + qn("w:t"))
                    )
                    lines.append(f"[{author} - {date}]: {comment_text}")
        except Exception:
            pass  # Comments not available

    # ── Join all lines ─────────────────────────────
    full_text = "\n".join(lines).strip()

    # ── Remove excessive blank lines ───────────────
    import re

    full_text = re.sub(r"\n{3,}", "\n\n", full_text)

    # ── Stats ──────────────────────────────────────
    word_count = len(full_text.split()) if full_text else 0
    char_count = len(full_text)

    return {
        "text": full_text,
        "paragraphs": para_count,
        "words": word_count,
        "chars": char_count,
        "tables": table_count,
        "headings": heading_count,
        "size_original": original_size,
        "size_original_kb": round(original_size / 1024, 2),
        "size_text": len(full_text.encode(encoding)),
        "size_text_kb": round(len(full_text.encode(encoding)) / 1024, 2),
        "encoding": encoding,
    }


def txt_to_word(
    source,
    filename: str = "document.txt",
    title: str = "",
    font_name: str = "Calibri",
    font_size: int = 11,
    line_spacing: float = 1.15,
    detect_headings: bool = True,
    detect_lists: bool = True,
    detect_tables: bool = True,
    page_size: str = "A4",
    encoding: str = "utf-8",
) -> bytes:
    """
    Convert plain text (.txt) → Word (.docx).

    Args:
        source          : uploaded file object | raw text string | bytes
        filename        : original filename
        title           : document title (metadata)
        font_name       : body font                  (default: Calibri)
        font_size       : body font size pt          (default: 11)
        line_spacing    : line spacing multiplier    (default: 1.15)
        detect_headings : auto-detect headings       (default: True)
        detect_lists    : auto-detect bullet lists   (default: True)
        detect_tables   : auto-detect ASCII tables   (default: True)
        page_size       : A4 | Letter                (default: A4)
        encoding        : input text encoding        (default: utf-8)

    Heading detection rules:
        - ALL CAPS short lines           → Heading 1
        - Lines ending with ===          → Heading 1
        - Lines ending with ---          → Heading 2
        - Lines ending with ~~~          → Heading 3
        - Lines starting with # ## ###  → Heading 1/2/3
        - Short lines followed by ===   → Heading 1
        - Short lines followed by ---   → Heading 2

    List detection rules:
        - Lines starting with - * + •   → Bullet list
        - Lines starting with 1. 2. 3.  → Numbered list

    Table detection rules:
        - Lines starting with |         → Table row
        - Lines starting with +---      → Table separator (skip)

    Returns:
        Raw bytes of the generated .docx file
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    import re

    # ── Read source ───────────────────────────────
    if hasattr(source, "read"):
        raw = source.read()
        if isinstance(raw, bytes):
            raw = raw.decode(encoding, errors="replace")
    elif isinstance(source, bytes):
        raw = source.decode(encoding, errors="replace")
    elif isinstance(source, str):
        raw = source
    else:
        raise ValueError("source must be a string, bytes, or file object.")

    if not raw.strip():
        raise ValueError("Empty input. Nothing to convert.")

    # ── Page size ─────────────────────────────────
    from docx.shared import Inches

    page_sizes = {
        "A4": (Inches(8.27), Inches(11.69)),
        "Letter": (Inches(8.5), Inches(11.0)),
        "Legal": (Inches(8.5), Inches(14.0)),
        "A3": (Inches(11.69), Inches(16.54)),
    }
    if page_size not in page_sizes:
        page_size = "A4"
    page_w, page_h = page_sizes[page_size]

    # ── Create document ───────────────────────────
    doc = Document()

    # ── Page setup ────────────────────────────────
    section = doc.sections[0]
    section.page_width = page_w
    section.page_height = page_h
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ── Document metadata ─────────────────────────
    core_props = doc.core_properties
    core_props.title = title or filename.replace(".txt", "")
    core_props.author = "Text to Word Converter"
    core_props.subject = "Converted from plain text"

    # ── Default style ─────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size)

    from docx.oxml import OxmlElement
    from lxml import etree

    def set_line_spacing(paragraph, spacing=1.15):
        pPr = paragraph._p.get_or_add_pPr()
        spacing_el = OxmlElement("w:spacing")
        spacing_el.set(qn("w:line"), str(int(spacing * 240)))
        spacing_el.set(qn("w:lineRule"), "auto")
        pPr.append(spacing_el)

    # ── Heading styles ─────────────────────────────
    heading_colors = {
        1: RGBColor(0x2E, 0x75, 0xB6),  # blue
        2: RGBColor(0x2E, 0x75, 0xB6),  # blue
        3: RGBColor(0x1F, 0x4E, 0x79),  # dark blue
    }

    def add_heading(text, level=1):
        p = doc.add_heading(text.strip(), level=level)
        for run in p.runs:
            run.font.color.rgb = heading_colors.get(level, RGBColor(0, 0, 0))
        return p

    def add_paragraph(text, bold=False, italic=False, align=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        run.italic = italic
        if align:
            p.alignment = align
        set_line_spacing(p, line_spacing)
        return p

    def add_bullet(text, numbered=False, num=1):
        if numbered:
            p = doc.add_paragraph(
                text.strip(),
                style="List Number",
            )
        else:
            p = doc.add_paragraph(
                text.strip(),
                style="List Bullet",
            )
        for run in p.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
        return p

    def add_table_from_rows(table_rows):
        """Convert ASCII table rows → Word table."""
        # Parse rows
        data = []
        for row in table_rows:
            row = row.strip()
            if row.startswith("+") or set(row.strip()) <= {"+", "-", "="}:
                continue  # skip separator rows
            cells = [c.strip() for c in row.strip("|").split("|")]
            if cells:
                data.append(cells)

        if not data:
            return

        num_cols = max(len(row) for row in data)
        num_rows = len(data)

        tbl = doc.add_table(rows=num_rows, cols=num_cols)
        tbl.style = "Table Grid"

        from docx.shared import RGBColor
        from docx.oxml import OxmlElement

        for r_idx, row_data in enumerate(data):
            row = tbl.rows[r_idx]
            for c_idx in range(num_cols):
                cell = row.cells[c_idx]
                text = row_data[c_idx] if c_idx < len(row_data) else ""
                cell.text = text

                # Style header row
                if r_idx == 0:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        # Blue background
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        shd = OxmlElement("w:shd")
                        shd.set(qn("w:fill"), "2E75B6")
                        shd.set(qn("w:color"), "auto")
                        shd.set(qn("w:val"), "clear")
                        tcPr.append(shd)
                else:
                    # Alternating rows
                    if r_idx % 2 == 0:
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        shd = OxmlElement("w:shd")
                        shd.set(qn("w:fill"), "DCE6F1")
                        shd.set(qn("w:color"), "auto")
                        shd.set(qn("w:val"), "clear")
                        tcPr.append(shd)

        doc.add_paragraph()  # spacing after table

    # ── Parse lines ───────────────────────────────
    lines = raw.splitlines()
    i = 0
    total_lines = len(lines)

    # Add title if provided
    if title:
        add_heading(title, level=1)
        doc.add_paragraph()

    while i < total_lines:
        line = lines[i]
        stripped = line.strip()

        # ── Empty line ─────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── Markdown headings # ## ### ─────────────
        if detect_headings and stripped.startswith("#"):
            level = min(len(stripped) - len(stripped.lstrip("#")), 3)
            text = stripped.lstrip("#").strip()
            if text:
                add_heading(text, level=level)
                i += 1
                continue

        # ── Underline-style headings ──────────────
        # Check next line for === or ---
        if detect_headings and i + 1 < total_lines:
            next_line = lines[i + 1].strip()
            if next_line and len(next_line) >= 3:
                if set(next_line) == {"="}:
                    add_heading(stripped, level=1)
                    i += 2
                    continue
                if set(next_line) == {"-"}:
                    add_heading(stripped, level=2)
                    i += 2
                    continue
                if set(next_line) == {"~"}:
                    add_heading(stripped, level=3)
                    i += 2
                    continue

        # ── ALL CAPS heading ──────────────────────
        if detect_headings:
            if (
                stripped.isupper()
                and len(stripped) > 3
                and len(stripped) < 80
                and not stripped.startswith("|")
            ):
                add_heading(stripped, level=1)
                i += 1
                continue

        # ── Inline === / --- heading markers ───────
        if detect_headings:
            if re.match(r"^={3,}$", stripped):
                i += 1
                continue
            if re.match(r"^-{3,}$", stripped):
                i += 1
                continue

        # ── ASCII table ───────────────────────────
        if detect_tables and (stripped.startswith("|") or stripped.startswith("+")):
            # Collect all consecutive table lines
            table_rows = []
            while i < total_lines and (
                lines[i].strip().startswith("|") or lines[i].strip().startswith("+")
            ):
                table_rows.append(lines[i])
                i += 1

            add_table_from_rows(table_rows)
            continue

        # ── Bullet list ───────────────────────────
        if detect_lists and re.match(r"^[-*+•]\s+", stripped):
            text = re.sub(r"^[-*+•]\s+", "", stripped)
            add_bullet(text, numbered=False)
            i += 1
            continue

        # ── Numbered list ─────────────────────────
        if detect_lists and re.match(r"^\d+[.)]\s+", stripped):
            text = re.sub(r"^\d+[.)]\s+", "", stripped)
            add_bullet(text, numbered=True)
            i += 1
            continue

        # ── Horizontal rule ───────────────────────
        if re.match(r"^[=\-_*]{3,}$", stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            from docx.oxml import OxmlElement

            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "2E75B6")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # ── Regular paragraph ─────────────────────
        add_paragraph(stripped)
        i += 1

    # ── Serialize ─────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def merge_docx(
    sources: list,
    page_break: bool = True,
    add_toc: bool = False,
    preserve_styles: bool = True,
    separator_text: str = "",
) -> bytes:
    """
    Merge multiple Word (.docx) files into one document.

    Args:
        sources         : list of file objects OR bytes OR paths
        page_break      : add page break between documents   (default: True)
        add_toc         : add table of contents              (default: False)
        preserve_styles : preserve each doc's styles         (default: True)
        separator_text  : text to insert between documents   (default: '')

    Returns:
        Raw bytes of the merged .docx file
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import copy

    if not sources:
        raise ValueError("No files provided to merge.")

    if len(sources) < 2:
        raise ValueError("At least 2 files are required to merge.")

    if len(sources) > 20:
        raise ValueError("Maximum 20 files can be merged at once.")

    # ── Read all documents ─────────────────────────
    documents = []
    for i, source in enumerate(sources):
        try:
            if hasattr(source, "read"):
                data = source.read()
            elif isinstance(source, bytes):
                data = source
            else:
                raise ValueError(f"Invalid source type at index {i}.")

            doc = Document(io.BytesIO(data))
            documents.append(
                {
                    "doc": doc,
                    "filename": getattr(source, "name", f"document_{i+1}.docx"),
                    "index": i,
                }
            )
        except Exception as e:
            raise ValueError(f"Cannot open file {i+1}: {e}")

    # ── Use first document as base ─────────────────
    merged = documents[0]["doc"]
    base_name = documents[0]["filename"]

    def add_page_break(doc):
        """Add page break to document."""
        para = doc.add_paragraph()
        run = para.add_run()
        run.add_break(
            __import__("docx.enum.text", fromlist=["WD_BREAK_TYPE"]).WD_BREAK_TYPE.PAGE
        )

    def add_separator(doc, text: str):
        """Add separator text between documents."""
        if not text:
            return
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add horizontal line above
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        top = OxmlElement("w:top")
        top.set(qn("w:val"), "single")
        top.set(qn("w:sz"), "6")
        top.set(qn("w:space"), "1")
        top.set(qn("w:color"), "2E75B6")
        pBdr.append(top)
        pPr.append(pBdr)

    def copy_element(element):
        """Deep copy an XML element."""
        return copy.deepcopy(element)

    def append_document(base_doc, src_doc, filename: str, index: int):
        """Append all content from src_doc into base_doc."""

        # ── Add separator if specified ─────────────
        if separator_text:
            add_separator(
                base_doc,
                separator_text.replace("{n}", str(index + 1)).replace(
                    "{filename}", filename
                ),
            )

        # ── Add page break between docs ────────────
        if page_break and index > 0:
            # Add page break to last paragraph of base
            if base_doc.paragraphs:
                last_para = base_doc.paragraphs[-1]
                run = last_para.add_run()
                br = OxmlElement("w:br")
                br.set(qn("w:type"), "page")
                run._r.append(br)
            else:
                add_page_break(base_doc)

        # ── Copy styles from source ────────────────
        if preserve_styles:
            _copy_styles(base_doc, src_doc)

        # ── Copy images / relationships ────────────
        _copy_relationships(base_doc, src_doc)

        # ── Copy body elements ─────────────────────
        for element in src_doc.element.body:
            tag = element.tag.split("}")[-1]

            # Skip section properties (sectPr) — keep base doc's layout
            if tag == "sectPr":
                continue

            # Deep copy element and append to base body
            new_element = copy_element(element)
            base_doc.element.body.append(new_element)

    def _copy_styles(base_doc, src_doc):
        """Copy missing styles from src_doc to base_doc."""
        base_style_names = {s.name for s in base_doc.styles}

        for style in src_doc.styles:
            if style.name not in base_style_names:
                try:
                    new_style = copy_element(style.element)
                    base_doc.styles.element.append(new_style)
                except Exception:
                    pass  # Skip problematic styles

    def _copy_relationships(base_doc, src_doc):
        """Copy image relationships from src to base."""
        try:
            src_part = src_doc.part
            base_part = base_doc.part

            for rel in src_part.rels.values():
                if "image" in rel.reltype:
                    try:
                        # Get image blob from source
                        img_part = rel.target_part
                        img_bytes = img_part.blob
                        img_ct = img_part.content_type

                        # Add to base document
                        base_part.relate_to(
                            img_part,
                            rel.reltype,
                        )
                    except Exception:
                        pass  # Skip problematic images
        except Exception:
            pass

    # ── Merge all documents ────────────────────────
    for i in range(1, len(documents)):
        append_document(
            merged,
            documents[i]["doc"],
            documents[i]["filename"],
            i,
        )

    # ── Add TOC if requested ───────────────────────
    if add_toc:
        _add_toc(merged)

    # ── Update metadata ────────────────────────────
    merged.core_properties.title = f"Merged Document ({len(documents)} files)"
    merged.core_properties.author = "Merge DOCX Service"
    merged.core_properties.subject = ", ".join(d["filename"] for d in documents)

    # ── Serialize ─────────────────────────────────
    buffer = io.BytesIO()
    merged.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _add_toc(doc):
    """
    Add a basic Table of Contents field to the beginning of the document.
    Note: TOC requires Word to refresh on open (F9).
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Insert TOC at beginning
    toc_para = doc.paragraphs[0]._p if doc.paragraphs else None

    # Add TOC heading
    toc_heading = doc.add_heading("Table of Contents", level=1)

    # Move heading to beginning
    doc.element.body.insert(0, toc_heading._p)

    # Add TOC field
    para = OxmlElement("w:p")
    run = OxmlElement("w:r")
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    run.append(fldChar)
    para.append(run)

    run2 = OxmlElement("w:r")
    instrText = OxmlElement("w:instrText")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2.append(instrText)
    para.append(run2)

    run3 = OxmlElement("w:r")
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run3.append(fldChar2)
    para.append(run3)

    doc.element.body.insert(1, para)


def split_docx(
    source,
    split_by: str = "page",
    pages: list = None,
    heading_level: int = 1,
    chunk_size: int = 1,
) -> list[dict]:
    """
    Split a Word (.docx) file into multiple documents.

    Args:
        source        : uploaded file object OR raw bytes
        split_by      : 'page'    → split by page break
                        'heading' → split at each heading
                        'chunk'   → split every N paragraphs
                        'range'   → extract specific page ranges
        pages         : list of page ranges for 'range' mode
                        e.g. [[1,3], [4,6]] or [1, 3, 5]
        heading_level : heading level to split on (1-6)  (default: 1)
        chunk_size    : paragraphs per chunk             (default: 1)

    Returns:
        list of {
            'index'    : int,
            'filename' : str,
            'bytes'    : bytes,
            'title'    : str,
            'paragraphs': int,
            'size_kb'  : float,
        }
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph as DocxPara
    from docx.table import Table as DocxTable
    import copy
    import re

    # ── Read source ───────────────────────────────
    if hasattr(source, "read"):
        file_bytes = source.read()
        filename = getattr(source, "name", "document.docx")
    elif isinstance(source, bytes):
        file_bytes = source
        filename = "document.docx"
    else:
        raise ValueError("Invalid source.")

    if not file_bytes:
        raise ValueError("Empty file.")

    base_name = (
        filename.replace(".docx", "")
        .replace(".doc", "")
        .replace(".DOCX", "")
        .replace(".DOC", "")
    )

    # ── Open document ─────────────────────────────
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"Cannot open Word document: {e}")

    # ── Validate split_by ─────────────────────────
    valid_split = ("page", "heading", "chunk", "range")
    if split_by not in valid_split:
        raise ValueError(f"split_by must be one of: {valid_split}")

    # ── Helper: create new doc with same styles ────
    def create_new_doc() -> Document:
        new_doc = Document()

        # Copy styles
        for style in doc.styles:
            try:
                if style.name not in {s.name for s in new_doc.styles}:
                    new_doc.styles.element.append(copy.deepcopy(style.element))
            except Exception:
                pass

        # Copy page layout from source
        try:
            src_section = doc.sections[0]
            dest_section = new_doc.sections[0]
            dest_section.page_width = src_section.page_width
            dest_section.page_height = src_section.page_height
            dest_section.left_margin = src_section.left_margin
            dest_section.right_margin = src_section.right_margin
            dest_section.top_margin = src_section.top_margin
            dest_section.bottom_margin = src_section.bottom_margin
        except Exception:
            pass

        # Remove default empty paragraph
        for p in new_doc.paragraphs:
            p._element.getparent().remove(p._element)

        return new_doc

    def doc_to_bytes(d: Document) -> bytes:
        buf = io.BytesIO()
        d.save(buf)
        buf.seek(0)
        return buf.read()

    def get_para_title(para) -> str:
        return para.text.strip()[:80] if para.text.strip() else ""

    def append_element(new_doc, element):
        new_doc.element.body.append(copy.deepcopy(element))

    def count_paragraphs(d: Document) -> int:
        return len([p for p in d.paragraphs if p.text.strip()])

    def is_page_break(element) -> bool:
        """Check if element contains a page break."""
        xml = element.xml if hasattr(element, "xml") else ""
        return (
            f'w:type="page"' in xml
            or f"w:type='page'" in xml
            or "w:lastRenderedPageBreak" in xml
        )

    def is_heading(para, level=1) -> bool:
        """Check if paragraph is a heading of given level."""
        style_name = para.style.name if para.style else ""
        return f"Heading {level}" in style_name

    def finalize_part(new_doc, index, title, base_name) -> dict:
        """Convert doc to bytes and build result dict."""
        raw = doc_to_bytes(new_doc)
        safe_t = re.sub(r"[^\w\s-]", "", title)[:40].strip()
        fname = f"{base_name}_part{index}"
        if safe_t:
            fname += f"_{safe_t}"
        fname += ".docx"

        return {
            "index": index,
            "filename": fname,
            "bytes": raw,
            "title": title or f"Part {index}",
            "paragraphs": count_paragraphs(new_doc),
            "size_kb": round(len(raw) / 1024, 2),
        }

    results = []

    # ────────────────────────────────────────────────
    # MODE 1: Split by page break
    # ────────────────────────────────────────────────
    if split_by == "page":
        current_doc = create_new_doc()
        current_title = ""
        part_index = 1

        for element in doc.element.body:
            tag = element.tag.split("}")[-1]

            if tag == "sectPr":
                continue

            # Check for page break
            if is_page_break(element):
                # Save current part (without the page break element)
                if (
                    count_paragraphs(current_doc) > 0
                    or len(current_doc.element.body) > 0
                ):
                    results.append(
                        finalize_part(current_doc, part_index, current_title, base_name)
                    )
                    part_index += 1
                    current_doc = create_new_doc()
                    current_title = ""
                continue

            # Set title from first heading or paragraph
            if not current_title and tag == "p":
                para = DocxPara(element, doc)
                if para.text.strip():
                    current_title = get_para_title(para)

            append_element(current_doc, element)

        # Save last part
        if len(current_doc.element.body) > 0:
            results.append(
                finalize_part(current_doc, part_index, current_title, base_name)
            )

    # ────────────────────────────────────────────────
    # MODE 2: Split by heading
    # ────────────────────────────────────────────────
    elif split_by == "heading":
        current_doc = create_new_doc()
        current_title = f"Part 1"
        part_index = 1
        found_first = False

        for element in doc.element.body:
            tag = element.tag.split("}")[-1]

            if tag == "sectPr":
                continue

            if tag == "p":
                para = DocxPara(element, doc)

                if is_heading(para, heading_level):
                    # Save previous part if it has content
                    if found_first and count_paragraphs(current_doc) > 0:
                        results.append(
                            finalize_part(
                                current_doc, part_index, current_title, base_name
                            )
                        )
                        part_index += 1
                        current_doc = create_new_doc()

                    current_title = get_para_title(para)
                    found_first = True

            append_element(current_doc, element)

        # Save last part
        if count_paragraphs(current_doc) > 0:
            results.append(
                finalize_part(current_doc, part_index, current_title, base_name)
            )

        # If no headings found
        if not results:
            raise ValueError(
                f"No Heading {heading_level} found in document. "
                f"Try a different heading_level or split_by mode."
            )

    # ────────────────────────────────────────────────
    # MODE 3: Split by chunk (every N paragraphs)
    # ────────────────────────────────────────────────
    elif split_by == "chunk":
        if chunk_size < 1:
            raise ValueError("chunk_size must be at least 1.")

        current_doc = create_new_doc()
        current_title = ""
        part_index = 1
        para_count = 0

        for element in doc.element.body:
            tag = element.tag.split("}")[-1]

            if tag == "sectPr":
                continue

            append_element(current_doc, element)

            if tag == "p":
                para = DocxPara(element, doc)
                if para.text.strip():
                    if not current_title:
                        current_title = get_para_title(para)
                    para_count += 1

                if para_count >= chunk_size:
                    results.append(
                        finalize_part(current_doc, part_index, current_title, base_name)
                    )
                    part_index += 1
                    current_doc = create_new_doc()
                    current_title = ""
                    para_count = 0

        # Save remaining
        if count_paragraphs(current_doc) > 0:
            results.append(
                finalize_part(current_doc, part_index, current_title, base_name)
            )

    # ────────────────────────────────────────────────
    # MODE 4: Split by range (extract specific pages)
    # ────────────────────────────────────────────────
    elif split_by == "range":
        if not pages:
            raise ValueError(
                '"pages" is required for range mode. '
                "e.g. [[1,3], [4,6]] or [1, 3, 5]"
            )

        # ── Collect all elements with page numbers ─
        all_elements = []
        current_page = 1

        for element in doc.element.body:
            tag = element.tag.split("}")[-1]
            if tag == "sectPr":
                continue

            all_elements.append(
                {
                    "element": element,
                    "page": current_page,
                }
            )

            if is_page_break(element):
                current_page += 1

        total_pages = current_page

        # ── Normalize ranges ───────────────────────
        normalized = []
        for r in pages:
            if isinstance(r, (int, float)):
                p = int(r)
                normalized.append((p, p))
            elif isinstance(r, (list, tuple)) and len(r) == 2:
                normalized.append((int(r[0]), int(r[1])))
            else:
                raise ValueError(
                    f"Invalid range: {r}. " f"Use integers or [start, end] pairs."
                )

        # ── Extract each range ─────────────────────
        for range_idx, (start, end) in enumerate(normalized, start=1):
            if start < 1:
                start = 1
            if end > total_pages:
                end = total_pages
            if start > end:
                raise ValueError(
                    f"Invalid range: [{start}, {end}]. " f"Start must be <= end."
                )

            range_doc = create_new_doc()
            range_title = f"Pages {start}-{end}"

            for item in all_elements:
                if start <= item["page"] <= end:
                    append_element(range_doc, item["element"])

            if count_paragraphs(range_doc) > 0 or len(range_doc.element.body) > 0:
                results.append(
                    finalize_part(range_doc, range_idx, range_title, base_name)
                )

    if not results:
        raise ValueError(
            "No content found to split. "
            "The document may be empty or the split criteria found nothing."
        )

    return results


def word_to_markdown(
    source,
    filename: str = "document.docx",
    include_tables: bool = True,
    include_images: bool = False,
    include_toc: bool = False,
    heading_style: str = "atx",
    code_block_style: str = "fenced",
    preserve_emphasis: bool = True,
    encoding: str = "utf-8",
) -> dict:
    """
    Convert Word (.docx) → Markdown text.

    Args:
        source            : uploaded file object OR raw bytes
        filename          : original filename
        include_tables    : convert tables to markdown     (default: True)
        include_images    : include image placeholders     (default: False)
        include_toc       : generate table of contents    (default: False)
        heading_style     : 'atx'  → # Heading
                            'setext' → Heading\n======    (default: atx)
        code_block_style  : 'fenced'  → ```code```
                            'indented' → 4-space indent   (default: fenced)
        preserve_emphasis : preserve bold/italic          (default: True)
        encoding          : output encoding               (default: utf-8)

    Returns:
        {
            'markdown'    : str,
            'headings'    : list,
            'word_count'  : int,
            'char_count'  : int,
            'table_count' : int,
            'image_count' : int,
            'toc'         : str,
        }
    """
    from docx import Document as DocxDocument
    from docx.text.paragraph import Paragraph as DocxPara
    from docx.table import Table as DocxTable
    from docx.oxml.ns import qn
    import re

    # ── Read source ───────────────────────────────
    if hasattr(source, "read"):
        file_bytes = source.read()
    elif isinstance(source, bytes):
        file_bytes = source
    else:
        raise ValueError("source must be a file object or bytes.")

    if not file_bytes:
        raise ValueError("Empty file.")

    # ── Open document ─────────────────────────────
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"Cannot open Word document: {e}")

    # ── Tracking ──────────────────────────────────
    lines = []
    headings = []
    table_count = 0
    image_count = 0

    # ── Helpers ────────────────────────────────────
    def escape_md(text: str) -> str:
        """Escape markdown special characters in plain text."""
        chars = r"\`*_{}[]()#+-.!"
        for ch in chars:
            text = text.replace(ch, "\\" + ch)
        return text

    def runs_to_md(para) -> str:
        """Convert paragraph runs to markdown with emphasis."""
        if not preserve_emphasis:
            return para.text

        parts = []
        for run in para.runs:
            text = run.text
            if not text:
                continue

            # Bold + Italic
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"

            # Strikethrough
            if run.font.strike:
                text = f"~~{text}~~"

            # Underline (no markdown equivalent — use HTML)
            if run.underline:
                text = f"<u>{text}</u>"

            # Inline code (Courier/monospace font)
            try:
                font_name = run.font.name or ""
                if "courier" in font_name.lower() or "mono" in font_name.lower():
                    text = f"`{text}`"
            except Exception:
                pass

            # Hyperlink
            try:
                if run._element.getparent().tag.endswith("}hyperlink"):
                    rId = run._element.getparent().get(qn("r:id"), "")
                    href = ""
                    try:
                        href = doc.part.rels[rId].target_ref
                    except Exception:
                        pass
                    if href:
                        text = f"[{text}]({href})"
            except Exception:
                pass

            parts.append(text)

        return "".join(parts)

    def process_heading(para, level: int) -> str:
        """Convert heading paragraph to markdown."""
        text = para.text.strip()
        slug = re.sub(r"[^\w\s-]", "", text.lower())
        slug = re.sub(r"[\s]+", "-", slug).strip("-")

        headings.append(
            {
                "level": level,
                "text": text,
                "slug": slug,
            }
        )

        if heading_style == "setext" and level <= 2:
            underline = ("=" if level == 1 else "-") * max(len(text), 3)
            return f"{text}\n{underline}"
        else:
            return "#" * level + " " + text

    def process_table(tbl) -> str:
        """Convert table to markdown."""
        nonlocal table_count
        table_count += 1

        rows_data = []
        for row in tbl.rows:
            row_data = []
            for cell in row.cells:
                cell_text = " ".join(
                    p.text.strip() for p in cell.paragraphs if p.text.strip()
                )
                # Escape pipes inside cells
                cell_text = cell_text.replace("|", "\\|")
                row_data.append(cell_text)
            rows_data.append(row_data)

        if not rows_data:
            return ""

        # Normalize column count
        num_cols = max(len(row) for row in rows_data)
        rows_data = [row + [""] * (num_cols - len(row)) for row in rows_data]

        # Calculate column widths
        col_widths = [
            max(len(rows_data[r][c]) for r in range(len(rows_data)))
            for c in range(num_cols)
        ]
        col_widths = [max(w, 3) for w in col_widths]

        def format_row(row):
            cells = [row[c].ljust(col_widths[c]) for c in range(num_cols)]
            return "| " + " | ".join(cells) + " |"

        # Header row
        md_lines = [format_row(rows_data[0])]

        # Separator row
        sep = "| " + " | ".join("-" * w for w in col_widths) + " |"
        md_lines.append(sep)

        # Data rows
        for row in rows_data[1:]:
            md_lines.append(format_row(row))

        return "\n".join(md_lines)

    def process_list_paragraph(para, name: str) -> str:
        """Convert list paragraph to markdown."""
        text = runs_to_md(para)
        level = 0

        # Detect indent level
        try:
            indent = para.paragraph_format.left_indent
            if indent:
                level = min(int(indent.pt // 18), 4)
        except Exception:
            pass

        prefix = "  " * level

        if "bullet" in name:
            return f"{prefix}- {text}"
        elif "number" in name:
            return f"{prefix}1. {text}"

        return f"- {text}"

    def generate_toc(headings: list) -> str:
        """Generate table of contents from headings."""
        if not headings:
            return ""

        toc_lines = ["## Table of Contents\n"]
        for h in headings:
            indent = "  " * (h["level"] - 1)
            toc_lines.append(f'{indent}- [{h["text"]}](#{h["slug"]})')
        return "\n".join(toc_lines)

    # ── Process document elements ──────────────────
    prev_was_list = False

    for element in doc.element.body:
        tag = element.tag.split("}")[-1]

        # ── Paragraph ─────────────────────────────
        if tag == "p":
            para = DocxPara(element, doc)
            text = para.text.strip()
            style_name = (para.style.name or "Normal").lower()

            # ── Page break ─────────────────────────
            if any(
                br.get(qn("w:type")) == "page"
                for br in element.findall(".//" + qn("w:br"))
            ):
                lines.append("\n---\n")
                continue

            # ── Headings ───────────────────────────
            if "heading 1" in style_name or style_name == "title":
                lines.append(process_heading(para, 1))
                lines.append("")
                prev_was_list = False
                continue

            if "heading 2" in style_name or style_name == "subtitle":
                lines.append(process_heading(para, 2))
                lines.append("")
                prev_was_list = False
                continue

            if "heading 3" in style_name:
                lines.append(process_heading(para, 3))
                lines.append("")
                prev_was_list = False
                continue

            if "heading 4" in style_name:
                lines.append(process_heading(para, 4))
                lines.append("")
                prev_was_list = False
                continue

            if "heading 5" in style_name:
                lines.append(process_heading(para, 5))
                lines.append("")
                prev_was_list = False
                continue

            if "heading 6" in style_name:
                lines.append(process_heading(para, 6))
                lines.append("")
                prev_was_list = False
                continue

            # ── Empty paragraph ────────────────────
            if not text:
                if prev_was_list:
                    lines.append("")
                    prev_was_list = False
                else:
                    lines.append("")
                continue

            # ── List items ─────────────────────────
            if "list bullet" in style_name or "list number" in style_name:
                lines.append(process_list_paragraph(para, style_name))
                prev_was_list = True
                continue

            # ── Code / preformatted ────────────────
            if "code" in style_name or "preformat" in style_name:
                if code_block_style == "fenced":
                    lines.append(f"```\n{text}\n```")
                else:
                    lines.append("\n".join("    " + l for l in text.splitlines()))
                lines.append("")
                prev_was_list = False
                continue

            # ── Quote / blockquote ─────────────────
            if "quote" in style_name or "block text" in style_name:
                quoted = "\n".join(f"> {l}" for l in text.splitlines())
                lines.append(quoted)
                lines.append("")
                prev_was_list = False
                continue

            # ── Horizontal rule ────────────────────
            if re.match(r"^[-_*]{3,}$", text):
                lines.append("---")
                lines.append("")
                prev_was_list = False
                continue

            # ── Image check ────────────────────────
            if include_images:
                drawing_elements = element.findall(".//" + qn("w:drawing"))
                if drawing_elements:
                    image_count += 1
                    lines.append(f"![Image {image_count}](image_{image_count}.png)")
                    lines.append("")
                    continue

            # ── Regular paragraph ──────────────────
            md_text = runs_to_md(para)
            if md_text.strip():
                lines.append(md_text)
                lines.append("")
            prev_was_list = False

        # ── Table ──────────────────────────────────
        elif tag == "tbl":
            if include_tables:
                try:
                    tbl = DocxTable(element, doc)
                    tbl_md = process_table(tbl)
                    if tbl_md:
                        lines.append("")
                        lines.append(tbl_md)
                        lines.append("")
                except Exception as e:
                    lines.append(f"<!-- Table error: {e} -->")
                    lines.append("")
            prev_was_list = False

        elif tag == "sectPr":
            pass

    # ── Join and clean up ──────────────────────────
    markdown = "\n".join(lines)

    # Remove excessive blank lines
    import re as _re

    markdown = _re.sub(r"\n{3,}", "\n\n", markdown)
    markdown = markdown.strip()

    # ── Generate TOC ──────────────────────────────
    toc = ""
    if include_toc and headings:
        toc = generate_toc(headings)
        markdown = toc + "\n\n---\n\n" + markdown

    # ── Stats ─────────────────────────────────────
    word_count = len(markdown.split())
    char_count = len(markdown)

    return {
        "markdown": markdown,
        "headings": headings,
        "word_count": word_count,
        "char_count": char_count,
        "table_count": table_count,
        "image_count": image_count,
        "toc": toc,
        "encoding": encoding,
    }


def markdown_to_word(
    source,
    filename      : str   = 'document.md',
    title         : str   = '',
    font_name     : str   = 'Calibri',
    font_size     : int   = 11,
    line_spacing  : float = 1.15,
    page_size     : str   = 'A4',
    encoding      : str   = 'utf-8',
) -> bytes:
    """
    Convert Markdown (.md) → Word (.docx).

    Handles:
        - Headings      # ## ### #### ##### ######
        - Setext        Title\n=== / Section\n---
        - Bold          **text** / __text__
        - Italic        *text* / _text_
        - Bold+Italic   ***text***
        - Strikethrough ~~text~~
        - Inline code   `code`
        - Code blocks   ```lang\ncode\n```
        - Blockquotes   > text
        - Bullet lists  - * +
        - Numbered lists 1. 2. 3.
        - Tables        | col | col |
        - Horizontal rules --- *** ___
        - Links         [text](url)
        - Images        ![alt](url)
        - HTML tags     <u> <br> <mark>
        - Nested lists
        - Task lists    - [ ] / - [x]

    Args:
        source       : file object | raw markdown string | bytes
        filename     : original filename
        title        : document title (metadata)
        font_name    : body font                  (default: Calibri)
        font_size    : body font size pt          (default: 11)
        line_spacing : line spacing multiplier    (default: 1.15)
        page_size    : A4 | Letter | Legal | A3  (default: A4)
        encoding     : input text encoding       (default: utf-8)

    Returns:
        Raw bytes of the generated .docx file
    """
    from docx                    import Document
    from docx.shared             import Pt, Inches, RGBColor
    from docx.oxml.ns            import qn
    from docx.oxml               import OxmlElement
    from docx.enum.text          import WD_ALIGN_PARAGRAPH
    import re

    # ── Read source ───────────────────────────────
    if hasattr(source, 'read'):
        raw = source.read()
        if isinstance(raw, bytes):
            raw = raw.decode(encoding, errors='replace')
    elif isinstance(source, bytes):
        raw = source.decode(encoding, errors='replace')
    elif isinstance(source, str):
        raw = source
    else:
        raise ValueError('source must be a string, bytes, or file object.')

    if not raw.strip():
        raise ValueError('Empty input.')

    # ── Page sizes ────────────────────────────────
    page_sizes = {
        'A4'    : (Inches(8.27),  Inches(11.69)),
        'Letter': (Inches(8.5),   Inches(11.0) ),
        'Legal' : (Inches(8.5),   Inches(14.0) ),
        'A3'    : (Inches(11.69), Inches(16.54)),
    }
    if page_size not in page_sizes:
        page_size = 'A4'
    page_w, page_h = page_sizes[page_size]

    # ── Create document ───────────────────────────
    doc     = Document()
    section = doc.sections[0]
    section.page_width    = page_w
    section.page_height   = page_h
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)

    # ── Metadata ──────────────────────────────────
    doc.core_properties.title   = title or filename.replace('.md', '')
    doc.core_properties.author  = 'Markdown to Word Converter'
    doc.core_properties.subject = 'Converted from Markdown'

    # ── Default style ─────────────────────────────
    normal_style          = doc.styles['Normal']
    normal_style.font.name = font_name
    normal_style.font.size = Pt(font_size)

    # ── Heading colors ─────────────────────────────
    heading_colors = {
        1: RGBColor(0x2E, 0x75, 0xB6),
        2: RGBColor(0x2E, 0x75, 0xB6),
        3: RGBColor(0x1F, 0x4E, 0x79),
        4: RGBColor(0x1F, 0x4E, 0x79),
        5: RGBColor(0x40, 0x40, 0x40),
        6: RGBColor(0x40, 0x40, 0x40),
    }

    # ── Line spacing helper ────────────────────────
    def set_spacing(para, spacing=None):
        spacing = spacing or line_spacing
        pPr     = para._p.get_or_add_pPr()
        sp      = OxmlElement('w:spacing')
        sp.set(qn('w:line'),     str(int(spacing * 240)))
        sp.set(qn('w:lineRule'), 'auto')
        pPr.append(sp)

    # ── Add heading ────────────────────────────────
    def add_heading(text: str, level: int):
        text  = text.strip()
        para  = doc.add_heading('', level=min(level, 6))
        run   = para.add_run(apply_inline(text, para))
        color = heading_colors.get(level, RGBColor(0, 0, 0))
        for r in para.runs:
            r.font.color.rgb = color
        return para

    # ── Add paragraph ──────────────────────────────
    def add_paragraph(text: str, style: str = None, align=None) -> 'Paragraph':
        para = doc.add_paragraph(style=style) if style else doc.add_paragraph()
        apply_inline_to_para(text.strip(), para)
        if align:
            para.alignment = align
        set_spacing(para)
        return para

    # ── Inline markdown parser ─────────────────────
    def apply_inline(text: str, para=None) -> str:
        """Parse inline markdown — returns plain text (side-effect: adds runs)."""
        if para is None:
            return text
        apply_inline_to_para(text, para)
        return ''

    def apply_inline_to_para(text: str, para):
        """Parse inline markdown and add styled runs to paragraph."""

        # Pattern: bold+italic, bold, italic, strikethrough,
        #          inline code, link, image, underline HTML
        pattern = re.compile(
            r'(\*\*\*(.+?)\*\*\*)'           # bold+italic ***
            r'|(__(.+?)__)'                   # bold __
            r'|(\*\*(.+?)\*\*)'              # bold **
            r'|(_(.+?)_)'                     # italic _
            r'|(\*(.+?)\*)'                   # italic *
            r'|(~~(.+?)~~)'                   # strikethrough
            r'|(`(.+?)`)'                     # inline code
            r'|(\[(.+?)\]\((.+?)\))'         # link
            r'|(!?\[(.+?)\]\((.+?)\))'       # image
            r'|(<u>(.+?)<\/u>)'              # underline HTML
            r'|(<mark>(.+?)<\/mark>)'        # highlight HTML
            r'|(<br\s*\/?>)',                 # line break
            re.DOTALL
        )

        last_end = 0
        for m in pattern.finditer(text):
            # Add plain text before this match
            if m.start() > last_end:
                plain = text[last_end:m.start()]
                if plain:
                    run = para.add_run(plain)
                    run.font.name = font_name
                    run.font.size = Pt(font_size)

            matched = m.group(0)

            # ── Bold + Italic ***text*** ───────────
            if m.group(1):
                run       = para.add_run(m.group(2))
                run.bold  = True
                run.italic= True
                run.font.name = font_name
                run.font.size = Pt(font_size)

            # ── Bold __text__ ──────────────────────
            elif m.group(3):
                run      = para.add_run(m.group(4))
                run.bold = True
                run.font.name = font_name
                run.font.size = Pt(font_size)

            # ── Bold **text** ──────────────────────
            elif m.group(5):
                run      = para.add_run(m.group(6))
                run.bold = True
                run.font.name = font_name
                run.font.size = Pt(font_size)

            # ── Italic _text_ ──────────────────────
            elif m.group(7):
                run        = para.add_run(m.group(8))
                run.italic = True
                run.font.name = font_name
                run.font.size = Pt(font_size)

            # ── Italic *text* ──────────────────────
            elif m.group(9):
                run        = para.add_run(m.group(10))
                run.italic = True
                run.font.name = font_name
                run.font.size = Pt(font_size)

            # ── Strikethrough ~~text~~ ─────────────
            elif m.group(11):
                run            = para.add_run(m.group(12))
                run.font.strike= True
                run.font.name  = font_name
                run.font.size  = Pt(font_size)

            # ── Inline code `code` ─────────────────
            elif m.group(13):
                run           = para.add_run(m.group(14))
                run.font.name = 'Courier New'
                run.font.size = Pt(font_size - 1)
                run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

            # ── Link [text](url) ───────────────────
            elif m.group(15):
                link_text = m.group(16)
                link_url  = m.group(17)
                run       = para.add_run(link_text)
                run.font.name  = font_name
                run.font.size  = Pt(font_size)
                run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
                run.font.underline = True

            # ── Image ![alt](url) ──────────────────
            elif m.group(18):
                alt_text = m.group(19)
                run      = para.add_run(f'[Image: {alt_text}]')
                run.font.name  = font_name
                run.font.size  = Pt(font_size)
                run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
                run.italic = True

            # ── Underline <u>text</u> ──────────────
            elif m.group(21):
                run               = para.add_run(m.group(22))
                run.font.underline = True
                run.font.name     = font_name
                run.font.size     = Pt(font_size)

            # ── Highlight <mark>text</mark> ────────
            elif m.group(23):
                run           = para.add_run(m.group(24))
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run.font.highlight_color = 7  # yellow

            # ── Line break <br> ────────────────────
            elif m.group(25):
                run = para.add_run()
                run.add_break()

            last_end = m.end()

        # Remaining plain text
        remaining = text[last_end:]
        if remaining:
            run           = para.add_run(remaining)
            run.font.name = font_name
            run.font.size = Pt(font_size)

    # ── Add code block ─────────────────────────────
    def add_code_block(code: str, language: str = ''):
        para           = doc.add_paragraph()
        run            = para.add_run(code)
        run.font.name  = 'Courier New'
        run.font.size  = Pt(font_size - 1)
        run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)

        # Grey background via paragraph shading
        pPr  = para._p.get_or_add_pPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F2F2F2')
        pPr.append(shd)

        # Left border (code style)
        pBdr   = OxmlElement('w:pBdr')
        left   = OxmlElement('w:left')
        left.set(qn('w:val'),   'single')
        left.set(qn('w:sz'),    '12')
        left.set(qn('w:space'), '4')
        left.set(qn('w:color'), '2E75B6')
        pBdr.append(left)
        pPr.append(pBdr)

        # Indentation
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '360')
        pPr.append(ind)

        set_spacing(para, 1.0)

        if language:
            # Add language label
            lang_para = doc.add_paragraph()
            lang_run  = lang_para.add_run(f'[{language}]')
            lang_run.font.name  = 'Courier New'
            lang_run.font.size  = Pt(8)
            lang_run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
            lang_run.italic = True

    # ── Add blockquote ─────────────────────────────
    def add_blockquote(text: str):
        para = doc.add_paragraph()
        apply_inline_to_para(text, para)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(para)

        pPr  = para._p.get_or_add_pPr()

        # Left blue border
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'),   'single')
        left.set(qn('w:sz'),    '16')
        left.set(qn('w:space'), '4')
        left.set(qn('w:color'), '2E75B6')
        pBdr.append(left)
        pPr.append(pBdr)

        # Left indent
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '720')
        pPr.append(ind)

        # Grey text
        for run in para.runs:
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            run.italic = True

    # ── Add horizontal rule ────────────────────────
    def add_horizontal_rule():
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after  = Pt(6)

        pPr  = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '6')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), '2E75B6')
        pBdr.append(bot)
        pPr.append(pBdr)

    # ── Add markdown table ─────────────────────────
    def add_table(lines: list):
        # Parse rows (skip separator row)
        rows_data = []
        for line in lines:
            line = line.strip()
            if re.match(r'^\|?[\s:]*-+[\s:]*(\|[\s:]*-+[\s:]*)*\|?$', line):
                continue  # skip separator
            cells = [
                c.strip()
                for c in line.strip('|').split('|')
            ]
            rows_data.append(cells)

        if not rows_data:
            return

        num_cols  = max(len(r) for r in rows_data)
        num_rows  = len(rows_data)
        rows_data = [r + [''] * (num_cols - len(r)) for r in rows_data]

        tbl = doc.add_table(rows=num_rows, cols=num_cols)
        tbl.style = 'Table Grid'

        from docx.oxml import OxmlElement as OXE

        for r_idx, row_data in enumerate(rows_data):
            row = tbl.rows[r_idx]
            for c_idx, cell_text in enumerate(row_data):
                cell = row.cells[c_idx]

                # Clear default paragraph
                for p in cell.paragraphs:
                    p._element.getparent().remove(p._element)

                new_para = cell.add_paragraph()
                apply_inline_to_para(cell_text, new_para)

                # Style header row
                if r_idx == 0:
                    for run in new_para.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

                    tc   = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd  = OXE('w:shd')
                    shd.set(qn('w:fill'),  '2E75B6')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:val'),   'clear')
                    tcPr.append(shd)
                else:
                    # Alternating rows
                    if r_idx % 2 == 0:
                        tc   = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        shd  = OXE('w:shd')
                        shd.set(qn('w:fill'),  'DCE6F1')
                        shd.set(qn('w:color'), 'auto')
                        shd.set(qn('w:val'),   'clear')
                        tcPr.append(shd)

        doc.add_paragraph()   # spacing after table

    # ── Add list item ──────────────────────────────
    def add_list_item(
        text    : str,
        ordered : bool = False,
        level   : int  = 0,
        checked : bool = None,
    ):
        style = 'List Number' if ordered else 'List Bullet'
        para  = doc.add_paragraph(style=style)

        # Indent nested lists
        if level > 0:
            from docx.shared import Inches as IN
            para.paragraph_format.left_indent = IN(0.5 * (level + 1))

        # Task list checkbox
        if checked is not None:
            checkbox = '☑ ' if checked else '☐ '
            run      = para.add_run(checkbox)
            run.font.name = font_name
            run.font.size = Pt(font_size)

        apply_inline_to_para(text.strip(), para)
        set_spacing(para)
        return para

    # ─────────────────────────────────────────────
    # Parse Markdown lines
    # ─────────────────────────────────────────────
    lines      = raw.splitlines()
    i          = 0
    total      = len(lines)

    # Add title if provided
    if title:
        add_heading(title, 1)
        doc.add_paragraph()

    while i < total:
        line     = lines[i]
        stripped = line.strip()

        # ── Empty line ─────────────────────────────
        if not stripped:
            doc.add_paragraph()
            i += 1
            continue

        # ── Fenced code block ```lang ──────────────
        if stripped.startswith('```') or stripped.startswith('~~~'):
            fence    = stripped[:3]
            language = stripped[3:].strip()
            i       += 1
            code_lines = []

            while i < total and not lines[i].strip().startswith(fence):
                code_lines.append(lines[i])
                i += 1

            code = '\n'.join(code_lines)
            add_code_block(code, language)
            i += 1   # skip closing fence
            continue

        # ── ATX Headings # ## ### ──────────────────
        m = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if m:
            level = len(m.group(1))
            text  = m.group(2).rstrip('#').strip()
            add_heading(text, level)
            i += 1
            continue

        # ── Setext H1 (===) ────────────────────────
        if i + 1 < total:
            next_line = lines[i + 1].strip()
            if next_line and set(next_line) == {'='}:
                add_heading(stripped, 1)
                i += 2
                continue
            if next_line and set(next_line) == {'-'} and len(next_line) >= 2:
                add_heading(stripped, 2)
                i += 2
                continue

        # ── Horizontal rule --- *** ___ ────────────
        if re.match(r'^[-*_]{3,}$', stripped.replace(' ', '')):
            add_horizontal_rule()
            i += 1
            continue

        # ── Blockquote > ───────────────────────────
        if stripped.startswith('>'):
            quote_lines = []
            while i < total and lines[i].strip().startswith('>'):
                quote_lines.append(
                    lines[i].strip().lstrip('>').strip()
                )
                i += 1
            add_blockquote(' '.join(quote_lines))
            continue

        # ── Table | col | col | ────────────────────
        if stripped.startswith('|') and '|' in stripped[1:]:
            table_lines = []
            while i < total and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            add_table(table_lines)
            continue

        # ── Bullet list - * + ──────────────────────
        m = re.match(r'^(\s*)[-*+]\s+(\[[ xX]\]\s+)?(.+)$', line)
        if m:
            indent  = len(m.group(1))
            level   = indent // 2
            checked = None
            text    = m.group(3)

            if m.group(2):
                checked = m.group(2).strip()[1].lower() == 'x'

            add_list_item(text, ordered=False, level=level, checked=checked)
            i += 1
            continue

        # ── Numbered list 1. 2. ────────────────────
        m = re.match(r'^(\s*)\d+[.)]\s+(.+)$', line)
        if m:
            indent = len(m.group(1))
            level  = indent // 3
            text   = m.group(2)
            add_list_item(text, ordered=True, level=level)
            i += 1
            continue

        # ── Regular paragraph ──────────────────────
        # Collect continuation lines
        para_lines = [stripped]
        i += 1

        while i < total:
            next_stripped = lines[i].strip()

            # Stop at blank line or block-level elements
            if not next_stripped:
                break
            if next_stripped.startswith(('#', '>', '|', '```', '~~~', '-', '*', '+')):
                break
            if re.match(r'^\d+[.)]\s', next_stripped):
                break
            if re.match(r'^[-*_]{3,}$', next_stripped.replace(' ', '')):
                break

            para_lines.append(next_stripped)
            i += 1

        full_text = ' '.join(para_lines)
        add_paragraph(full_text)

    # ── Serialize ─────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def word_to_latex(
    source,
    filename: str = "document.docx",
    document_class: str = "article",
    font_size: int = 12,
    paper_size: str = "a4paper",
    include_packages: bool = True,
    include_toc: bool = False,
    include_title: bool = True,
    encoding: str = "utf-8",
) -> dict:
    """
    Convert Word (.docx) → LaTeX (.tex).

    Handles:
        - Headings      → \\section \\subsection \\subsubsection
        - Paragraphs    → plain text paragraphs
        - Bold          → \\textbf{}
        - Italic        → \\textit{}
        - Underline     → \\underline{}
        - Strikethrough → \\sout{}
        - Bullet lists  → itemize
        - Numbered lists→ enumerate
        - Tables        → tabular environment
        - Code blocks   → verbatim / lstlisting
        - Hyperlinks    → \\href{}
        - Page breaks   → \\newpage
        - Horizontal rules → \\hrule

    Args:
        source           : file object | raw bytes
        filename         : original filename
        document_class   : article | report | book    (default: article)
        font_size        : 10 | 11 | 12              (default: 12)
        paper_size       : a4paper | letterpaper     (default: a4paper)
        include_packages : include usepackage lines  (default: True)
        include_toc      : include tableofcontents   (default: False)
        include_title    : include title block       (default: True)
        encoding         : input encoding            (default: utf-8)

    Returns:
        {
            'latex'       : str,
            'headings'    : list,
            'word_count'  : int,
            'char_count'  : int,
            'table_count' : int,
            'list_count'  : int,
        }
    """
    from docx import Document as DocxDocument
    from docx.text.paragraph import Paragraph as DocxPara
    from docx.table import Table as DocxTable
    from docx.oxml.ns import qn
    import re

    # ── Read source ───────────────────────────────
    if hasattr(source, "read"):
        file_bytes = source.read()
    elif isinstance(source, bytes):
        file_bytes = source
    else:
        raise ValueError("source must be a file object or bytes.")

    if not file_bytes:
        raise ValueError("Empty file.")

    # ── Open document ─────────────────────────────
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"Cannot open Word document: {e}")

    # ── Validate ──────────────────────────────────
    if document_class not in ("article", "report", "book"):
        document_class = "article"
    if font_size not in (10, 11, 12):
        font_size = 12
    if paper_size not in ("a4paper", "letterpaper", "legalpaper", "a3paper"):
        paper_size = "a4paper"

    # ── Tracking ──────────────────────────────────
    headings = []
    table_count = 0
    list_count = 0

    # ── LaTeX escape ──────────────────────────────
    def escape_latex(text: str) -> str:
        """Escape LaTeX special characters."""
        replacements = [
            ("\\", "\\textbackslash{}"),
            ("&", "\\&"),
            ("%", "\\%"),
            ("$", "\\$"),
            ("#", "\\#"),
            ("^", "\\^{}"),
            ("_", "\\_"),
            ("{", "\\{"),
            ("}", "\\}"),
            ("~", "\\textasciitilde{}"),
            ("<", "\\textless{}"),
            (">", "\\textgreater{}"),
            ("|", "\\textbar{}"),
        ]
        # Replace backslash first to avoid double escaping
        text = text.replace("\\", "\\textbackslash{}")
        for old, new in replacements[1:]:
            text = text.replace(old, new)
        return text

    # ── Process runs → LaTeX ──────────────────────
    def runs_to_latex(para) -> str:
        parts = []
        for run in para.runs:
            text = escape_latex(run.text)
            if not text:
                continue

            # Bold + Italic
            if run.bold and run.italic:
                text = f"\\textbf{{\\textit{{{text}}}}}"
            elif run.bold:
                text = f"\\textbf{{{text}}}"
            elif run.italic:
                text = f"\\textit{{{text}}}"

            # Underline
            if run.underline:
                text = f"\\underline{{{text}}}"

            # Strikethrough
            if run.font.strike:
                text = f"\\sout{{{text}}}"

            # Monospace font
            try:
                font_name = run.font.name or ""
                if "courier" in font_name.lower() or "mono" in font_name.lower():
                    text = f"\\texttt{{{text}}}"
            except Exception:
                pass

            # Hyperlink
            try:
                parent = run._element.getparent()
                if parent.tag.endswith("}hyperlink"):
                    rId = parent.get(qn("r:id"), "")
                    href = ""
                    try:
                        href = doc.part.rels[rId].target_ref
                    except Exception:
                        pass
                    if href:
                        text = f"\\href{{{href}}}{{{text}}}"
            except Exception:
                pass

            parts.append(text)

        return "".join(parts)

    # ── Process heading ────────────────────────────
    def process_heading(para, level: int) -> str:
        text = escape_latex(para.text.strip())
        headings.append({"level": level, "text": para.text.strip()})

        cmd_map = {
            1: "\\section",
            2: "\\subsection",
            3: "\\subsubsection",
            4: "\\paragraph",
            5: "\\subparagraph",
            6: "\\subparagraph",
        }
        # report/book use chapter for H1
        if document_class in ("report", "book") and level == 1:
            return f"\\chapter{{{text}}}"

        cmd = cmd_map.get(level, "\\paragraph")
        return f"{cmd}{{{text}}}"

    # ── Process table ──────────────────────────────
    def process_table(tbl) -> str:
        nonlocal table_count
        table_count += 1

        rows_data = []
        for row in tbl.rows:
            row_data = []
            for cell in row.cells:
                cell_text = escape_latex(
                    " ".join(p.text.strip() for p in cell.paragraphs)
                )
                row_data.append(cell_text)
            rows_data.append(row_data)

        if not rows_data:
            return ""

        num_cols = max(len(r) for r in rows_data)
        rows_data = [r + [""] * (num_cols - len(r)) for r in rows_data]

        col_spec = " | ".join(["l"] * num_cols)

        latex_lines = [
            "",
            "\\begin{table}[h!]",
            "\\centering",
            f"\\begin{{tabular}}{{| {col_spec} |}}",
            "\\hline",
        ]

        for r_idx, row in enumerate(rows_data):
            row_str = " & ".join(
                f"\\textbf{{{cell}}}" if r_idx == 0 else cell for cell in row
            )
            latex_lines.append(f"{row_str} \\\\")
            if r_idx == 0:
                latex_lines.append("\\hline")

        latex_lines += [
            "\\hline",
            "\\end{tabular}",
            f"\\caption{{Table {table_count}}}",
            f"\\label{{tab:table{table_count}}}",
            "\\end{table}",
            "",
        ]

        return "\n".join(latex_lines)

    # ── Process list ───────────────────────────────
    def process_list_block(items: list) -> str:
        """
        items: list of (style_name, latex_text, level)
        Returns LaTeX itemize/enumerate block.
        """
        nonlocal list_count
        list_count += 1

        if not items:
            return ""

        # Detect type from first item
        first_style = items[0][0]
        env = "enumerate" if "number" in first_style else "itemize"

        lines = [f"\\begin{{{env}}}"]
        for _, text, _ in items:
            lines.append(f"  \\item {text}")
        lines.append(f"\\end{{{env}}}")

        return "\n".join(lines)

    # ─────────────────────────────────────────────
    # Process document body
    # ─────────────────────────────────────────────
    body_lines = []
    list_buffer = []  # collect consecutive list items

    def flush_list():
        """Flush accumulated list items to body."""
        if list_buffer:
            body_lines.append(process_list_block(list_buffer))
            list_buffer.clear()

    for element in doc.element.body:
        tag = element.tag.split("}")[-1]

        if tag == "p":
            para = DocxPara(element, doc)
            text = para.text.strip()
            style_name = (para.style.name or "Normal").lower()

            # ── Page break ─────────────────────────
            if any(
                br.get(qn("w:type")) == "page"
                for br in element.findall(".//" + qn("w:br"))
            ):
                flush_list()
                body_lines.append("\\newpage")
                continue

            # ── Headings ───────────────────────────
            if "heading 1" in style_name or style_name == "title":
                flush_list()
                body_lines.append(process_heading(para, 1))
                body_lines.append("")
                continue

            if "heading 2" in style_name or style_name == "subtitle":
                flush_list()
                body_lines.append(process_heading(para, 2))
                body_lines.append("")
                continue

            if "heading 3" in style_name:
                flush_list()
                body_lines.append(process_heading(para, 3))
                body_lines.append("")
                continue

            if "heading 4" in style_name:
                flush_list()
                body_lines.append(process_heading(para, 4))
                body_lines.append("")
                continue

            if "heading 5" in style_name:
                flush_list()
                body_lines.append(process_heading(para, 5))
                body_lines.append("")
                continue

            if "heading 6" in style_name:
                flush_list()
                body_lines.append(process_heading(para, 6))
                body_lines.append("")
                continue

            # ── Empty paragraph ────────────────────
            if not text:
                flush_list()
                body_lines.append("")
                continue

            # ── List items ─────────────────────────
            if "list bullet" in style_name or "list number" in style_name:
                latex_text = runs_to_latex(para)
                list_buffer.append((style_name, latex_text, 0))
                continue

            # ── Code / preformatted ────────────────
            if "code" in style_name or "preformat" in style_name:
                flush_list()
                body_lines.append("\\begin{verbatim}")
                body_lines.append(para.text)
                body_lines.append("\\end{verbatim}")
                body_lines.append("")
                continue

            # ── Blockquote ─────────────────────────
            if "quote" in style_name or "block text" in style_name:
                flush_list()
                latex_text = runs_to_latex(para)
                body_lines.append("\\begin{quote}")
                body_lines.append(latex_text)
                body_lines.append("\\end{quote}")
                body_lines.append("")
                continue

            # ── Horizontal rule ────────────────────
            if re.match(r"^[-_*]{3,}$", text.replace(" ", "")):
                flush_list()
                body_lines.append("")
                body_lines.append("\\hrule")
                body_lines.append("")
                continue

            # ── Regular paragraph ──────────────────
            flush_list()
            latex_text = runs_to_latex(para)
            if latex_text.strip():
                body_lines.append(latex_text)
                body_lines.append("")

        elif tag == "tbl":
            flush_list()
            try:
                tbl = DocxTable(element, doc)
                body_lines.append(process_table(tbl))
            except Exception as e:
                body_lines.append(f"% Table error: {e}")

        elif tag == "sectPr":
            pass

    flush_list()

    # ─────────────────────────────────────────────
    # Build LaTeX document
    # ─────────────────────────────────────────────
    doc_title = filename.replace(".docx", "").replace(".doc", "")
    latex_parts = []

    # ── Preamble ──────────────────────────────────
    latex_parts.append(
        f"\\documentclass[{font_size}pt,{paper_size}]{{{document_class}}}"
    )

    if include_packages:
        latex_parts.append("")
        latex_parts.append("% ── Packages ─────────────────────────────")
        latex_parts.append("\\usepackage[utf8]{inputenc}")
        latex_parts.append("\\usepackage[T1]{fontenc}")
        latex_parts.append("\\usepackage{lmodern}")
        latex_parts.append("\\usepackage{amsmath}")
        latex_parts.append("\\usepackage{amssymb}")
        latex_parts.append("\\usepackage{graphicx}")
        latex_parts.append("\\usepackage{hyperref}")
        latex_parts.append("\\usepackage{booktabs}")
        latex_parts.append("\\usepackage{array}")
        latex_parts.append("\\usepackage{longtable}")
        latex_parts.append("\\usepackage{xcolor}")
        latex_parts.append("\\usepackage[normalem]{ulem}")  # for \\sout
        latex_parts.append("\\usepackage{listings}")
        latex_parts.append("\\usepackage{geometry}")
        latex_parts.append(f"\\geometry{{{paper_size}, margin=1in}}")

        # hyperref setup
        latex_parts.append("")
        latex_parts.append("% ── Hyperref setup ───────────────────────")
        latex_parts.append("\\hypersetup{")
        latex_parts.append("    colorlinks=true,")
        latex_parts.append("    linkcolor=blue,")
        latex_parts.append("    urlcolor=blue,")
        latex_parts.append("    citecolor=green,")
        latex_parts.append("}")

        # lstlisting setup
        latex_parts.append("")
        latex_parts.append("% ── Code listing setup ───────────────────")
        latex_parts.append("\\lstset{")
        latex_parts.append("    basicstyle=\\ttfamily\\small,")
        latex_parts.append("    breaklines=true,")
        latex_parts.append("    frame=single,")
        latex_parts.append("    backgroundcolor=\\color{gray!10},")
        latex_parts.append("    keywordstyle=\\color{blue},")
        latex_parts.append("    commentstyle=\\color{green!60!black},")
        latex_parts.append("    stringstyle=\\color{red},")
        latex_parts.append("}")

    # ── Title block ───────────────────────────────
    if include_title:
        latex_parts.append("")
        latex_parts.append("% ── Document info ────────────────────────")
        latex_parts.append(f"\\title{{{escape_latex(doc_title)}}}")
        latex_parts.append("\\author{}")
        latex_parts.append("\\date{\\today}")

    # ── Begin document ────────────────────────────
    latex_parts.append("")
    latex_parts.append("% ── Begin document ───────────────────────")
    latex_parts.append("\\begin{document}")

    if include_title:
        latex_parts.append("")
        latex_parts.append("\\maketitle")

    if include_toc:
        latex_parts.append("")
        latex_parts.append("\\tableofcontents")
        latex_parts.append("\\newpage")

    # ── Body ──────────────────────────────────────
    latex_parts.append("")
    latex_parts.extend(body_lines)

    # ── End document ──────────────────────────────
    latex_parts.append("")
    latex_parts.append("\\end{document}")

    # ── Join ──────────────────────────────────────
    import re as _re

    latex = "\n".join(latex_parts)
    latex = _re.sub(r"\n{3,}", "\n\n", latex)

    word_count = len(latex.split())
    char_count = len(latex)

    return {
        "latex": latex,
        "headings": headings,
        "word_count": word_count,
        "char_count": char_count,
        "table_count": table_count,
        "list_count": list_count,
    }

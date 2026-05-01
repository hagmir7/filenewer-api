"""
PDF service functions.
"""

import io
import os
import re
import logging
from pathlib import Path

import pdfplumber
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────
# PDF TO WORD
# ─────────────────────────────────────────────


def _looks_like_heading(text: str) -> bool:
    """Heuristic: short, no period at end, possible ALL CAPS."""
    text = text.strip()
    if not text:
        return False
    if len(text) > 120:
        return False
    if text.endswith("."):
        return False
    if text.isupper() and len(text) > 3:
        return True
    if re.match(r"^(\d+[\.\)]\s+|\d+\.\d+\s+)", text):  # "1. Title" / "1.2 Section"
        return True
    return False


def _add_table_to_doc(doc: Document, table_data: list[list]) -> None:
    """Add a pdfplumber table into the Word document."""
    if not table_data:
        return

    rows = len(table_data)
    cols = max(len(row) for row in table_data)
    if rows == 0 or cols == 0:
        return

    word_table = doc.add_table(rows=rows, cols=cols)
    word_table.style = "Table Grid"

    for r_idx, row in enumerate(table_data):
        for c_idx, cell_text in enumerate(row):
            if c_idx >= cols:
                break
            cell = word_table.cell(r_idx, c_idx)
            cell.text = str(cell_text) if cell_text is not None else ""
            # Bold first row (header)
            if r_idx == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def convert_pdf_to_docx(pdf_bytes: bytes) -> bytes:
    """
    Convert PDF bytes → DOCX bytes.

    Returns the raw bytes of the generated .docx file.
    """
    doc = Document()

    # ── Default style tweaks ──────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        total_pages = len(pdf.pages)
        logger.info("Converting PDF: %d page(s)", total_pages)

        for page_num, page in enumerate(pdf.pages, start=1):
            # ── Extract tables first ─────────────
            tables = page.extract_tables() or []
            table_bboxes = []

            for table_data in tables:
                if table_data:
                    _add_table_to_doc(doc, table_data)
                    doc.add_paragraph()  # breathing room after table

            # ── Extract text (excluding table regions) ──
            # Crop away table areas to avoid duplicating content
            cropped_page = page
            for table in page.find_tables():
                try:
                    cropped_page = cropped_page.outside_bbox(table.bbox)
                except Exception:
                    pass  # some versions don't support outside_bbox; fall through

            raw_text = cropped_page.extract_text(x_tolerance=3, y_tolerance=3) or ""

            if not raw_text.strip() and not tables:
                # Blank page – add a soft separator
                doc.add_paragraph()
                continue

            for line in raw_text.splitlines():
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                    continue

                if _looks_like_heading(line):
                    heading_level = 1 if line.isupper() else 2
                    doc.add_heading(line, level=heading_level)
                else:
                    para = doc.add_paragraph(line)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # ── Page break between pages (except last) ──
            if page_num < total_pages:
                doc.add_page_break()

    # ── Serialize to bytes ────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def pdf_to_excel(source, password: str = None) -> bytes:
    """
    Convert PDF (file or bytes) → Excel bytes.

    Strategy:
        1. Extract tables  → each table gets its own sheet
        2. Extract text    → one 'Text' sheet with all plain text
        3. Style headers   → bold white on blue, alternating rows, auto-width

    Args:
        source   : uploaded file object OR raw bytes
        password : PDF password if encrypted (default: None)

    Returns:
        Raw bytes of the generated .xlsx file
    """
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl import Workbook

    # ── Read PDF bytes ────────────────────────────
    if hasattr(source, "read"):
        pdf_bytes = source.read()
    else:
        pdf_bytes = source

    # ── Styles ────────────────────────────────────
    def make_header_style():
        return {
            "font": Font(bold=True, color="FFFFFF", size=11),
            "fill": PatternFill(fill_type="solid", fgColor="2E75B6"),
            "align": Alignment(horizontal="center", vertical="center", wrap_text=True),
            "border": Border(
                left=Side(style="thin"),
                right=Side(style="thin"),
                top=Side(style="thin"),
                bottom=Side(style="thin"),
            ),
        }

    def apply_header(ws):
        style = make_header_style()
        for cell in ws[1]:
            cell.font = style["font"]
            cell.fill = style["fill"]
            cell.alignment = style["align"]
            cell.border = style["border"]
        ws.row_dimensions[1].height = 25

    def apply_rows(ws):
        thin = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin"),
        )
        for row_idx, row in enumerate(ws.iter_rows(min_row=2), start=2):
            color = "DCE6F1" if row_idx % 2 == 0 else "FFFFFF"
            for cell in row:
                cell.fill = PatternFill(fill_type="solid", fgColor=color)
                cell.border = thin
                cell.alignment = Alignment(vertical="center")

    def auto_fit(ws):
        for col_cells in ws.columns:
            max_len = max(
                len(str(cell.value)) if cell.value is not None else 0
                for cell in col_cells
            )
            ws.column_dimensions[col_cells[0].column_letter].width = min(
                max_len + 4, 60
            )

    # ── Parse PDF ─────────────────────────────────
    wb = Workbook()
    wb.remove(wb.active)  # remove default empty sheet

    table_count = 0
    all_text = []

    open_kwargs = {"stream": io.BytesIO(pdf_bytes)}
    pdf_stream = io.BytesIO(pdf_bytes)
    if password:
        open_kwargs["password"] = password

    with pdfplumber.open(pdf_stream, password=password) as pdf:
        total_pages = len(pdf.pages)

        for page_num, page in enumerate(pdf.pages, start=1):

            # ── Extract tables ─────────────────────
            tables = page.extract_tables() or []
            for table in tables:
                if not table:
                    continue

                # clean None cells
                cleaned = [
                    [str(cell).strip() if cell is not None else "" for cell in row]
                    for row in table
                ]

                table_count += 1
                sheet_label = f"Table_{table_count}_P{page_num}"[:31]
                ws = wb.create_sheet(title=sheet_label)

                for row in cleaned:
                    ws.append(row)

                apply_header(ws)
                apply_rows(ws)
                auto_fit(ws)
                ws.freeze_panes = "A2"

            # ── Extract plain text ─────────────────
            # Exclude table regions to avoid duplication
            cropped = page
            for tbl in page.find_tables():
                try:
                    cropped = cropped.outside_bbox(tbl.bbox)
                except Exception:
                    pass

            text = cropped.extract_text(x_tolerance=3, y_tolerance=3) or ""
            if text.strip():
                all_text.append(f"--- Page {page_num} ---")
                all_text.append(text.strip())
                all_text.append("")

    # ── Text sheet ────────────────────────────────
    if all_text:
        ws_text = wb.create_sheet(title="Text")
        ws_text.column_dimensions["A"].width = 100

        # Header
        ws_text.append(["Extracted Text"])
        apply_header(ws_text)

        for line in all_text:
            ws_text.append([line])

        # Light style for text rows
        thin = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin"),
        )
        for row in ws_text.iter_rows(min_row=2):
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical="top")
                cell.border = thin

    # ── Summary sheet (first sheet) ───────────────
    ws_summary = wb.create_sheet(title="Summary", index=0)
    ws_summary.column_dimensions["A"].width = 25
    ws_summary.column_dimensions["B"].width = 40

    summary_data = [
        ["Property", "Value"],
        ["Total Pages", total_pages],
        ["Tables Found", table_count],
        ["Has Text", "Yes" if all_text else "No"],
        ["Sheets", wb.sheetnames.__len__()],
    ]

    for row in summary_data:
        ws_summary.append(row)

    apply_header(ws_summary)
    apply_rows(ws_summary)
    auto_fit(ws_summary)

    # ── Serialize ─────────────────────────────────
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.read()


def pdf_to_jpg(source, dpi: int = 200, quality: int = 85, password: str = None) -> list[dict]:
    """
    Convert PDF → JPG using pymupdf (no poppler / no system deps needed).

    Args:
        source   : uploaded file object OR raw bytes
        dpi      : image resolution (default: 200)
        quality  : JPG quality 1-95 (default: 85)
        password : PDF password if encrypted (default: None)

    Returns:
        list of {
            'page'    : int,
            'bytes'   : bytes,
            'width'   : int,
            'height'  : int,
            'filename': str,
        }
    """
    import fitz          # pymupdf
    from PIL import Image

    # ── Read PDF bytes ────────────────────────────
    if hasattr(source, 'read'):
        pdf_bytes = source.read()
    else:
        pdf_bytes = source

    # ── Open PDF ──────────────────────────────────
    doc = fitz.open(stream=pdf_bytes, filetype='pdf')

    if password:
        if not doc.authenticate(password):
            raise ValueError('Invalid PDF password.')

    # ── Convert each page → JPG ───────────────────
    zoom   = dpi / 72        # pymupdf base DPI is 72
    matrix = fitz.Matrix(zoom, zoom)

    results = []

    for page_num in range(len(doc)):
        page    = doc[page_num]
        pixmap  = page.get_pixmap(matrix=matrix, alpha=False)

        # Pixmap → PIL Image → JPG bytes
        image = Image.frombytes(
            'RGB',
            [pixmap.width, pixmap.height],
            pixmap.samples,
        )

        buffer = io.BytesIO()
        image.save(
            buffer,
            format  ='JPEG',
            quality =quality,
            optimize=True,
        )
        buffer.seek(0)
        jpg_bytes = buffer.read()

        results.append({
            'page'    : page_num + 1,
            'bytes'   : jpg_bytes,
            'width'   : pixmap.width,
            'height'  : pixmap.height,
            'filename': f'page_{page_num + 1}.jpg',
        })

    doc.close()
    return results


# ── Arabic support ─────────────────────────────────
try:
    from bidi.algorithm import get_display
    import arabic_reshaper

    ARABIC_SUPPORT = True
except ImportError:
    ARABIC_SUPPORT = False


def is_arabic(text: str) -> bool:
    """Check if text contains Arabic characters."""
    return bool(re.search(r"[؀-ۿݐ-ݿࢠ-ࣿ]", text))


def process_arabic_text(text: str) -> str:
    """
    Reshape and reorder Arabic text for correct PDF rendering.
    Arabic needs two steps:
        1. Reshape  → connect letters correctly
        2. Bidi     → reverse for RTL display
    """
    if not ARABIC_SUPPORT or not text.strip():
        return text

    try:
        if is_arabic(text):
            reshaped = arabic_reshaper.reshape(text)
            return get_display(reshaped)
        return text
    except Exception:
        return text


def register_arabic_fonts(fonts_dir: str = None) -> dict:
    """
    Register Arabic fonts with reportlab.
    Returns dict of available font names.
    """
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    if fonts_dir is None:
        # Default: fonts/ folder next to manage.py
        fonts_dir = os.path.join(
            os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "fonts"
        )

    registered = {}
    font_files = {
        "Amiri": "Amiri-Regular.ttf",
        "Amiri-Bold": "Amiri-Bold.ttf",
        "Amiri-Italic": "Amiri-Italic.ttf",
        "Amiri-BoldItalic": "Amiri-BoldItalic.ttf",
    }

    for font_name, font_file in font_files.items():
        font_path = os.path.join(fonts_dir, font_file)
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont(font_name, font_path))
                registered[font_name] = font_path
            except Exception:
                pass

    return registered


def encrypt_pdf(
    source,
    user_password: str = "",
    owner_password: str = None,
    allow_printing: bool = True,
    allow_copying: bool = True,
    allow_editing: bool = True,
    allow_annotations: bool = True,
) -> bytes:
    """
    Encrypt a PDF with password protection and permissions.

    Args:
        source           : uploaded file object OR raw bytes
        user_password    : password required to open the PDF
        owner_password   : password for full access (default: same as user)
        allow_printing   : allow printing                (default: True)
        allow_copying    : allow copying text            (default: True)
        allow_editing    : allow editing                 (default: True)
        allow_annotations: allow adding annotations      (default: True)

    Returns:
        Raw bytes of the encrypted PDF
    """
    from pypdf import PdfReader, PdfWriter
    from pypdf.generic import NameObject
    import pypdf.constants as pdfconst

    # ── Read PDF bytes ────────────────────────────
    if hasattr(source, "read"):
        pdf_bytes = source.read()
    else:
        pdf_bytes = source

    # ── Validate it is a real PDF ──────────────────
    if not pdf_bytes.startswith(b"%PDF"):
        raise ValueError("Invalid PDF file.")

    # ── Read + Write ──────────────────────────────
    reader = PdfReader(io.BytesIO(pdf_bytes))

    # Check if already encrypted
    if reader.is_encrypted:
        raise ValueError(
            "PDF is already encrypted. " "Please decrypt it first before re-encrypting."
        )

    writer = PdfWriter()

    # ── Copy all pages ─────────────────────────────
    for page in reader.pages:
        writer.add_page(page)

    # ── Copy metadata ──────────────────────────────
    if reader.metadata:
        writer.add_metadata(reader.metadata)

    # ── Set owner password ─────────────────────────
    if owner_password is None:
        owner_password = user_password

    # ── Build permissions ─────────────────────────
    permissions = build_permissions(
        allow_printing=allow_printing,
        allow_copying=allow_copying,
        allow_editing=allow_editing,
        allow_annotations=allow_annotations,
    )

    # ── Encrypt ────────────────────────────────────
    writer.encrypt(
        user_password=user_password,
        owner_password=owner_password,
        use_128bit=True,
        permissions_flag=permissions,
    )

    # ── Serialize ─────────────────────────────────
    buffer = io.BytesIO()
    writer.write(buffer)
    buffer.seek(0)
    return buffer.read()


def build_permissions(
    allow_printing: bool = True,
    allow_copying: bool = True,
    allow_editing: bool = True,
    allow_annotations: bool = True,
) -> int:
    """
    Build PDF permissions flag.

    PDF permission bits (128-bit RC4):
        Bit 3  → Print
        Bit 4  → Modify contents
        Bit 5  → Copy / extract text
        Bit 6  → Annotations
        Bit 9  → Fill forms
        Bit 10 → Extract for accessibility
        Bit 11 → Assemble document
        Bit 12 → Print high quality
    """
    # Start with base permissions (bits 1,2 always 0; rest 1)
    permissions = 0b11111111111111111111111100000000

    if not allow_printing:
        permissions &= ~(1 << 2)  # clear bit 3
        permissions &= ~(1 << 11)  # clear bit 12 (high quality print)

    if not allow_editing:
        permissions &= ~(1 << 3)  # clear bit 4
        permissions &= ~(1 << 8)  # clear bit 9  (forms)
        permissions &= ~(1 << 10)  # clear bit 11 (assemble)

    if not allow_copying:
        permissions &= ~(1 << 4)  # clear bit 5
        permissions &= ~(1 << 9)  # clear bit 10 (accessibility)

    if not allow_annotations:
        permissions &= ~(1 << 5)  # clear bit 6

    return permissions


def decrypt_pdf(source, password: str) -> bytes:
    """
    Decrypt / remove password from a PDF.

    Args:
        source   : uploaded file object OR raw bytes
        password : password to decrypt the PDF

    Returns:
        Raw bytes of the decrypted PDF
    """
    from pypdf import PdfReader, PdfWriter

    # ── Read PDF bytes ────────────────────────────
    if hasattr(source, "read"):
        pdf_bytes = source.read()
    else:
        pdf_bytes = source

    reader = PdfReader(io.BytesIO(pdf_bytes))

    # ── Check if encrypted ─────────────────────────
    if not reader.is_encrypted:
        raise ValueError("PDF is not encrypted.")

    # ── Decrypt ────────────────────────────────────
    result = reader.decrypt(password)
    if result == 0:
        raise ValueError("Wrong password. Could not decrypt PDF.")

    # ── Copy to new writer (removes encryption) ────
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)

    if reader.metadata:
        writer.add_metadata(reader.metadata)

    # ── Serialize ─────────────────────────────────
    buffer = io.BytesIO()
    writer.write(buffer)
    buffer.seek(0)
    return buffer.read()


def get_pdf_info(source, password: str = None) -> dict:
    """
    Get PDF metadata and encryption info.

    Args:
        source   : uploaded file object OR raw bytes
        password : password if encrypted (optional)

    Returns:
        dict with PDF info
    """
    from pypdf import PdfReader

    if hasattr(source, "read"):
        pdf_bytes = source.read()
    else:
        pdf_bytes = source

    reader = PdfReader(io.BytesIO(pdf_bytes))

    # ── Decrypt if needed ──────────────────────────
    if reader.is_encrypted:
        if not password:
            return {
                "encrypted": True,
                "pages": None,
                "metadata": None,
                "permissions": None,
                "message": "PDF is encrypted. Provide password to get full info.",
            }
        result = reader.decrypt(password)
        if result == 0:
            raise ValueError("Wrong password.")

    # ── Collect metadata ───────────────────────────
    meta = {}
    if reader.metadata:
        meta = {
            "title": reader.metadata.get("/Title", ""),
            "author": reader.metadata.get("/Author", ""),
            "subject": reader.metadata.get("/Subject", ""),
            "creator": reader.metadata.get("/Creator", ""),
            "producer": reader.metadata.get("/Producer", ""),
            "created": str(reader.metadata.get("/CreationDate", "")),
            "modified": str(reader.metadata.get("/ModDate", "")),
        }

    return {
        "encrypted": reader.is_encrypted,
        "pages": len(reader.pages),
        "metadata": meta,
        "file_size": len(pdf_bytes),
        "pdf_version": reader.pdf_header,
    }


def compress_pdf(
    source,
    compression_level: str = "medium",
    password: str = None,
) -> dict:
    """
    Compress a PDF file by reducing image quality and removing redundant data.

    Args:
        source            : uploaded file object OR raw bytes
        compression_level : 'low' | 'medium' | 'high' | 'extreme'
        password          : PDF password if encrypted

    Compression levels:
        low     → light compression, best quality  (image DPI: 150, quality: 85)
        medium  → balanced compression             (image DPI: 120, quality: 72)
        high    → aggressive compression           (image DPI: 96,  quality: 60)
        extreme → maximum compression, low quality (image DPI: 72,  quality: 40)

    Returns:
        {
            'bytes'            : bytes,
            'original_size'    : int,
            'compressed_size'  : int,
            'reduction_percent': float,
            'compression_level': str,
        }
    """
    import fitz  # pymupdf
    from PIL import Image

    # ── Read PDF bytes ────────────────────────────
    if hasattr(source, "read"):
        pdf_bytes = source.read()
    else:
        pdf_bytes = source

    original_size = len(pdf_bytes)

    # ── Validate PDF ──────────────────────────────
    if not pdf_bytes.startswith(b"%PDF"):
        raise ValueError("Invalid PDF file.")

    # ── Compression level config ───────────────────
    levels = {
        "low": {"dpi": 150, "quality": 85, "deflate": 3},
        "medium": {"dpi": 120, "quality": 72, "deflate": 6},
        "high": {"dpi": 96, "quality": 60, "deflate": 7},
        "extreme": {"dpi": 72, "quality": 40, "deflate": 9},
    }

    if compression_level not in levels:
        raise ValueError(
            f'Invalid compression_level: "{compression_level}". '
            f"Must be one of: {list(levels.keys())}"
        )

    config = levels[compression_level]
    dpi = config["dpi"]
    quality = config["quality"]

    # ── Open PDF ──────────────────────────────────
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    if doc.is_encrypted:
        if not password:
            raise ValueError("PDF is encrypted. Provide a password.")
        if not doc.authenticate(password):
            raise ValueError("Wrong password.")

    # ── Step 1: Compress images on each page ───────
    for page_num in range(len(doc)):
        page = doc[page_num]

        # Get all images on this page
        image_list = page.get_images(full=True)

        for img_info in image_list:
            xref = img_info[0]  # image reference number
            base_image = doc.extract_image(xref)

            if not base_image:
                continue

            img_bytes = base_image["image"]
            img_ext = base_image["ext"]

            # Skip non-raster images (masks, etc.)
            if img_ext not in ("jpeg", "jpg", "png", "bmp", "tiff"):
                continue

            try:
                # ── Open image with PIL ───────────
                pil_img = Image.open(io.BytesIO(img_bytes))

                # Convert to RGB if needed
                if pil_img.mode in ("RGBA", "P", "LA"):
                    pil_img = pil_img.convert("RGB")
                elif pil_img.mode == "L":
                    pass  # keep grayscale as is

                # ── Resize if DPI is too high ──────
                orig_w, orig_h = pil_img.size
                if orig_w > dpi * 8 or orig_h > dpi * 11:
                    scale = min(dpi * 8 / orig_w, dpi * 11 / orig_h)
                    new_w = max(1, int(orig_w * scale))
                    new_h = max(1, int(orig_h * scale))
                    pil_img = pil_img.resize((new_w, new_h), Image.LANCZOS)

                # ── Save as JPEG with compression ──
                img_buffer = io.BytesIO()
                pil_img.save(
                    img_buffer,
                    format="JPEG",
                    quality=quality,
                    optimize=True,
                )
                img_buffer.seek(0)
                new_img_bytes = img_buffer.read()

                # Only replace if new image is smaller
                if len(new_img_bytes) < len(img_bytes):
                    doc.update_stream(xref, new_img_bytes)

            except Exception:
                continue  # skip problematic images silently

    # ── Step 2: Clean and compress PDF structure ───
    buffer = io.BytesIO()
    doc.save(
        buffer,
        garbage=4,  # remove unused objects (0-4, 4=most aggressive)
        deflate=True,  # compress streams
        deflate_images=True,  # compress images
        deflate_fonts=True,  # compress fonts
        clean=True,  # clean content streams
        pretty=False,  # no pretty printing (saves space)
        linear=False,  # no linearization (saves space)
    )
    doc.close()

    buffer.seek(0)
    compressed_bytes = buffer.read()
    compressed_size = len(compressed_bytes)

    # ── Calculate reduction ────────────────────────
    reduction = (
        (original_size - compressed_size) / original_size * 100
        if original_size > 0
        else 0
    )

    return {
        "bytes": compressed_bytes,
        "original_size": original_size,
        "compressed_size": compressed_size,
        "reduction_percent": round(reduction, 2),
        "compression_level": compression_level,
        "original_size_kb": round(original_size / 1024, 2),
        "compressed_size_kb": round(compressed_size / 1024, 2),
        "original_size_mb": round(original_size / (1024 * 1024), 2),
        "compressed_size_mb": round(compressed_size / (1024 * 1024), 2),
    }


def pdf_to_png(
    source,
    dpi: int = 200,
    pages: list = None,
    password: str = None,
) -> list[dict]:
    """
    Convert PDF (file or bytes) → list of PNG images (one per page).

    Args:
        source   : uploaded file object OR raw bytes
        dpi      : image resolution (default: 200)
        pages    : list of page numbers to convert (1-based)
                   None = convert all pages
        password : PDF password if encrypted

    Returns:
        list of {
            'page'    : int,
            'bytes'   : bytes,
            'width'   : int,
            'height'  : int,
            'filename': str,
        }
    """
    import fitz  # pymupdf
    from PIL import Image

    # ── Read PDF bytes ────────────────────────────
    if hasattr(source, "read"):
        pdf_bytes = source.read()
    else:
        pdf_bytes = source

    # ── Validate PDF ──────────────────────────────
    if not pdf_bytes.startswith(b"%PDF"):
        raise ValueError("Invalid PDF file.")

    # ── Open PDF ──────────────────────────────────
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    if doc.is_encrypted:
        if not password:
            raise ValueError("PDF is encrypted. Provide a password.")
        if not doc.authenticate(password):
            raise ValueError("Wrong password.")

    total_pages = len(doc)

    # ── Validate pages ────────────────────────────
    if pages is not None:
        invalid = [p for p in pages if not (1 <= p <= total_pages)]
        if invalid:
            raise ValueError(
                f"Invalid page numbers: {invalid}. "
                f"PDF has {total_pages} pages (1-{total_pages})."
            )
        pages_to_convert = pages
    else:
        pages_to_convert = list(range(1, total_pages + 1))

    # ── Convert each page → PNG ───────────────────
    zoom = dpi / 72  # pymupdf base DPI is 72
    matrix = fitz.Matrix(zoom, zoom)
    results = []

    for page_num in pages_to_convert:
        page = doc[page_num - 1]  # 0-based index

        # ── Render page → pixmap ──────────────────
        pixmap = page.get_pixmap(
            matrix=matrix,
            alpha=True,  # PNG supports transparency
        )

        # ── Pixmap → PIL Image → PNG bytes ────────
        image = Image.frombytes(
            "RGBA",
            [pixmap.width, pixmap.height],
            pixmap.samples,
        )

        buffer = io.BytesIO()
        image.save(
            buffer,
            format="PNG",
            optimize=True,
            compress_level=6,  # 0-9, 6 is balanced
        )
        buffer.seek(0)
        png_bytes = buffer.read()

        results.append(
            {
                "page": page_num,
                "bytes": png_bytes,
                "width": pixmap.width,
                "height": pixmap.height,
                "filename": f"page_{page_num}.png",
                "size_kb": round(len(png_bytes) / 1024, 2),
            }
        )

    doc.close()
    return results


def rotate_pdf(
    source,
    rotation : int  = 90,
    pages    : list = None,
    password : str  = None,
) -> tuple:
    """
    Rotate pages in a PDF.

    Args:
        source   : uploaded file object OR raw bytes
        rotation : degrees to rotate (90, 180, 270, -90)
        pages    : list of page numbers to rotate (1-based)
                   None = rotate all pages
        password : PDF password if encrypted

    Returns:
        (rotated_bytes, total_pages, rotated_count)
    """
    from pypdf import PdfReader, PdfWriter

    # ── Read PDF bytes ────────────────────────────
    if hasattr(source, 'read'):
        pdf_bytes = source.read()
    else:
        pdf_bytes = source

    # ── Validate PDF ──────────────────────────────
    if not pdf_bytes.startswith(b'%PDF'):
        raise ValueError('Invalid PDF file.')

    # ── Normalize rotation ─────────────────────────
    valid_rotations = [90, 180, 270, -90, -180, -270]
    if rotation not in valid_rotations:
        raise ValueError(
            f'Invalid rotation: {rotation}. '
            f'Must be one of: {valid_rotations}'
        )
    rotation = rotation % 360

    # ── Read PDF ──────────────────────────────────
    reader = PdfReader(io.BytesIO(pdf_bytes))

    # ── Decrypt if needed ──────────────────────────
    if reader.is_encrypted:
        if not password:
            raise ValueError('PDF is encrypted. Provide a password.')
        if reader.decrypt(password) == 0:
            raise ValueError('Wrong password.')

    total_pages = len(reader.pages)

    # ── Validate page numbers ──────────────────────
    if pages is not None:
        invalid = [p for p in pages if not (1 <= p <= total_pages)]
        if invalid:
            raise ValueError(
                f'Invalid page numbers: {invalid}. '
                f'PDF has {total_pages} pages (1-{total_pages}).'
            )
        pages_to_rotate = set(pages)
    else:
        pages_to_rotate = set(range(1, total_pages + 1))

    # ── Rotate + write ─────────────────────────────
    writer = PdfWriter()

    for i, page in enumerate(reader.pages, start=1):
        if i in pages_to_rotate:
            page.rotate(rotation)
        writer.add_page(page)

    # ── Copy metadata ──────────────────────────────
    if reader.metadata:
        writer.add_metadata(reader.metadata)

    # ── Serialize ─────────────────────────────────
    buffer = io.BytesIO()
    writer.write(buffer)
    buffer.seek(0)

    return buffer.read(), total_pages, len(pages_to_rotate)


def get_pdf_page_info(source, password: str = None) -> list:
    """
    Get rotation and dimension info for each page.

    Args:
        source   : uploaded file object OR raw bytes
        password : PDF password if encrypted

    Returns:
        list of { page, width, height, rotation }
    """
    from pypdf import PdfReader

    if hasattr(source, 'read'):
        pdf_bytes = source.read()
    else:
        pdf_bytes = source

    reader = PdfReader(io.BytesIO(pdf_bytes))

    if reader.is_encrypted:
        if not password:
            raise ValueError('PDF is encrypted. Provide a password.')
        if reader.decrypt(password) == 0:
            raise ValueError('Wrong password.')

    result = []
    for i, page in enumerate(reader.pages, start=1):
        result.append({
            'page'    : i,
            'width'   : float(page.mediabox.width),
            'height'  : float(page.mediabox.height),
            'rotation': page.rotation,
        })

    return result


def watermark_pdf(
    source,
    watermark_text: str = "CONFIDENTIAL",
    watermark_type: str = "text",
    watermark_image: bytes = None,
    opacity: float = 0.3,
    font_size: int = 60,
    color: str = "red",
    angle: int = 45,
    position: str = "center",
    pages: list = None,
    password: str = None,
) -> bytes:
    """
    Add text or image watermark to a PDF.
    Returns raw bytes of the watermarked PDF.
    """
    from pypdf import PdfReader, PdfWriter
    from reportlab.lib import colors as rl_colors
    from reportlab.pdfgen import canvas as rl_canvas

    # ── Read PDF bytes ────────────────────────────
    if hasattr(source, "read"):
        pdf_bytes = source.read()
    else:
        pdf_bytes = source

    if not pdf_bytes.startswith(b"%PDF"):
        raise ValueError("Invalid PDF file.")

    # ── Validate ──────────────────────────────────
    valid_positions = (
        "center",
        "top",
        "bottom",
        "top-left",
        "top-right",
        "bottom-left",
        "bottom-right",
    )
    if position not in valid_positions:
        raise ValueError(f"Invalid position. Must be one of: {valid_positions}")
    if not (0.0 <= opacity <= 1.0):
        raise ValueError("opacity must be between 0.0 and 1.0.")
    if watermark_type not in ("text", "image"):
        raise ValueError('watermark_type must be "text" or "image".')
    if watermark_type == "image" and not watermark_image:
        raise ValueError("watermark_image bytes are required for image watermark.")

    # ── Color map ─────────────────────────────────
    color_map = {
        "red": rl_colors.red,
        "blue": rl_colors.blue,
        "grey": rl_colors.grey,
        "gray": rl_colors.grey,
        "black": rl_colors.black,
        "green": rl_colors.green,
        "yellow": rl_colors.yellow,
        "white": rl_colors.white,
    }
    wm_color = color_map.get(color.lower(), rl_colors.red)

    # ── Read PDF ──────────────────────────────────
    reader = PdfReader(io.BytesIO(pdf_bytes))

    if reader.is_encrypted:
        if not password:
            raise ValueError("PDF is encrypted. Provide a password.")
        if reader.decrypt(password) == 0:
            raise ValueError("Wrong password.")

    total_pages = len(reader.pages)

    # ── Pages to watermark ────────────────────────
    if pages is not None:
        invalid = [p for p in pages if not (1 <= p <= total_pages)]
        if invalid:
            raise ValueError(
                f"Invalid page numbers: {invalid}. "
                f"PDF has {total_pages} pages (1-{total_pages})."
            )
        pages_to_watermark = set(pages)
    else:
        pages_to_watermark = set(range(1, total_pages + 1))

    # ── Position calculator ────────────────────────
    def get_position(pos, pw, ph):
        padding = 60
        return {
            "center": (pw / 2, ph / 2),
            "top": (pw / 2, ph - padding),
            "bottom": (pw / 2, padding),
            "top-left": (padding, ph - padding),
            "top-right": (pw - padding, ph - padding),
            "bottom-left": (padding, padding),
            "bottom-right": (pw - padding, padding),
        }.get(pos, (pw / 2, ph / 2))

    # ── Build text watermark page ──────────────────
    def make_text_watermark(pw, ph) -> bytes:
        buf = io.BytesIO()
        c = rl_canvas.Canvas(buf, pagesize=(pw, ph))

        x, y = get_position(position, pw, ph)

        c.saveState()
        c.setFillColor(wm_color, alpha=opacity)
        c.setFont("Helvetica-Bold", font_size)
        c.translate(x, y)
        c.rotate(angle)
        c.drawCentredString(0, 0, watermark_text)
        c.restoreState()

        c.save()
        buf.seek(0)
        return buf.read()

    # ── Build image watermark page ─────────────────
    def make_image_watermark(pw, ph) -> bytes:
        from PIL import Image as PILImage
        from reportlab.lib.utils import ImageReader

        buf = io.BytesIO()
        c = rl_canvas.Canvas(buf, pagesize=(pw, ph))

        # ── Open + convert image ───────────────────
        pil_img = PILImage.open(io.BytesIO(watermark_image))

        # Convert palette/RGBA → RGBA for consistent handling
        if pil_img.mode == "P":
            pil_img = pil_img.convert("RGBA")
        if pil_img.mode not in ("RGB", "RGBA", "L"):
            pil_img = pil_img.convert("RGBA")

        # ── Apply opacity to image ─────────────────
        if pil_img.mode == "RGBA":
            r, g, b, a = pil_img.split()
            # Scale alpha channel by opacity
            a = a.point(lambda px: int(px * opacity))
            pil_img = PILImage.merge("RGBA", (r, g, b, a))
        else:
            # For RGB images, convert to RGBA and apply opacity
            pil_img = pil_img.convert("RGBA")
            r, g, b, a = pil_img.split()
            a = a.point(lambda px: int(px * opacity))
            pil_img = PILImage.merge("RGBA", (r, g, b, a))

        # ── Save processed image to buffer ─────────
        img_buf = io.BytesIO()
        pil_img.save(img_buf, format="PNG")
        img_buf.seek(0)

        # ── Scale image to fit page (max 40%) ──────
        orig_w, orig_h = pil_img.size
        max_w = pw * 0.4
        max_h = ph * 0.4
        scale = min(max_w / orig_w, max_h / orig_h)
        draw_w = orig_w * scale
        draw_h = orig_h * scale

        # ── Get anchor position ────────────────────
        x, y = get_position(position, pw, ph)

        # ── Draw image ─────────────────────────────
        c.saveState()
        c.translate(x, y)
        c.rotate(angle)

        # Draw centered on anchor point
        c.drawImage(
            ImageReader(img_buf),
            x=-draw_w / 2,
            y=-draw_h / 2,
            width=draw_w,
            height=draw_h,
            mask="auto",
        )
        c.restoreState()

        c.save()
        buf.seek(0)
        return buf.read()

    # ── Apply watermark to each page ──────────────
    writer = PdfWriter()

    for i, page in enumerate(reader.pages, start=1):
        if i in pages_to_watermark:
            pw = float(page.mediabox.width)
            ph = float(page.mediabox.height)

            try:
                if watermark_type == "image":
                    wm_bytes = make_image_watermark(pw, ph)
                else:
                    wm_bytes = make_text_watermark(pw, ph)

                wm_page = PdfReader(io.BytesIO(wm_bytes)).pages[0]
                page.merge_page(wm_page)

            except Exception as e:
                # If watermark fails on a page, keep original page
                raise RuntimeError(f"Watermark failed on page {i}: {e}")

        writer.add_page(page)

    # ── Copy metadata ──────────────────────────────
    if reader.metadata:
        writer.add_metadata(reader.metadata)

    # ── Serialize → always return bytes ───────────
    buffer = io.BytesIO()
    writer.write(buffer)
    buffer.seek(0)
    return buffer.read()  # ← always bytes, never None


def merge_pdfs(
    sources: list,
    add_bookmarks: bool = True,
    add_page_numbers: bool = False,
    password: str = None,
) -> dict:
    """
    Merge multiple PDF files into one document.

    Args:
        sources          : list of file objects OR bytes
        add_bookmarks    : add bookmark per file      (default: True)
        add_page_numbers : add page numbers           (default: False)
        password         : output PDF password        (default: None)

    Returns:
        {
            'bytes'      : bytes,
            'total_pages': int,
            'files_merged': int,
            'bookmarks'  : list,
        }
    """
    from pypdf import PdfReader, PdfWriter

    if not sources:
        raise ValueError("No files provided to merge.")
    if len(sources) < 2:
        raise ValueError("At least 2 PDF files are required.")
    if len(sources) > 50:
        raise ValueError("Maximum 50 files can be merged at once.")

    writer = PdfWriter()
    total_pages = 0
    bookmarks = []
    file_metadata = []

    # ── Process each PDF ──────────────────────────
    for i, source in enumerate(sources):
        try:
            if hasattr(source, "read"):
                pdf_bytes = source.read()
                filename = getattr(source, "name", f"file_{i+1}.pdf")
            elif isinstance(source, bytes):
                pdf_bytes = source
                filename = f"file_{i+1}.pdf"
            else:
                raise ValueError(f"Invalid source at index {i}.")

            if not pdf_bytes.startswith(b"%PDF"):
                raise ValueError(f'"{filename}" is not a valid PDF.')

            reader = PdfReader(io.BytesIO(pdf_bytes))

            # ── Decrypt if needed ──────────────────
            if reader.is_encrypted:
                raise ValueError(
                    f'"{filename}" is encrypted. ' f"Please decrypt it first."
                )

            page_count = len(reader.pages)

            # ── Add bookmark for this file ─────────
            if add_bookmarks:
                bookmarks.append(
                    {
                        "title": filename.replace(".pdf", "").replace(".PDF", ""),
                        "page": total_pages + 1,
                        "page_count": page_count,
                        "filename": filename,
                    }
                )

                writer.add_outline_item(
                    title=filename.replace(".pdf", "").replace(".PDF", ""),
                    page_number=total_pages,
                )

            # ── Copy all pages ─────────────────────
            for page in reader.pages:
                writer.add_page(page)

            file_metadata.append(
                {
                    "index": i + 1,
                    "filename": filename,
                    "pages": page_count,
                    "start_page": total_pages + 1,
                    "end_page": total_pages + page_count,
                }
            )

            total_pages += page_count

        except ValueError:
            raise
        except Exception as e:
            raise ValueError(f"Error processing file {i+1}: {e}")

    # ── Add page numbers if requested ──────────────
    if add_page_numbers:
        _add_page_numbers_to_writer(writer, total_pages)

    # ── Add merged metadata ────────────────────────
    writer.add_metadata(
        {
            "/Title": f"Merged PDF ({len(sources)} files)",
            "/Author": "Merge PDF Service",
            "/Subject": ", ".join(m["filename"] for m in file_metadata),
            "/Creator": "PDF Merge Tool",
        }
    )

    # ── Encrypt if password provided ───────────────
    if password:
        writer.encrypt(
            user_password=password,
            owner_password=password,
            use_128bit=True,
        )

    # ── Serialize ─────────────────────────────────
    buffer = io.BytesIO()
    writer.write(buffer)
    buffer.seek(0)
    merged_bytes = buffer.read()

    return {
        "bytes": merged_bytes,
        "total_pages": total_pages,
        "files_merged": len(sources),
        "size_kb": round(len(merged_bytes) / 1024, 2),
        "size_mb": round(len(merged_bytes) / (1024 * 1024), 2),
        "bookmarks": bookmarks,
        "files": file_metadata,
    }


def _add_page_numbers_to_writer(writer, total_pages: int):
    """
    Add page number overlay to each page.
    Uses reportlab to create a transparent overlay.
    """
    from pypdf import PdfReader
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.pagesizes import A4

    for page_num in range(total_pages):
        page = writer.pages[page_num]

        # Get page dimensions
        try:
            width = float(page.mediabox.width)
            height = float(page.mediabox.height)
        except Exception:
            width, height = A4

        # Create overlay with page number
        overlay_buffer = io.BytesIO()
        c = rl_canvas.Canvas(overlay_buffer, pagesize=(width, height))
        c.setFont("Helvetica", 9)
        c.setFillColorRGB(0.5, 0.5, 0.5, alpha=0.8)
        c.drawCentredString(
            width / 2,
            20,
            f"Page {page_num + 1} of {total_pages}",
        )
        c.save()
        overlay_buffer.seek(0)

        # Merge overlay onto page
        overlay_reader = PdfReader(overlay_buffer)
        overlay_page = overlay_reader.pages[0]
        page.merge_page(overlay_page)


def split_pdf(
    source,
    split_by: str = "page",
    pages: list = None,
    chunk_size: int = 1,
    ranges: list = None,
    password: str = None,
) -> list[dict]:
    """
    Split a PDF file into multiple documents.

    Args:
        source     : uploaded file object OR raw bytes
        split_by   : 'page'   → one PDF per page
                     'chunk'  → split every N pages
                     'range'  → extract specific page ranges
                     'pages'  → extract specific individual pages
        pages      : list of page numbers for 'pages' mode
                     e.g. [1, 3, 5] → extract pages 1, 3, 5
        chunk_size : pages per chunk for 'chunk' mode  (default: 1)
        ranges     : list of [start, end] for 'range' mode
                     e.g. [[1,3], [4,6], [7,10]]
        password   : PDF password if encrypted

    Returns:
        list of {
            'index'      : int,
            'filename'   : str,
            'bytes'      : bytes,
            'pages'      : int,
            'start_page' : int,
            'end_page'   : int,
            'size_kb'    : float,
        }
    """
    from pypdf import PdfReader, PdfWriter

    # ── Read source ───────────────────────────────
    if hasattr(source, "read"):
        pdf_bytes = source.read()
        filename = getattr(source, "name", "document.pdf")
    elif isinstance(source, bytes):
        pdf_bytes = source
        filename = "document.pdf"
    else:
        raise ValueError("Invalid source.")

    if not pdf_bytes:
        raise ValueError("Empty file.")

    if not pdf_bytes.startswith(b"%PDF"):
        raise ValueError("Invalid PDF file.")

    base_name = filename.replace(".pdf", "").replace(".PDF", "")

    # ── Validate split_by ─────────────────────────
    valid_modes = ("page", "chunk", "range", "pages")
    if split_by not in valid_modes:
        raise ValueError(f"split_by must be one of: {valid_modes}")

    # ── Open PDF ──────────────────────────────────
    reader = PdfReader(io.BytesIO(pdf_bytes))

    if reader.is_encrypted:
        if not password:
            raise ValueError("PDF is encrypted. Provide a password.")
        if reader.decrypt(password) == 0:
            raise ValueError("Wrong password.")

    total_pages = len(reader.pages)

    if total_pages == 0:
        raise ValueError("PDF has no pages.")

    # ── Helper: write pages to bytes ───────────────
    def pages_to_bytes(
        page_numbers: list,
        start_page: int,
        end_page: int,
        part_index: int,
        label: str = "",
    ) -> dict:
        writer = PdfWriter()

        for pn in page_numbers:
            writer.add_page(reader.pages[pn - 1])

        # Copy metadata
        if reader.metadata:
            writer.add_metadata(dict(reader.metadata))

        buffer = io.BytesIO()
        writer.write(buffer)
        buffer.seek(0)
        raw = buffer.read()

        fname = f"{base_name}_part{part_index}"
        if label:
            fname += f"_{label}"
        fname += ".pdf"

        return {
            "index": part_index,
            "filename": fname,
            "bytes": raw,
            "pages": len(page_numbers),
            "start_page": start_page,
            "end_page": end_page,
            "size_kb": round(len(raw) / 1024, 2),
        }

    results = []

    # ────────────────────────────────────────────────
    # MODE 1: One PDF per page
    # ────────────────────────────────────────────────
    if split_by == "page":
        for page_num in range(1, total_pages + 1):
            result = pages_to_bytes(
                page_numbers=[page_num],
                start_page=page_num,
                end_page=page_num,
                part_index=page_num,
                label=f"page{page_num}",
            )
            results.append(result)

    # ────────────────────────────────────────────────
    # MODE 2: Split every N pages (chunk)
    # ────────────────────────────────────────────────
    elif split_by == "chunk":
        if chunk_size < 1:
            raise ValueError("chunk_size must be at least 1.")

        if chunk_size >= total_pages:
            raise ValueError(
                f"chunk_size ({chunk_size}) must be less than "
                f"total pages ({total_pages})."
            )

        part_index = 1
        for start in range(1, total_pages + 1, chunk_size):
            end = min(start + chunk_size - 1, total_pages)
            page_numbers = list(range(start, end + 1))

            result = pages_to_bytes(
                page_numbers=page_numbers,
                start_page=start,
                end_page=end,
                part_index=part_index,
                label=f"pages{start}-{end}",
            )
            results.append(result)
            part_index += 1

    # ────────────────────────────────────────────────
    # MODE 3: Extract specific page ranges
    # ────────────────────────────────────────────────
    elif split_by == "range":
        if not ranges:
            raise ValueError(
                '"ranges" is required for range mode. ' "e.g. [[1,3], [4,6]]"
            )

        for part_index, r in enumerate(ranges, start=1):
            if isinstance(r, (int, float)):
                start = end = int(r)
            elif isinstance(r, (list, tuple)) and len(r) == 2:
                start, end = int(r[0]), int(r[1])
            else:
                raise ValueError(
                    f"Invalid range: {r}. " f"Use integers or [start, end] pairs."
                )

            if start < 1 or end > total_pages:
                raise ValueError(
                    f"Range [{start}, {end}] is out of bounds. "
                    f"PDF has {total_pages} pages (1-{total_pages})."
                )
            if start > end:
                raise ValueError(
                    f"Invalid range [{start}, {end}]: " f"start must be <= end."
                )

            page_numbers = list(range(start, end + 1))

            result = pages_to_bytes(
                page_numbers=page_numbers,
                start_page=start,
                end_page=end,
                part_index=part_index,
                label=f"pages{start}-{end}",
            )
            results.append(result)

    # ────────────────────────────────────────────────
    # MODE 4: Extract specific individual pages
    # ────────────────────────────────────────────────
    elif split_by == "pages":
        if not pages:
            raise ValueError('"pages" is required for pages mode. ' "e.g. [1, 3, 5]")

        # ── Validate page numbers ─────────────────
        invalid = [p for p in pages if not (1 <= p <= total_pages)]
        if invalid:
            raise ValueError(
                f"Invalid page numbers: {invalid}. "
                f"PDF has {total_pages} pages (1-{total_pages})."
            )

        # ── Option A: each page as separate PDF ────
        for part_index, page_num in enumerate(sorted(pages), start=1):
            result = pages_to_bytes(
                page_numbers=[page_num],
                start_page=page_num,
                end_page=page_num,
                part_index=part_index,
                label=f"page{page_num}",
            )
            results.append(result)

    if not results:
        raise ValueError("No content found to split.")

    return results


def latex_to_pdf(
    source,
    filename   : str = 'document.tex',
    engine     : str = 'pdflatex',
    runs       : int = 2,
    encoding   : str = 'utf-8',
) -> dict:
    """
    Convert LaTeX (.tex) → PDF using available LaTeX engine.

    Strategy:
        1. Try pdflatex / xelatex / lualatex (system LaTeX)
        2. Fallback → pure Python renderer (basic, no math)

    Args:
        source   : file object | raw LaTeX string | bytes
        filename : original filename
        engine   : pdflatex | xelatex | lualatex    (default: pdflatex)
        runs     : number of compilation runs        (default: 2)
                   2 runs needed for TOC/refs
        encoding : input encoding                   (default: utf-8)

    Returns:
        {
            'bytes'    : bytes,
            'engine'   : str,
            'log'      : str,
            'warnings' : list,
            'errors'   : list,
            'pages'    : int,
            'method'   : 'latex' | 'fallback',
        }
    """
    import subprocess
    import tempfile
    import shutil
    import re

    # ── Read source ───────────────────────────────
    if hasattr(source, 'read'):
        raw = source.read()
        if isinstance(raw, bytes):
            raw = raw.decode(encoding, errors='replace')
    elif isinstance(source, bytes):
        raw = source.decode(encoding, errors='replace')
    elif isinstance(source, str):
        raw = source
    else:
        raise ValueError('source must be a string, bytes, or file object.')

    if not raw.strip():
        raise ValueError('Empty input.')

    # ── Validate engine ───────────────────────────
    valid_engines = ('pdflatex', 'xelatex', 'lualatex')
    if engine not in valid_engines:
        engine = 'pdflatex'

    # ── Try LaTeX engines ─────────────────────────
    engines_to_try = [engine] + [
        e for e in valid_engines if e != engine
    ]

    for eng in engines_to_try:
        eng_path = shutil.which(eng)
        if eng_path:
            try:
                result = _compile_latex(
                    raw, filename, eng_path, eng, runs
                )
                return result
            except Exception:
                continue

    # ── Fallback: pure Python ─────────────────────
    return _latex_to_pdf_fallback(raw, filename)


def _compile_latex(
    latex_str : str,
    filename  : str,
    engine_path: str,
    engine_name: str,
    runs      : int,
) -> dict:
    """
    Compile LaTeX using system engine in a temp directory.
    """
    import subprocess
    import tempfile
    import re

    base_name = filename.replace('.tex', '').replace('.TEX', '')

    with tempfile.TemporaryDirectory() as tmp_dir:
        tex_path = os.path.join(tmp_dir, f'{base_name}.tex')
        pdf_path = os.path.join(tmp_dir, f'{base_name}.pdf')
        log_path = os.path.join(tmp_dir, f'{base_name}.log')

        # ── Write .tex file ────────────────────────
        with open(tex_path, 'w', encoding='utf-8') as f:
            f.write(latex_str)

        # ── Compile (multiple runs for TOC/refs) ───
        full_log = ''
        for run_num in range(runs):
            result = subprocess.run(
                [
                    engine_path,
                    '-interaction=nonstopmode',
                    '-halt-on-error',
                    f'-output-directory={tmp_dir}',
                    tex_path,
                ],
                capture_output=True,
                text=True,
                timeout=60,
                cwd=tmp_dir,
            )
            full_log += result.stdout + result.stderr

        # ── Read log file ──────────────────────────
        log_content = ''
        if os.path.exists(log_path):
            with open(log_path, 'r', encoding='utf-8', errors='replace') as f:
                log_content = f.read()

        # ── Parse warnings and errors ──────────────
        warnings = re.findall(
            r'LaTeX Warning: (.+?)(?:\n|$)', log_content
        )
        errors = re.findall(
            r'! (.+?)(?:\n|$)', log_content
        )

        # ── Check PDF was generated ────────────────
        if not os.path.exists(pdf_path):
            error_msg = errors[0] if errors else 'Compilation failed.'
            raise RuntimeError(
                f'{engine_name} failed: {error_msg}\n\n'
                f'Log:\n{log_content[-2000:]}'
            )

        with open(pdf_path, 'rb') as f:
            pdf_bytes = f.read()

    # ── Get page count ─────────────────────────────
    pages = _count_pdf_pages(pdf_bytes)

    return {
        'bytes'   : pdf_bytes,
        'engine'  : engine_name,
        'log'     : log_content[-3000:] if log_content else full_log[-3000:],
        'warnings': warnings[:20],
        'errors'  : errors[:10],
        'pages'   : pages,
        'method'  : 'latex',
        'size_kb' : round(len(pdf_bytes) / 1024, 2),
        'size_mb' : round(len(pdf_bytes) / (1024 * 1024), 2),
    }


def _latex_to_pdf_fallback(latex_str: str, filename: str) -> dict:
    """
    Pure Python fallback — converts LaTeX → PDF without system LaTeX.
    Extracts text content from LaTeX and renders as PDF using reportlab.

    Handles: basic text, sections, lists, tables, bold, italic, code
    Does NOT handle: math equations, complex macros, custom packages
    """
    from reportlab.lib.pagesizes  import A4
    from reportlab.lib.styles     import ParagraphStyle
    from reportlab.lib.units      import inch
    from reportlab.lib            import colors
    from reportlab.lib.enums      import TA_LEFT, TA_CENTER, TA_JUSTIFY
    from reportlab.platypus       import (
        SimpleDocTemplate, Paragraph, Spacer,
        Table, TableStyle, PageBreak, HRFlowable,
    )
    import re

    # ── Strip LaTeX preamble ──────────────────────
    # Extract only document body
    body_match = re.search(
        r'\\begin\{document\}(.+?)\\end\{document\}',
        latex_str,
        re.DOTALL,
    )
    body = body_match.group(1).strip() if body_match else latex_str

    # ── Extract title ─────────────────────────────
    title_match = re.search(r'\\title\{(.+?)\}', latex_str, re.DOTALL)
    doc_title   = title_match.group(1).strip() if title_match else \
                  filename.replace('.tex', '')

    # ── Strip LaTeX commands ───────────────────────
    def strip_latex(text: str) -> str:
        """Convert LaTeX markup to plain/HTML for reportlab."""
        # Bold
        text = re.sub(r'\\textbf\{(.+?)\}',    r'<b>\1</b>',  text)
        # Italic
        text = re.sub(r'\\textit\{(.+?)\}',    r'<i>\1</i>',  text)
        text = re.sub(r'\\emph\{(.+?)\}',      r'<i>\1</i>',  text)
        # Underline
        text = re.sub(r'\\underline\{(.+?)\}', r'<u>\1</u>',  text)
        # Strikethrough
        text = re.sub(r'\\sout\{(.+?)\}',      r'<strike>\1</strike>', text)
        # Monospace
        text = re.sub(r'\\texttt\{(.+?)\}',    r'<font name="Courier">\1</font>', text)
        # href links
        text = re.sub(
            r'\\href\{[^}]+\}\{(.+?)\}',
            r'<font color="blue"><u>\1</u></font>',
            text,
        )
        # Unescape LaTeX special chars
        text = text.replace('\\&',  '&amp;')
        text = text.replace('\\%',  '%')
        text = text.replace('\\$',  '$')
        text = text.replace('\\#',  '#')
        text = text.replace('\\_',  '_')
        text = text.replace('\\{',  '{')
        text = text.replace('\\}',  '}')
        text = text.replace('\\textbackslash{}', '\\')
        text = text.replace('\\textless{}',  '<')
        text = text.replace('\\textgreater{}', '>')
        text = text.replace('\\textbar{}',   '|')
        text = text.replace('\\textasciitilde{}', '~')
        text = text.replace('\\^{}',  '^')
        text = text.replace('---',    '—')
        text = text.replace('--',     '–')
        text = text.replace('``',     '"')
        text = text.replace("''",     '"')
        # Remove remaining unknown commands
        text = re.sub(r'\\[a-zA-Z]+\*?\{([^}]*)\}', r'\1', text)
        text = re.sub(r'\\[a-zA-Z]+\*?',             '',    text)
        return text.strip()

    # ── Styles ────────────────────────────────────
    def make_style(name, font, size, bold=False,
                   before=6, after=6, align=TA_JUSTIFY, color=None):
        s = ParagraphStyle(
            name,
            fontName   =font,
            fontSize   =size,
            leading    =size * 1.4,
            spaceBefore=before,
            spaceAfter =after,
            alignment  =align,
        )
        if color:
            s.textColor = color
        return s

    styles = {
        'title'   : make_style('title',    'Helvetica-Bold', 20,
                               before=12, after=12, align=TA_CENTER,
                               color=colors.HexColor('#1F4E79')),
        'h1'      : make_style('h1',       'Helvetica-Bold', 16,
                               before=14, after=8,
                               color=colors.HexColor('#2E75B6')),
        'h2'      : make_style('h2',       'Helvetica-Bold', 14,
                               before=12, after=6,
                               color=colors.HexColor('#2E75B6')),
        'h3'      : make_style('h3',       'Helvetica-Bold', 12,
                               before=10, after=4,
                               color=colors.HexColor('#1F4E79')),
        'h4'      : make_style('h4',       'Helvetica-Bold', 11,
                               before=8,  after=4,
                               color=colors.HexColor('#1F4E79')),
        'body'    : make_style('body',     'Helvetica',      11,
                               before=4,  after=4),
        'bullet'  : make_style('bullet',   'Helvetica',      11,
                               before=2,  after=2,
                               align=TA_LEFT),
        'code'    : make_style('code',     'Courier',         9,
                               before=4,  after=4,
                               align=TA_LEFT),
        'quote'   : make_style('quote',    'Helvetica-Oblique', 11,
                               before=6,  after=6,
                               align=TA_LEFT,
                               color=colors.HexColor('#555555')),
    }

    # ── Parse body content ─────────────────────────
    story  = []
    buffer = io.BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize    =A4,
        rightMargin =inch,
        leftMargin  =inch,
        topMargin   =inch,
        bottomMargin=inch,
        title       =doc_title,
    )

    # Remove comments
    body = re.sub(r'%[^\n]*', '', body)

    # Add title
    story.append(Paragraph(
        doc_title.replace('&', '&amp;').replace('<', '&lt;'),
        styles['title'],
    ))
    story.append(Spacer(1, 12))

    # Tokenize content
    lines   = body.split('\n')
    i       = 0
    total   = len(lines)
    in_list = False
    list_items = []

    def flush_list(items):
        if not items:
            return
        for item in items:
            try:
                story.append(Paragraph(f'• {item}', styles['bullet']))
            except Exception:
                story.append(Paragraph(f'• {re.sub("<[^>]+>","",item)}',
                                       styles['bullet']))
        story.append(Spacer(1, 4))

    while i < total:
        line     = lines[i].strip()

        # ── Skip empty ─────────────────────────────
        if not line:
            if list_items:
                flush_list(list_items)
                list_items = []
            story.append(Spacer(1, 6))
            i += 1
            continue

        # ── maketitle ──────────────────────────────
        if '\\maketitle' in line:
            i += 1
            continue

        # ── tableofcontents ────────────────────────
        if '\\tableofcontents' in line:
            story.append(Paragraph(
                '<b>Table of Contents</b>', styles['h2']
            ))
            story.append(Spacer(1, 8))
            i += 1
            continue

        # ── newpage ────────────────────────────────
        if '\\newpage' in line or '\\clearpage' in line:
            flush_list(list_items)
            list_items = []
            story.append(PageBreak())
            i += 1
            continue

        # ── hrule ──────────────────────────────────
        if '\\hrule' in line:
            flush_list(list_items)
            list_items = []
            story.append(HRFlowable(
                width='100%', thickness=1,
                color=colors.HexColor('#2E75B6'),
                spaceAfter=6, spaceBefore=6,
            ))
            i += 1
            continue

        # ── Sections ───────────────────────────────
        m = re.match(r'\\(chapter|section|subsection|subsubsection|paragraph)\*?\{(.+?)\}', line)
        if m:
            flush_list(list_items)
            list_items = []
            cmd  = m.group(1)
            text = strip_latex(m.group(2))
            style_map = {
                'chapter'        : styles['h1'],
                'section'        : styles['h1'],
                'subsection'     : styles['h2'],
                'subsubsection'  : styles['h3'],
                'paragraph'      : styles['h4'],
            }
            s = style_map.get(cmd, styles['h1'])
            try:
                story.append(Paragraph(text, s))
            except Exception:
                story.append(Paragraph(
                    re.sub('<[^>]+>', '', text), s
                ))
            i += 1
            continue

        # ── itemize / enumerate ────────────────────
        if '\\begin{itemize}' in line or '\\begin{enumerate}' in line:
            i += 1
            continue
        if '\\end{itemize}' in line or '\\end{enumerate}' in line:
            flush_list(list_items)
            list_items = []
            i += 1
            continue
        if line.startswith('\\item'):
            text = re.sub(r'^\\item\s*', '', line)
            list_items.append(strip_latex(text))
            i += 1
            continue

        # ── verbatim / lstlisting ──────────────────
        if '\\begin{verbatim}' in line or '\\begin{lstlisting}' in line:
            flush_list(list_items)
            list_items = []
            i += 1
            code_lines = []
            end_tag = 'verbatim' if 'verbatim' in line else 'lstlisting'
            while i < total and f'\\end{{{end_tag}}}' not in lines[i]:
                code_lines.append(lines[i])
                i += 1
            code_text = '\n'.join(code_lines)
            try:
                story.append(Paragraph(
                    code_text.replace('&', '&amp;')
                             .replace('<', '&lt;')
                             .replace('>', '&gt;')
                             .replace('\n', '<br/>'),
                    styles['code'],
                ))
            except Exception:
                pass
            i += 1
            continue

        # ── quote / blockquote ─────────────────────
        if '\\begin{quote}' in line or '\\begin{quotation}' in line:
            flush_list(list_items)
            list_items = []
            i += 1
            quote_lines = []
            end_tag = 'quote' if 'quote' in line else 'quotation'
            while i < total and f'\\end{{{end_tag}}}' not in lines[i]:
                quote_lines.append(strip_latex(lines[i]))
                i += 1
            quote_text = ' '.join(quote_lines)
            try:
                story.append(Paragraph(quote_text, styles['quote']))
            except Exception:
                story.append(Paragraph(
                    re.sub('<[^>]+>', '', quote_text),
                    styles['quote'],
                ))
            i += 1
            continue

        # ── table ──────────────────────────────────
        if '\\begin{tabular}' in line:
            flush_list(list_items)
            list_items = []
            i += 1
            tbl_lines = []
            while i < total and '\\end{tabular}' not in lines[i]:
                tbl_lines.append(lines[i])
                i += 1

            tbl_data = []
            for tl in tbl_lines:
                tl = tl.strip()
                if tl in ('\\hline', ''):
                    continue
                if tl.endswith('\\\\'):
                    tl = tl[:-2].strip()
                cells = [
                    strip_latex(c.strip())
                    for c in tl.split('&')
                ]
                if cells:
                    tbl_data.append(cells)

            if tbl_data:
                num_cols = max(len(r) for r in tbl_data)
                tbl_data = [
                    r + [''] * (num_cols - len(r))
                    for r in tbl_data
                ]
                col_w    = (A4[0] - 2 * inch) / num_cols

                tbl_rows = []
                for r_idx, row in enumerate(tbl_data):
                    tbl_row = []
                    for cell in row:
                        try:
                            p = Paragraph(
                                cell if cell else '',
                                ParagraphStyle(
                                    f'tc_{r_idx}',
                                    fontName='Helvetica-Bold'
                                             if r_idx == 0
                                             else 'Helvetica',
                                    fontSize=9,
                                    textColor=colors.white
                                              if r_idx == 0
                                              else colors.black,
                                    leading=12,
                                )
                            )
                        except Exception:
                            p = Paragraph(
                                re.sub('<[^>]+>', '', cell),
                                ParagraphStyle(
                                    f'tc_plain_{r_idx}',
                                    fontName='Helvetica',
                                    fontSize=9,
                                    leading=12,
                                )
                            )
                        tbl_row.append(p)
                    tbl_rows.append(tbl_row)

                pdf_tbl = Table(
                    tbl_rows,
                    colWidths=[col_w] * num_cols,
                )
                pdf_tbl.setStyle(TableStyle([
                    ('BACKGROUND',    (0,0), (-1,0),
                     colors.HexColor('#2E75B6')),
                    ('TEXTCOLOR',     (0,0), (-1,0), colors.white),
                    ('ROWBACKGROUNDS',(0,1), (-1,-1),
                     [colors.HexColor('#DCE6F1'), colors.white]),
                    ('GRID',   (0,0), (-1,-1), 0.5,
                     colors.HexColor('#AAAAAA')),
                    ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                    ('PADDING',(0,0), (-1,-1), 6),
                ]))
                story.append(Spacer(1, 8))
                story.append(pdf_tbl)
                story.append(Spacer(1, 12))
            i += 1
            continue

        # ── Skip other environments ────────────────
        if line.startswith('\\begin{') or line.startswith('\\end{'):
            i += 1
            continue

        # ── Skip pure LaTeX commands ───────────────
        if re.match(r'^\\[a-zA-Z]+\s*$', line):
            i += 1
            continue

        # ── Regular paragraph ──────────────────────
        flush_list(list_items)
        list_items = []
        text = strip_latex(line)
        if text.strip():
            try:
                story.append(Paragraph(text, styles['body']))
            except Exception:
                story.append(Paragraph(
                    re.sub('<[^>]+>', '', text),
                    styles['body'],
                ))
        i += 1

    flush_list(list_items)

    # ── Build PDF ─────────────────────────────────
    doc.build(story)
    buffer.seek(0)
    pdf_bytes = buffer.read()

    pages = _count_pdf_pages(pdf_bytes)

    return {
        'bytes'   : pdf_bytes,
        'engine'  : 'reportlab (fallback)',
        'log'     : 'No LaTeX engine found — used pure Python fallback.',
        'warnings': ['Math equations, custom packages not supported in fallback mode.'],
        'errors'  : [],
        'pages'   : pages,
        'method'  : 'fallback',
        'size_kb' : round(len(pdf_bytes) / 1024, 2),
        'size_mb' : round(len(pdf_bytes) / (1024 * 1024), 2),
    }


def _count_pdf_pages(pdf_bytes: bytes) -> int:
    """Count pages in PDF bytes using pymupdf."""
    try:
        import fitz
        doc = fitz.open(stream=pdf_bytes, filetype='pdf')
        pages = len(doc)
        doc.close()
        return pages
    except Exception:
        return 0


def pdf_to_mobi(
    source,
    filename: str = "document.pdf",
    title: str = "",
    author: str = "",
    password: str = None,
) -> dict:
    """
    Convert PDF → MOBI (Kindle format).

    Strategy:
        1. ebook-convert (Calibre) → best quality
        2. Fallback → extract text + ebooklib → basic MOBI

    Args:
        source   : uploaded file object OR raw bytes
        filename : original filename
        title    : book title          (default: filename)
        author   : book author         (default: '')
        password : PDF password        (default: None)

    Returns:
        {
            'bytes'   : bytes,
            'method'  : str,
            'pages'   : int,
            'size_kb' : float,
        }
    """
    import shutil
    import tempfile
    import subprocess

    # ── Read PDF ──────────────────────────────────
    if hasattr(source, "read"):
        pdf_bytes = source.read()
        filename = getattr(source, "name", filename)
    elif isinstance(source, bytes):
        pdf_bytes = source
    else:
        raise ValueError("Invalid source.")

    if not pdf_bytes.startswith(b"%PDF"):
        raise ValueError("Invalid PDF file.")

    title = title or filename.replace(".pdf", "").replace(".PDF", "")
    author = author or "Unknown"

    # ── Try Calibre ebook-convert ─────────────────
    calibre = shutil.which("ebook-convert")
    if calibre:
        return _pdf_to_mobi_calibre(pdf_bytes, filename, title, author, calibre)

    # ── Fallback: pymupdf + ebooklib ──────────────
    return _pdf_to_mobi_python(pdf_bytes, filename, title, author, password)


def _pdf_to_mobi_calibre(
    pdf_bytes: bytes,
    filename: str,
    title: str,
    author: str,
    calibre: str,
) -> dict:
    """Convert PDF → MOBI using Calibre ebook-convert."""
    import subprocess
    import tempfile

    base = filename.replace(".pdf", "").replace(".PDF", "")

    with tempfile.TemporaryDirectory() as tmp:
        pdf_path = os.path.join(tmp, f"{base}.pdf")
        mobi_path = os.path.join(tmp, f"{base}.mobi")

        with open(pdf_path, "wb") as f:
            f.write(pdf_bytes)

        subprocess.run(
            [
                calibre,
                pdf_path,
                mobi_path,
                "--title",
                title,
                "--authors",
                author,
                "--output-profile",
                "kindle",
            ],
            capture_output=True,
            timeout=120,
            check=True,
        )

        with open(mobi_path, "rb") as f:
            mobi_bytes = f.read()

    return {
        "bytes": mobi_bytes,
        "method": "calibre",
        "size_kb": round(len(mobi_bytes) / 1024, 2),
        "size_mb": round(len(mobi_bytes) / (1024 * 1024), 2),
    }


def _pdf_to_mobi_python(
    pdf_bytes: bytes,
    filename: str,
    title: str,
    author: str,
    password: str = None,
) -> dict:
    """
    Pure Python fallback — PDF text → EPUB → MOBI bytes.
    Uses pymupdf for extraction + ebooklib for packaging.
    Note: Images not included in fallback mode.
    """
    import fitz
    from ebooklib import epub

    # ── Open PDF ──────────────────────────────────
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    if doc.is_encrypted:
        if not password:
            raise ValueError("PDF is encrypted. Provide a password.")
        if not doc.authenticate(password):
            raise ValueError("Wrong password.")

    total_pages = len(doc)

    # ── Extract text per page ─────────────────────
    pages_text = []
    for page_num in range(total_pages):
        page = doc[page_num]
        text = page.get_text("text").strip()
        if text:
            pages_text.append(
                {
                    "page": page_num + 1,
                    "text": text,
                }
            )
    doc.close()

    if not pages_text:
        raise ValueError("No text found in PDF.")

    # ── Build EPUB (MOBI-compatible) ──────────────
    book = epub.EpubBook()
    book.set_identifier(f"id_{filename}")
    book.set_title(title)
    book.set_language("en")
    book.add_author(author)

    # ── CSS ───────────────────────────────────────
    css = epub.EpubItem(
        uid="style",
        file_name="style.css",
        media_type="text/css",
        content=b"""
            body { font-family: Georgia, serif; margin: 2em; line-height: 1.6; }
            h1   { color: #2E75B6; border-bottom: 1px solid #ccc; padding-bottom: 0.3em; }
            p    { margin: 0.8em 0; text-align: justify; }
            .page-break { page-break-after: always; }
        """,
    )
    book.add_item(css)

    # ── Chapters (one per page) ────────────────────
    chapters = []
    spine = ["nav"]

    for p in pages_text:
        chapter = epub.EpubHtml(
            title=f'Page {p["page"]}',
            file_name=f'page_{p["page"]}.xhtml',
            lang="en",
        )

        # Convert plain text → HTML paragraphs
        html_paras = "".join(
            f"<p>{line.strip()}</p>" for line in p["text"].splitlines() if line.strip()
        )

        chapter.content = (
            f'<html><head><link rel="stylesheet" href="style.css"/></head>'
            f'<body><h1>Page {p["page"]}</h1>{html_paras}</body></html>'
        ).encode("utf-8")

        chapter.add_item(css)
        book.add_item(chapter)
        chapters.append(chapter)
        spine.append(chapter)

    # ── TOC + Spine ───────────────────────────────
    book.toc = tuple(chapters)
    book.spine = spine
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    # ── Write EPUB to bytes ───────────────────────
    epub_buffer = io.BytesIO()
    epub.write_epub(epub_buffer, book)
    epub_buffer.seek(0)
    epub_bytes = epub_buffer.read()

    # ── Note: true MOBI needs Calibre ─────────────
    # ebooklib cannot write native MOBI — we return EPUB
    # which is compatible with most Kindle apps via Send to Kindle
    return {
        "bytes": epub_bytes,
        "method": "python (epub — kindle compatible)",
        "size_kb": round(len(epub_bytes) / 1024, 2),
        "size_mb": round(len(epub_bytes) / (1024 * 1024), 2),
        "pages": total_pages,
        "note": (
            "Calibre not found. Output is EPUB format which is "
            "compatible with Kindle apps. "
            "Install Calibre for native MOBI output."
        ),
    }


def html_to_epub(
    source,
    filename: str = "document.html",
    title: str = "",
    author: str = "",
    language: str = "en",
    description: str = "",
    publisher: str = "",
    cover_image: bytes = None,
    encoding: str = "utf-8",
) -> dict:
    """
    Convert HTML (file or text) → EPUB.

    Args:
        source      : file object | raw HTML string | bytes
        filename    : original filename
        title       : book title                (default: from <title> tag)
        author      : book author               (default: from <meta author>)
        language    : book language code        (default: en)
        description : book description          (default: from <meta description>)
        publisher   : book publisher            (default: '')
        cover_image : cover image bytes PNG/JPG (default: None)
        encoding    : input encoding            (default: utf-8)

    Returns:
        {
            'bytes'       : bytes,
            'title'       : str,
            'author'      : str,
            'chapters'    : int,
            'images'      : int,
            'size_kb'     : float,
            'toc'         : list,
        }
    """
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup, Tag
    import re
    import base64
    import hashlib

    # ── Read source ───────────────────────────────
    if hasattr(source, "read"):
        raw = source.read()
        if isinstance(raw, bytes):
            raw = raw.decode(encoding, errors="replace")
    elif isinstance(source, bytes):
        raw = source.decode(encoding, errors="replace")
    elif isinstance(source, str):
        raw = source
    else:
        raise ValueError("source must be a string, bytes, or file object.")

    if not raw.strip():
        raise ValueError("Empty input.")

    # ── Parse HTML ────────────────────────────────
    soup = BeautifulSoup(raw, "lxml")

    # ── Extract metadata from HTML ────────────────
    if not title:
        title_tag = soup.find("title")
        title = (
            title_tag.get_text(strip=True)
            if title_tag
            else filename.replace(".html", "").replace(".htm", "")
        )

    if not author:
        author_meta = soup.find("meta", attrs={"name": "author"})
        author = author_meta.get("content", "") if author_meta else ""

    if not description:
        desc_meta = soup.find("meta", attrs={"name": "description"})
        description = desc_meta.get("content", "") if desc_meta else ""

    # ── Extract inline CSS ─────────────────────────
    inline_css = ""
    for style_tag in soup.find_all("style"):
        inline_css += style_tag.get_text() + "\n"
        style_tag.decompose()

    # ── Base CSS ───────────────────────────────────
    base_css = (
        """
body  {
    font-family   : Georgia, serif;
    font-size     : 1em;
    line-height   : 1.7;
    margin        : 1em;
    color         : #222;
    text-align    : justify;
}
h1, h2, h3, h4, h5, h6 {
    font-family   : Arial, sans-serif;
    line-height   : 1.3;
    margin        : 1em 0 0.5em;
    color         : #111;
}
h1 { font-size: 1.6em; border-bottom: 2px solid #ccc; padding-bottom: 0.3em; }
h2 { font-size: 1.4em; }
h3 { font-size: 1.2em; }
p  { margin: 0.6em 0; }
a  { color: #2E75B6; text-decoration: underline; }
img { max-width: 100%; height: auto; display: block; margin: 1em auto; }
table {
    width         : 100%;
    border-collapse: collapse;
    margin        : 1em 0;
}
th, td {
    border   : 1px solid #ccc;
    padding  : 0.4em 0.6em;
    text-align: left;
}
th { background: #f0f0f0; font-weight: bold; }
tr:nth-child(even) { background: #f9f9f9; }
blockquote {
    border-left  : 4px solid #2E75B6;
    margin       : 1em 2em;
    padding      : 0.5em 1em;
    color        : #555;
    font-style   : italic;
}
code, pre {
    font-family  : Courier, monospace;
    font-size    : 0.9em;
    background   : #f5f5f5;
    padding      : 0.2em 0.4em;
    border-radius: 3px;
}
pre {
    padding      : 0.8em;
    overflow-x   : auto;
    white-space  : pre-wrap;
}
ul, ol { margin: 0.5em 0 0.5em 1.5em; }
li     { margin: 0.3em 0; }
hr     { border: none; border-top: 1px solid #ccc; margin: 1.5em 0; }
"""
        + inline_css
    )

    # ── Create EPUB ───────────────────────────────
    book = epub.EpubBook()
    book.set_title(title)
    book.set_language(language)
    book.set_identifier(f"html-{hashlib.md5(raw.encode()).hexdigest()[:12]}")
    if author:
        book.add_author(author)
    if description:
        book.add_metadata("DC", "description", description)
    if publisher:
        book.add_metadata("DC", "publisher", publisher)

    # ── Add CSS ───────────────────────────────────
    css_item = epub.EpubItem(
        uid="style_main",
        file_name="styles/main.css",
        media_type="text/css",
        content=base_css.encode("utf-8"),
    )
    book.add_item(css_item)

    # ── Add cover image ───────────────────────────
    if cover_image:
        try:
            from PIL import Image as PILImage

            img = PILImage.open(io.BytesIO(cover_image))
            ext = img.format.lower() if img.format else "jpeg"
            mt = f"image/{ext}"
            cov_item = epub.EpubItem(
                uid="cover_img",
                file_name=f"images/cover.{ext}",
                media_type=mt,
                content=cover_image,
            )
            book.add_item(cov_item)
            book.set_cover(f"images/cover.{ext}", cover_image)
        except Exception:
            pass

    # ── Extract embedded images ────────────────────
    image_count = 0
    image_map = {}  # src → epub filename

    for img_tag in soup.find_all("img"):
        src = img_tag.get("src", "")
        if not src or src in image_map:
            continue

        img_bytes = None
        media_type = "image/jpeg"

        # ── Base64 embedded image ──────────────────
        if src.startswith("data:image"):
            try:
                header, b64data = src.split(",", 1)
                mt_match = re.search(r"data:([^;]+)", header)
                media_type = mt_match.group(1) if mt_match else "image/jpeg"
                img_bytes = base64.b64decode(b64data)
            except Exception:
                continue

        if img_bytes:
            image_count += 1
            ext = media_type.split("/")[-1].replace("jpeg", "jpg")
            epub_src = f"images/img_{image_count}.{ext}"

            img_item = epub.EpubItem(
                uid=f"img_{image_count}",
                file_name=epub_src,
                media_type=media_type,
                content=img_bytes,
            )
            book.add_item(img_item)
            image_map[src] = epub_src
            img_tag["src"] = epub_src

    # ── Split HTML into chapters ───────────────────
    # Strategy: split on H1/H2 headings
    body = soup.find("body") or soup

    chapters_data = _split_html_into_chapters(body, title)

    if not chapters_data:
        # No headings — treat as single chapter
        chapters_data = [
            {
                "title": title,
                "content": str(body),
            }
        ]

    # ── Build EPUB chapters ────────────────────────
    epub_chapters = []
    toc = []

    for ch_idx, ch in enumerate(chapters_data):
        ch_title = ch["title"]
        ch_content = ch["content"]

        # Fix image src in chapter content
        for old_src, new_src in image_map.items():
            ch_content = ch_content.replace(old_src, new_src)

        # Build XHTML content
        xhtml = (
            '<?xml version="1.0" encoding="utf-8"?>\n'
            "<!DOCTYPE html>\n"
            '<html xmlns="http://www.w3.org/1999/xhtml">\n'
            "<head>\n"
            f"  <title>{_escape_xml(ch_title)}</title>\n"
            '  <link rel="stylesheet" type="text/css" '
            'href="../styles/main.css"/>\n'
            "</head>\n"
            "<body>\n"
            f"{ch_content}\n"
            "</body>\n"
            "</html>"
        )

        chapter = epub.EpubHtml(
            title=ch_title,
            file_name=f"chapters/chapter_{ch_idx:03d}.xhtml",
            lang=language,
            content=xhtml.encode("utf-8"),
        )
        chapter.add_item(css_item)
        book.add_item(chapter)
        epub_chapters.append(chapter)

        toc.append(
            epub.Link(
                f"chapters/chapter_{ch_idx:03d}.xhtml",
                ch_title,
                f"chapter_{ch_idx}",
            )
        )

    # ── Set spine + TOC ───────────────────────────
    book.spine = ["nav"] + epub_chapters
    book.toc = toc
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    # ── Write EPUB ────────────────────────────────
    buf = io.BytesIO()
    epub.write_epub(buf, book, {})
    buf.seek(0)
    epub_bytes = buf.read()

    return {
        "bytes": epub_bytes,
        "title": title,
        "author": author,
        "language": language,
        "chapters": len(epub_chapters),
        "images": image_count,
        "size_kb": round(len(epub_bytes) / 1024, 2),
        "size_mb": round(len(epub_bytes) / (1024 * 1024), 2),
        "toc": [t.title for t in toc],
        "description": description,
    }


def _split_html_into_chapters(body, default_title: str) -> list:
    """
    Split HTML body into chapters based on H1/H2 headings.
    Returns list of { 'title': str, 'content': str }
    """
    from bs4 import Tag, NavigableString
    import re

    chapters = []
    current_title = default_title
    current_parts = []

    for element in body.children:
        if not isinstance(element, Tag):
            if isinstance(element, NavigableString) and str(element).strip():
                current_parts.append(str(element))
            continue

        tag = element.name.lower() if element.name else ""

        if tag in ("h1", "h2"):
            # Save previous chapter
            if current_parts:
                chapters.append(
                    {
                        "title": current_title,
                        "content": "".join(current_parts),
                    }
                )
            current_title = element.get_text(strip=True) or current_title
            current_parts = [str(element)]
        else:
            current_parts.append(str(element))

    # Save last chapter
    if current_parts:
        chapters.append(
            {
                "title": current_title,
                "content": "".join(current_parts),
            }
        )

    return chapters


def _escape_xml(text: str) -> str:
    """Escape XML special characters."""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&apos;")
    )


def pdf_to_epub(
    source,
    filename: str = "document.pdf",
    title: str = "",
    author: str = "",
    language: str = "en",
    description: str = "",
    publisher: str = "",
    password: str = None,
    image_dpi: int = 150,
    image_quality: int = 75,
    extract_text: bool = True,
) -> dict:
    """
    Convert PDF → EPUB.

    Strategy:
        1. Extract text blocks per page using pymupdf
        2. Render each page as JPEG image
        3. Build EPUB with text + image per chapter (page)
        4. Full metadata + TOC + CSS

    Args:
        source         : file object OR raw bytes
        filename       : original filename
        title          : book title           (default: filename)
        author         : book author          (default: '')
        language       : language code        (default: en)
        description    : book description     (default: '')
        publisher      : publisher name       (default: '')
        password       : PDF password         (default: None)
        image_dpi      : page render DPI      (default: 150)
        image_quality  : JPEG quality 1-95    (default: 75)
        extract_text   : extract text content (default: True)

    Returns:
        {
            'bytes'   : bytes,
            'title'   : str,
            'author'  : str,
            'pages'   : int,
            'chapters': int,
            'size_kb' : float,
            'size_mb' : float,
            'toc'     : list,
        }
    """
    import fitz
    import ebooklib
    from ebooklib import epub
    from PIL import Image
    import hashlib
    import re

    # ── Read source ───────────────────────────────
    if hasattr(source, "read"):
        pdf_bytes = source.read()
        if not title:
            title = (
                getattr(source, "name", filename)
                .replace(".pdf", "")
                .replace(".PDF", "")
            )
    elif isinstance(source, bytes):
        pdf_bytes = source
    else:
        raise ValueError("source must be a file object or bytes.")

    if not pdf_bytes:
        raise ValueError("Empty file.")

    if not pdf_bytes.startswith(b"%PDF"):
        raise ValueError("Invalid PDF file.")

    if not title:
        title = filename.replace(".pdf", "").replace(".PDF", "")

    # ── Decrypt if needed ──────────────────────────
    if password:
        from pypdf import PdfReader, PdfWriter

        reader = PdfReader(io.BytesIO(pdf_bytes))
        if reader.is_encrypted:
            if reader.decrypt(password) == 0:
                raise ValueError("Wrong PDF password.")
            writer = PdfWriter()
            for page in reader.pages:
                writer.add_page(page)
            buf = io.BytesIO()
            writer.write(buf)
            buf.seek(0)
            pdf_bytes = buf.read()

    # ── Open PDF ──────────────────────────────────
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    total_pages = len(doc)

    if total_pages == 0:
        raise ValueError("PDF has no pages.")

    # ── Extract PDF metadata ───────────────────────
    meta = doc.metadata or {}
    if not author:
        author = meta.get("author", "")
    if not description:
        description = meta.get("subject", "")
    if not publisher:
        publisher = meta.get("creator", "")

    # ── Create EPUB ───────────────────────────────
    book = epub.EpubBook()
    book.set_title(title)
    book.set_language(language)
    book.set_identifier(f"pdf-{hashlib.md5(pdf_bytes[:1024]).hexdigest()[:12]}")
    if author:
        book.add_author(author)
    if description:
        book.add_metadata("DC", "description", description)
    if publisher:
        book.add_metadata("DC", "publisher", publisher)

    # ── Add CSS ───────────────────────────────────
    css_content = b"""
body {
    font-family   : Georgia, serif;
    font-size     : 1em;
    line-height   : 1.8;
    margin        : 1em;
    color         : #222;
    text-align    : justify;
}
h1 {
    font-family   : Arial, sans-serif;
    font-size     : 1.4em;
    color         : #1F4E79;
    border-bottom : 2px solid #2E75B6;
    padding-bottom: 0.3em;
    margin        : 0.5em 0 1em;
}
h2 {
    font-family : Arial, sans-serif;
    font-size   : 1.1em;
    color       : #2E75B6;
    margin      : 0.8em 0 0.4em;
}
p {
    margin : 0.5em 0;
}
img {
    max-width  : 100%;
    height     : auto;
    display    : block;
    margin     : 1em auto;
    border     : 1px solid #ddd;
    box-shadow : 0 2px 4px rgba(0,0,0,0.1);
}
.page-image {
    width   : 100%;
    margin  : 0;
    padding : 0;
    border  : none;
    box-shadow: none;
}
.page-num {
    font-size   : 0.8em;
    color       : #999;
    text-align  : center;
    margin      : 0.5em 0;
}
.text-content {
    margin        : 1em 0;
    padding       : 0.5em;
    background    : #fafafa;
    border-left   : 3px solid #2E75B6;
}
.no-text {
    color      : #aaa;
    font-style : italic;
    text-align : center;
    font-size  : 0.9em;
}
"""

    css_item = epub.EpubItem(
        uid="style_main",
        file_name="styles/main.css",
        media_type="text/css",
        content=css_content,
    )
    book.add_item(css_item)

    # ── Process cover from first page ─────────────
    cover_page = doc[0]
    cover_pixmap = cover_page.get_pixmap(matrix=fitz.Matrix(1.0, 1.0), alpha=False)
    cover_img = Image.frombytes(
        "RGB",
        [cover_pixmap.width, cover_pixmap.height],
        cover_pixmap.samples,
    )
    cover_buf = io.BytesIO()
    cover_img.save(cover_buf, format="JPEG", quality=85)
    cover_buf.seek(0)
    cover_bytes = cover_buf.read()
    book.set_cover("images/cover.jpg", cover_bytes)

    # ── Process each page ──────────────────────────
    chapters = []
    toc = []
    spine = ["nav"]

    matrix = fitz.Matrix(image_dpi / 72, image_dpi / 72)

    for page_num in range(total_pages):
        page = doc[page_num]
        page_label = f"Page {page_num + 1}"

        # ── Render page as JPEG ─────────────────────
        pixmap = page.get_pixmap(matrix=matrix, alpha=False)
        img = Image.frombytes(
            "RGB",
            [pixmap.width, pixmap.height],
            pixmap.samples,
        )
        img_buf = io.BytesIO()
        img.save(
            img_buf,
            format="JPEG",
            quality=image_quality,
            optimize=True,
        )
        img_buf.seek(0)
        img_bytes = img_buf.read()

        # Add image to EPUB
        img_item = epub.EpubItem(
            uid=f"img_page_{page_num}",
            file_name=f"images/page_{page_num:04d}.jpg",
            media_type="image/jpeg",
            content=img_bytes,
        )
        book.add_item(img_item)

        # ── Extract text ────────────────────────────
        text_html = ""
        if extract_text:
            blocks = page.get_text("blocks")
            text_parts = []

            for block in blocks:
                if block[6] == 0:  # text block type
                    block_text = block[4].strip()
                    if not block_text:
                        continue

                    # Clean text
                    block_text = (
                        block_text.replace("&", "&amp;")
                        .replace("<", "&lt;")
                        .replace(">", "&gt;")
                        .replace("\n", " ")
                        .strip()
                    )

                    if block_text:
                        text_parts.append(f"<p>{block_text}</p>")

            if text_parts:
                text_html = (
                    '<div class="text-content">\n' + "\n".join(text_parts) + "\n</div>"
                )
            else:
                text_html = (
                    '<p class="no-text">' "[No extractable text on this page]" "</p>"
                )

        # ── Build chapter XHTML ─────────────────────
        xhtml = (
            '<?xml version="1.0" encoding="utf-8"?>\n'
            "<!DOCTYPE html>\n"
            '<html xmlns="http://www.w3.org/1999/xhtml">\n'
            "<head>\n"
            f"  <title>{_escape_xml(page_label)}</title>\n"
            '  <link rel="stylesheet" type="text/css" '
            'href="../styles/main.css"/>\n'
            "</head>\n"
            "<body>\n"
            f"  <h1>{_escape_xml(page_label)}</h1>\n"
            f'  <p class="page-num">— {page_num + 1} / {total_pages} —</p>\n'
            f'  <img class="page-image" '
            f'src="../images/page_{page_num:04d}.jpg" '
            f'alt="{_escape_xml(page_label)}"/>\n'
            f"{text_html}\n"
            "</body>\n"
            "</html>"
        )

        chapter = epub.EpubHtml(
            title=page_label,
            file_name=f"chapters/page_{page_num:04d}.xhtml",
            lang=language,
            content=xhtml.encode("utf-8"),
        )
        chapter.add_item(css_item)
        book.add_item(chapter)
        chapters.append(chapter)
        spine.append(chapter)
        toc.append(
            epub.Link(
                f"chapters/page_{page_num:04d}.xhtml",
                page_label,
                f"page_{page_num}",
            )
        )

    doc.close()

    # ── Set spine + TOC ───────────────────────────
    book.spine = spine
    book.toc = toc
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    # ── Write EPUB ────────────────────────────────
    buf = io.BytesIO()
    epub.write_epub(buf, book, {})
    buf.seek(0)
    epub_bytes = buf.read()

    return {
        "bytes": epub_bytes,
        "title": title,
        "author": author,
        "language": language,
        "pages": total_pages,
        "chapters": len(chapters),
        "size_kb": round(len(epub_bytes) / 1024, 2),
        "size_mb": round(len(epub_bytes) / (1024 * 1024), 2),
        "toc": [t.title for t in toc],
        "description": description,
        "publisher": publisher,
    }


def ipynb_to_pdf(
    source,
    filename: str = "notebook.ipynb",
    title: str = "",
    author: str = "",
    include_input: bool = True,
    include_output: bool = True,
    include_markdown: bool = True,
    theme: str = "light",
    page_size: str = "A4",
) -> dict:
    """
    Convert Jupyter Notebook (.ipynb) → PDF.

    Strategy:
        1. Parse .ipynb with nbformat
        2. Walk cells manually — code, markdown, raw, outputs
        3. Render to PDF using reportlab (pure Python, zero C deps)
        4. Supports light/dark theme, A4/Letter/A3 page size

    Args:
        source           : file object OR raw bytes
        filename         : original filename
        title            : document title        (default: filename)
        author           : author name           (default: notebook metadata)
        include_input    : render code inputs     (default: True)
        include_output   : render cell outputs    (default: True)
        include_markdown : render markdown cells  (default: True)
        theme            : 'light' or 'dark'      (default: 'light')
        page_size        : 'A4', 'Letter', 'A3'  (default: 'A4')

    Returns:
        {
            'bytes'         : bytes,
            'title'         : str,
            'author'        : str,
            'language'      : str,
            'total_cells'   : int,
            'code_cells'    : int,
            'markdown_cells': int,
            'size_kb'       : float,
            'size_mb'       : float,
        }
    """
    import io
    import re
    import nbformat
    from reportlab.lib.pagesizes import A4, letter, A3
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        HRFlowable,
        KeepTogether,
        ListFlowable,
        ListItem,
    )
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

    # ── Read source ───────────────────────────────
    if hasattr(source, "read"):
        raw = source.read()
        if not title:
            title = getattr(source, "name", filename)
    elif isinstance(source, (bytes, str)):
        raw = source
    else:
        raise ValueError("source must be a file object, bytes, or str.")

    if not raw:
        raise ValueError("Empty file.")
    if isinstance(raw, bytes):
        raw = raw.decode("utf-8", errors="replace")
    if not title:
        title = filename.replace(".ipynb", "").replace(".IPYNB", "")

    # ── Validate options ──────────────────────────
    if theme not in ("light", "dark"):
        raise ValueError('theme must be "light" or "dark".')
    if page_size not in ("A4", "Letter", "A3"):
        raise ValueError('page_size must be "A4", "Letter", or "A3".')

    # ── Parse notebook ────────────────────────────
    try:
        nb = nbformat.reads(raw, as_version=4)
    except Exception as e:
        raise ValueError(f"Invalid .ipynb file: {e}")

    # ── Extract metadata ──────────────────────────
    nb_meta = nb.get("metadata", {})
    language = nb_meta.get("kernelspec", {}).get("language", "python").capitalize()

    if not author:
        nb_authors = nb_meta.get("authors", [])
        if isinstance(nb_authors, list) and nb_authors:
            author = nb_authors[0].get("name", "")

    total_cells = len(nb.cells)
    code_cells = sum(1 for c in nb.cells if c.cell_type == "code")
    md_cells = sum(1 for c in nb.cells if c.cell_type == "markdown")

    # ── Colour palette ────────────────────────────
    if theme == "dark":
        FG = colors.HexColor("#cdd6f4")
        HEAD_COL = colors.HexColor("#89b4fa")
        CODE_BG = colors.HexColor("#181825")
        PRMPT_C = colors.HexColor("#42a5f5")
    else:
        FG = colors.HexColor("#212121")
        HEAD_COL = colors.HexColor("#1565C0")
        CODE_BG = colors.HexColor("#f5f5f5")
        PRMPT_C = colors.HexColor("#42a5f5")

    BLUE = colors.HexColor("#1565C0")
    GREY_TXT = colors.HexColor("#757575")
    RED_BG = colors.HexColor("#fff3f3")
    RED_TXT = colors.HexColor("#c62828")

    # ── Paragraph styles ──────────────────────────
    base = ParagraphStyle(
        "nb_base",
        fontName="Helvetica",
        fontSize=10,
        leading=15,
        textColor=FG,
        spaceAfter=4,
    )
    st = {
        "h1": ParagraphStyle(
            "nb_h1",
            parent=base,
            fontName="Helvetica-Bold",
            fontSize=18,
            textColor=HEAD_COL,
            spaceAfter=6,
            spaceBefore=12,
        ),
        "h2": ParagraphStyle(
            "nb_h2",
            parent=base,
            fontName="Helvetica-Bold",
            fontSize=14,
            textColor=HEAD_COL,
            spaceAfter=4,
            spaceBefore=10,
        ),
        "h3": ParagraphStyle(
            "nb_h3",
            parent=base,
            fontName="Helvetica-Bold",
            fontSize=12,
            textColor=HEAD_COL,
            spaceAfter=4,
            spaceBefore=8,
        ),
        "h4": ParagraphStyle(
            "nb_h4",
            parent=base,
            fontName="Helvetica-Bold",
            fontSize=11,
            textColor=HEAD_COL,
            spaceAfter=3,
            spaceBefore=6,
        ),
        "body": ParagraphStyle("nb_body", parent=base, alignment=TA_JUSTIFY),
        "li": ParagraphStyle("nb_li", parent=base, leftIndent=16, spaceAfter=2),
        "code": ParagraphStyle(
            "nb_code",
            fontName="Courier",
            fontSize=8,
            leading=11,
            textColor=FG,
            backColor=CODE_BG,
            leftIndent=8,
            rightIndent=8,
            spaceAfter=2,
        ),
        "prompt": ParagraphStyle(
            "nb_prompt",
            fontName="Courier-Bold",
            fontSize=8,
            leading=11,
            textColor=PRMPT_C,
            spaceAfter=1,
        ),
        "stderr": ParagraphStyle(
            "nb_stderr",
            fontName="Courier",
            fontSize=8,
            leading=11,
            textColor=RED_TXT,
            backColor=RED_BG,
            leftIndent=8,
            rightIndent=8,
            spaceAfter=2,
        ),
        "cover_t": ParagraphStyle(
            "nb_cover_t",
            fontName="Helvetica-Bold",
            fontSize=24,
            textColor=HEAD_COL,
            alignment=TA_CENTER,
            spaceAfter=6,
        ),
        "cover_m": ParagraphStyle(
            "nb_cover_m",
            fontName="Helvetica",
            fontSize=10,
            textColor=GREY_TXT,
            alignment=TA_CENTER,
            spaceAfter=3,
        ),
    }

    # ── Helpers ───────────────────────────────────
    def _esc(text: str) -> str:
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def _inline(text: str) -> str:
        """Basic inline markdown → ReportLab XML."""
        text = _esc(text)
        text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
        text = re.sub(r"__(.+?)__", r"<b>\1</b>", text)
        text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
        text = re.sub(r"_(.+?)_", r"<i>\1</i>", text)
        text = re.sub(r"`(.+?)`", r'<font face="Courier">\1</font>', text)
        return text

    def _src(cell) -> str:
        s = cell.get("source", "")
        return "".join(s) if isinstance(s, list) else s

    def _render_markdown(text: str) -> list:
        """Markdown block → list of ReportLab flowables."""
        elems = []
        lines = text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]
            # Headings
            m = re.match(r"^(#{1,4})\s+(.*)", line)
            if m:
                level = min(len(m.group(1)), 4)
                elems.append(Paragraph(_inline(m.group(2)), st[f"h{level}"]))
                i += 1
                continue
            # Unordered list
            if re.match(r"^[\-\*]\s+", line):
                items = []
                while i < len(lines) and re.match(r"^[\-\*]\s+(.*)", lines[i]):
                    lm = re.match(r"^[\-\*]\s+(.*)", lines[i])
                    items.append(
                        ListItem(
                            Paragraph(_inline(lm.group(1)), st["li"]),
                            bulletColor=BLUE,
                            leftIndent=16,
                        )
                    )
                    i += 1
                elems.append(ListFlowable(items, bulletType="bullet", start="•"))
                continue
            # Ordered list
            if re.match(r"^\d+\.\s+", line):
                items = []
                idx = 1
                while i < len(lines) and re.match(r"^\d+\.\s+(.*)", lines[i]):
                    lm = re.match(r"^\d+\.\s+(.*)", lines[i])
                    items.append(
                        ListItem(
                            Paragraph(_inline(lm.group(1)), st["li"]),
                            value=idx,
                        )
                    )
                    i += 1
                    idx += 1
                elems.append(ListFlowable(items, bulletType="1"))
                continue
            # Blank line
            if not line.strip():
                elems.append(Spacer(1, 4))
                i += 1
                continue
            # Normal paragraph
            elems.append(Paragraph(_inline(line), st["body"]))
            i += 1
        return elems

    def _render_output(output: dict, num) -> list:
        """Cell output → list of ReportLab flowables."""
        elems = []
        otype = output.get("output_type", "")
        label = f"Out [{num}]:"

        if otype == "stream":
            text = "".join(output.get("text", []))
            sty = st["stderr"] if output.get("name") == "stderr" else st["code"]
            elems.append(Paragraph(_esc(label), st["prompt"]))
            for line in text.splitlines():
                elems.append(Paragraph(_esc(line) or " ", sty))

        elif otype in ("display_data", "execute_result"):
            data = output.get("data", {})
            if "text/plain" in data:
                txt = data["text/plain"]
                txt = "".join(txt) if isinstance(txt, list) else txt
                elems.append(Paragraph(_esc(label), st["prompt"]))
                for line in txt.splitlines():
                    elems.append(Paragraph(_esc(line) or " ", st["code"]))

        elif otype == "error":
            elems.append(Paragraph(_esc(label), st["prompt"]))
            elems.append(
                Paragraph(
                    f"{_esc(output.get('ename',''))}: {_esc(output.get('evalue',''))}",
                    st["stderr"],
                )
            )

        return elems

    # ── Build PDF story ───────────────────────────
    PAGE_MAP = {"A4": A4, "Letter": letter, "A3": A3}

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=PAGE_MAP[page_size],
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    story = []

    # ── Cover block ───────────────────────────────
    story.append(Spacer(1, 1 * cm))
    story.append(Paragraph(_esc(title), st["cover_t"]))
    story.append(HRFlowable(width="100%", thickness=2, color=BLUE, spaceAfter=6))
    if author:
        story.append(Paragraph(f"<b>{_esc(author)}</b>", st["cover_m"]))
    story.append(
        Paragraph(
            f"{language} kernel &nbsp;•&nbsp; "
            f"{total_cells} cells ({code_cells} code, {md_cells} markdown)",
            st["cover_m"],
        )
    )
    story.append(Spacer(1, 0.6 * cm))
    story.append(
        HRFlowable(
            width="100%",
            thickness=0.5,
            color=colors.HexColor("#e0e0e0"),
            spaceAfter=8,
        )
    )
    story.append(Spacer(1, 0.2 * cm))

    # ── Cells ─────────────────────────────────────
    exec_count = 0
    for cell in nb.cells:
        ctype = cell.cell_type
        src = _src(cell)

        if ctype == "code" and (include_input or include_output):
            exec_count += 1
            num = cell.get("execution_count") or exec_count
            block = []

            if include_input and src.strip():
                block.append(Paragraph(f"In [{num}]:", st["prompt"]))
                for line in src.splitlines():
                    block.append(Paragraph(_esc(line) or " ", st["code"]))
                block.append(Spacer(1, 3))

            if include_output:
                for out in cell.get("outputs", []):
                    block.extend(_render_output(out, num))

            if block:
                story.append(KeepTogether(block))
                story.append(Spacer(1, 6))

        elif ctype == "markdown" and include_markdown and src.strip():
            story.extend(_render_markdown(src))
            story.append(Spacer(1, 4))

        elif ctype == "raw" and src.strip():
            for line in src.splitlines():
                story.append(Paragraph(_esc(line) or " ", st["code"]))
            story.append(Spacer(1, 4))

    # ── Render ────────────────────────────────────
    doc.build(story)
    buf.seek(0)
    pdf_bytes = buf.read()

    if not pdf_bytes or not pdf_bytes.startswith(b"%PDF"):
        raise ValueError("PDF rendering failed.")

    return {
        "bytes": pdf_bytes,
        "title": title,
        "author": author,
        "language": language,
        "total_cells": total_cells,
        "code_cells": code_cells,
        "markdown_cells": md_cells,
        "size_kb": round(len(pdf_bytes) / 1024, 2),
        "size_mb": round(len(pdf_bytes) / (1024 * 1024), 2),
    }

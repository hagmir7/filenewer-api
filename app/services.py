"""
PDF → Word conversion service.

Strategy:
  1. pdfplumber  – extract text with layout + tables
  2. pypdf       – fallback for metadata / basic text
  3. python-docx – build the .docx output
"""

import io
import os
import re
import logging
from pathlib import Path
import pandas as pd
import json

import pdfplumber
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────
# PDF TO WORD
# ─────────────────────────────────────────────


def _looks_like_heading(text: str) -> bool:
    """Heuristic: short, no period at end, possible ALL CAPS."""
    text = text.strip()
    if not text:
        return False
    if len(text) > 120:
        return False
    if text.endswith("."):
        return False
    if text.isupper() and len(text) > 3:
        return True
    if re.match(r"^(\d+[\.\)]\s+|\d+\.\d+\s+)", text):  # "1. Title" / "1.2 Section"
        return True
    return False


def _add_table_to_doc(doc: Document, table_data: list[list]) -> None:
    """Add a pdfplumber table into the Word document."""
    if not table_data:
        return

    rows = len(table_data)
    cols = max(len(row) for row in table_data)
    if rows == 0 or cols == 0:
        return

    word_table = doc.add_table(rows=rows, cols=cols)
    word_table.style = "Table Grid"

    for r_idx, row in enumerate(table_data):
        for c_idx, cell_text in enumerate(row):
            if c_idx >= cols:
                break
            cell = word_table.cell(r_idx, c_idx)
            cell.text = str(cell_text) if cell_text is not None else ""
            # Bold first row (header)
            if r_idx == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def convert_pdf_to_docx(pdf_bytes: bytes) -> bytes:
    """
    Convert PDF bytes → DOCX bytes.

    Returns the raw bytes of the generated .docx file.
    """
    doc = Document()

    # ── Default style tweaks ──────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        total_pages = len(pdf.pages)
        logger.info("Converting PDF: %d page(s)", total_pages)

        for page_num, page in enumerate(pdf.pages, start=1):
            # ── Extract tables first ─────────────
            tables = page.extract_tables() or []
            table_bboxes = []

            for table_data in tables:
                if table_data:
                    _add_table_to_doc(doc, table_data)
                    doc.add_paragraph()  # breathing room after table

            # ── Extract text (excluding table regions) ──
            # Crop away table areas to avoid duplicating content
            cropped_page = page
            for table in page.find_tables():
                try:
                    cropped_page = cropped_page.outside_bbox(table.bbox)
                except Exception:
                    pass  # some versions don't support outside_bbox; fall through

            raw_text = cropped_page.extract_text(x_tolerance=3, y_tolerance=3) or ""

            if not raw_text.strip() and not tables:
                # Blank page – add a soft separator
                doc.add_paragraph()
                continue

            for line in raw_text.splitlines():
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                    continue

                if _looks_like_heading(line):
                    heading_level = 1 if line.isupper() else 2
                    doc.add_heading(line, level=heading_level)
                else:
                    para = doc.add_paragraph(line)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # ── Page break between pages (except last) ──
            if page_num < total_pages:
                doc.add_page_break()

    # ── Serialize to bytes ────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# ─────────────────────────────────────────────
# CSV TO SQL
# ─────────────────────────────────────────────

def sanitize_name(name: str) -> str:
    name = name.strip().lower()
    name = re.sub(r"[^a-z0-9_]", "_", name)
    if name[0].isdigit():
        name = "col_" + name
    return name


def infer_sql_type(dtype) -> str:
    dtype = str(dtype)
    if "int" in dtype:
        return "INTEGER"
    if "float" in dtype:
        return "REAL"
    if "bool" in dtype:
        return "BOOLEAN"
    if "datetime" in dtype:
        return "TIMESTAMP"
    return "TEXT"


def csv_to_sql(source, table_name: str, dialect: str = 'sqlite', separator: str = None) -> str:
    """
    Convert CSV (file or raw text string) → SQL statements.

    Args:
        source     : uploaded file object OR raw CSV string
        table_name : name for the SQL table
        dialect    : 'sqlite' | 'postgresql' | 'mysql'

    Returns:
        SQL string (CREATE TABLE + INSERT statements)
    """
    # ── Read CSV ─────────────────────────────────
    if isinstance(source, str):
        df = pd.read_csv(io.StringIO(source), sep=separator, engine='python')
    else:
        # ── File: read bytes first, then decode ──
        raw = source.read().decode('utf-8')          # ← read raw bytes
        df  = pd.read_csv(io.StringIO(raw), sep=separator, engine='python') 

    df.columns = [sanitize_name(c) for c in df.columns]
    table_name = sanitize_name(table_name)

    # ── Dialect config ────────────────────────────
    if dialect == "postgresql":
        pk = "id SERIAL PRIMARY KEY"
        str_quote = "E'{}'"
    elif dialect == "mysql":
        pk = "id INT AUTO_INCREMENT PRIMARY KEY"
        str_quote = "'{}'"
    else:  # sqlite (default)
        pk = "id INTEGER PRIMARY KEY AUTOINCREMENT"
        str_quote = "'{}'"

    # ── CREATE TABLE ──────────────────────────────
    col_defs = ",\n  ".join(
        f'"{col}" {infer_sql_type(dtype)}' for col, dtype in df.dtypes.items()
    )
    sql_lines = [
        f"-- Generated SQL ({dialect}) for table: {table_name}",
        f"-- Rows: {len(df)}  |  Columns: {len(df.columns)}\n",
        f'CREATE TABLE IF NOT EXISTS "{table_name}" (',
        f"  {pk},",
        f"  {col_defs}",
        ");\n",
    ]

    # ── INSERT statements ─────────────────────────
    col_str = ", ".join(f'"{c}"' for c in df.columns)

    for _, row in df.iterrows():
        values = []
        for val in row:
            if pd.isna(val):
                values.append("NULL")
            elif isinstance(val, bool):
                values.append("TRUE" if val else "FALSE")
            elif isinstance(val, (int, float)):
                values.append(str(val))
            else:
                escaped = str(val).replace("'", "''")
                values.append(f"'{escaped}'")

        val_str = ", ".join(values)
        sql_lines.append(f'INSERT INTO "{table_name}" ({col_str}) VALUES ({val_str});')

    return "\n".join(sql_lines)


# ─────────────────────────────────────────────
# CSV TO JSON
# ─────────────────────────────────────────────
def csv_to_json(source, separator: str = None, orient: str = "records") -> list | dict:
    """
    Convert CSV (file or raw text string) → JSON.

    Args:
        source    : uploaded file object OR raw CSV string
        separator : column separator (None = auto-detect)
        orient    : 'records'  → [ {col: val}, ... ]         (default)
                    'columns'  → { col: {index: val}, ... }
                    'values'   → [ [val, val], ... ]
                    'index'    → { index: {col: val}, ... }

    Returns:
        parsed Python object (list or dict) ready for JSON response
    """
    # ── Read CSV ──────────────────────────────────
    if isinstance(source, str):
        df = pd.read_csv(io.StringIO(source), sep=separator, engine="python")
    else:
        raw = source.read().decode("utf-8")
        df = pd.read_csv(io.StringIO(raw), sep=separator, engine="python")

    # ── Clean column names ─────────────────────────
    df.columns = [sanitize_name(c) for c in df.columns]

    # ── Replace NaN with None (JSON null) ─────────
    df = df.where(pd.notnull(df), None)

    # ── Convert to JSON-serializable object ────────
    return df.to_dict(orient=orient)


def json_to_csv(source, separator: str = ",") -> str:
    """
    Convert JSON (file, raw text, or dict/list) → CSV string.

    Args:
        source    : uploaded file object | raw JSON string | list | dict
        separator : column separator (default: ',')

    Returns:
        CSV string
    """
    # ── Parse input ───────────────────────────────
    if isinstance(source, (list, dict)):
        data = source  # already parsed
    elif isinstance(source, str):
        data = json.loads(source)  # raw JSON string
    else:
        raw = source.read().decode("utf-8")  # file object
        data = json.loads(raw)

    # ── Handle nested/wrapped JSON ─────────────────
    # e.g. { "data": [ {...}, {...} ] }
    if isinstance(data, dict):
        # find the first key whose value is a list
        for key, val in data.items():
            if isinstance(val, list):
                data = val
                break
        else:
            data = [data]  # single object → wrap in list

    if not isinstance(data, list):
        raise ValueError("JSON must be an array of objects or a wrapped array.")

    # ── Convert to DataFrame → CSV ─────────────────
    df = pd.json_normalize(data)  # flattens nested objects
    df.columns = [sanitize_name(c) for c in df.columns]

    return df.to_csv(index=False, sep=separator)


def excel_to_csv(source, sheet_name=None, separator: str = ",") -> dict:
    """
    Convert Excel file (.xlsx / .xls) → CSV string(s).

    Args:
        source     : uploaded file object or file path
        sheet_name : specific sheet name or index (None = all sheets)
        separator  : column separator (default: ',')

    Returns:
        {
            'sheets': {
                'Sheet1': 'csv string...',
                'Sheet2': 'csv string...',
            },
            'total_sheets': 2,
        }
    """
    # ── Read Excel ────────────────────────────────
    if hasattr(source, "read"):
        raw = source.read()
        xls = pd.ExcelFile(io.BytesIO(raw))  # ← BytesIO for binary
    else:
        xls = pd.ExcelFile(source)

    available_sheets = xls.sheet_names

    # ── Decide which sheets to convert ────────────
    if sheet_name is not None:
        # single sheet by name or index
        if isinstance(sheet_name, int):
            if sheet_name >= len(available_sheets):
                raise ValueError(
                    f"Sheet index {sheet_name} out of range. "
                    f"Available: 0-{len(available_sheets)-1}"
                )
            sheet_name = available_sheets[sheet_name]

        if sheet_name not in available_sheets:
            raise ValueError(
                f"Sheet '{sheet_name}' not found. "
                f"Available sheets: {available_sheets}"
            )
        sheets_to_convert = [sheet_name]
    else:
        sheets_to_convert = available_sheets

    # ── Convert each sheet → CSV ──────────────────
    result = {}
    for name in sheets_to_convert:
        df = pd.read_excel(xls, sheet_name=name, engine=None)

        # Clean column names
        df.columns = [sanitize_name(str(c)) for c in df.columns]

        # Replace NaN with empty string for clean CSV
        df = df.fillna("")

        result[name] = df.to_csv(index=False, sep=separator)

    return {
        "sheets": result,
        "total_sheets": len(result),
        "sheet_names": list(result.keys()),
    }


def json_to_excel(source, sheet_name: str = "Sheet1") -> bytes:
    """
    Convert JSON (file, raw text, list, or dict) → Excel bytes.

    Args:
        source     : uploaded file object | raw JSON string | list | dict
        sheet_name : name of the Excel sheet (default: Sheet1)

    Returns:
        Raw bytes of the generated .xlsx file
    """
    # ── Parse input ───────────────────────────────
    if isinstance(source, (list, dict)):
        data = source  # already parsed
    elif isinstance(source, str):
        data = json.loads(source)  # raw JSON string
    else:
        raw = source.read().decode("utf-8")  # file object
        data = json.loads(raw)

    # ── Handle all JSON shapes ─────────────────────
    if isinstance(data, dict):
        # { "data": [{...}] }  → unwrap
        for key, val in data.items():
            if isinstance(val, list):
                data = val
                break
        else:
            data = [data]  # single object → wrap in list

    if not isinstance(data, list):
        raise ValueError("JSON must be an array of objects or a wrapped array.")

    # ── Flatten nested objects ─────────────────────
    df = pd.json_normalize(data)
    df.columns = [sanitize_name(c) for c in df.columns]

    # ── Write to Excel in memory ───────────────────
    buffer = io.BytesIO()

    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name=sheet_name)

        # ── Auto-fit column widths ─────────────────
        worksheet = writer.sheets[sheet_name]
        for col_cells in worksheet.columns:
            max_len = max(
                len(str(cell.value)) if cell.value is not None else 0
                for cell in col_cells
            )
            col_letter = col_cells[0].column_letter
            worksheet.column_dimensions[col_letter].width = min(max_len + 4, 50)

        # ── Style header row ───────────────────────
        from openpyxl.styles import Font, PatternFill, Alignment

        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(fill_type="solid", fgColor="2E75B6")
        header_align = Alignment(horizontal="center", vertical="center")

        for cell in worksheet[1]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_align

    buffer.seek(0)
    return buffer.read()


def json_to_excel_multisheets(sources: dict) -> bytes:
    """
    Convert multiple JSON arrays → multi-sheet Excel file.

    Args:
        sources : { 'SheetName': list_or_dict, ... }

    Returns:
        Raw bytes of the generated .xlsx file
    """
    from openpyxl.styles import Font, PatternFill, Alignment

    buffer = io.BytesIO()

    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        for sheet, data in sources.items():

            # ── Handle shapes ──────────────────────
            if isinstance(data, dict):
                for key, val in data.items():
                    if isinstance(val, list):
                        data = val
                        break
                else:
                    data = [data]

            df = pd.json_normalize(data)
            df.columns = [sanitize_name(c) for c in df.columns]
            df.to_excel(
                writer, index=False, sheet_name=sheet[:31]
            )  # Excel max 31 chars

            # ── Auto-fit + style header ────────────
            worksheet = writer.sheets[sheet[:31]]

            for col_cells in worksheet.columns:
                max_len = max(
                    len(str(cell.value)) if cell.value is not None else 0
                    for cell in col_cells
                )
                col_letter = col_cells[0].column_letter
                worksheet.column_dimensions[col_letter].width = min(max_len + 4, 50)

            for cell in worksheet[1]:
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill(fill_type="solid", fgColor="2E75B6")
                cell.alignment = Alignment(horizontal="center", vertical="center")

    buffer.seek(0)
    return buffer.read()


def csv_to_excel(source, sheet_name: str = "Sheet1", separator: str = None) -> bytes:
    """
    Convert CSV (file or raw text string) → Excel bytes.

    Args:
        source     : uploaded file object OR raw CSV string
        sheet_name : name of the Excel sheet (default: Sheet1)
        separator  : column separator (None = auto-detect)

    Returns:
        Raw bytes of the generated .xlsx file
    """
    # ── Read CSV ──────────────────────────────────
    if isinstance(source, str):
        df = pd.read_csv(io.StringIO(source), sep=separator, engine="python")
    else:
        raw = source.read().decode("utf-8")
        df = pd.read_csv(io.StringIO(raw), sep=separator, engine="python")

    # ── Clean column names ─────────────────────────
    df.columns = [sanitize_name(str(c)) for c in df.columns]

    # ── Replace NaN with empty string ─────────────
    df = df.fillna("")

    # ── Write to Excel in memory ───────────────────
    buffer = io.BytesIO()

    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name=sheet_name)

        worksheet = writer.sheets[sheet_name]

        # ── Auto-fit column widths ─────────────────
        for col_cells in worksheet.columns:
            max_len = max(
                len(str(cell.value)) if cell.value is not None else 0
                for cell in col_cells
            )
            col_letter = col_cells[0].column_letter
            worksheet.column_dimensions[col_letter].width = min(max_len + 4, 50)

        # ── Style header row ───────────────────────
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

        header_font = Font(bold=True, color="FFFFFF", size=11)
        header_fill = PatternFill(fill_type="solid", fgColor="2E75B6")
        header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
        thin_border = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin"),
        )

        for cell in worksheet[1]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_align
            cell.border = thin_border

        # ── Style data rows (alternating colors) ───
        from openpyxl.styles import PatternFill as PF

        even_fill = PF(fill_type="solid", fgColor="DCE6F1")  # light blue
        odd_fill = PF(fill_type="solid", fgColor="FFFFFF")  # white

        for row_idx, row in enumerate(worksheet.iter_rows(min_row=2), start=2):
            fill = even_fill if row_idx % 2 == 0 else odd_fill
            for cell in row:
                cell.fill = fill
                cell.border = thin_border
                cell.alignment = Alignment(vertical="center")

        # ── Freeze header row ──────────────────────
        worksheet.freeze_panes = "A2"

        # ── Set row height for header ──────────────
        worksheet.row_dimensions[1].height = 30

    buffer.seek(0)
    return buffer.read()


def csv_to_excel_multisheets(sources: dict, separator: str = None) -> bytes:
    """
    Convert multiple CSV strings → multi-sheet Excel file.

    Args:
        sources   : { 'SheetName': 'csv string...', ... }
        separator : column separator (None = auto-detect)

    Returns:
        Raw bytes of the generated .xlsx file
    """
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    buffer = io.BytesIO()

    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        for sheet, csv_data in sources.items():

            # ── Read CSV ───────────────────────────
            if isinstance(csv_data, str):
                df = pd.read_csv(io.StringIO(csv_data), sep=separator, engine="python")
            else:
                raw = csv_data.read().decode("utf-8")
                df = pd.read_csv(io.StringIO(raw), sep=separator, engine="python")

            df.columns = [sanitize_name(str(c)) for c in df.columns]
            df = df.fillna("")

            sheet_label = sheet[:31]  # Excel max 31 chars
            df.to_excel(writer, index=False, sheet_name=sheet_label)

            worksheet = writer.sheets[sheet_label]

            # ── Auto-fit columns ───────────────────
            for col_cells in worksheet.columns:
                max_len = max(
                    len(str(cell.value)) if cell.value is not None else 0
                    for cell in col_cells
                )
                col_letter = col_cells[0].column_letter
                worksheet.column_dimensions[col_letter].width = min(max_len + 4, 50)

            # ── Style header ───────────────────────
            thin = Border(
                left=Side(style="thin"),
                right=Side(style="thin"),
                top=Side(style="thin"),
                bottom=Side(style="thin"),
            )
            for cell in worksheet[1]:
                cell.font = Font(bold=True, color="FFFFFF", size=11)
                cell.fill = PatternFill(fill_type="solid", fgColor="2E75B6")
                cell.alignment = Alignment(horizontal="center", vertical="center")
                cell.border = thin

            # ── Alternating row colors ─────────────
            for row_idx, row in enumerate(worksheet.iter_rows(min_row=2), start=2):
                fill_color = "DCE6F1" if row_idx % 2 == 0 else "FFFFFF"
                for cell in row:
                    cell.fill = PatternFill(fill_type="solid", fgColor=fill_color)
                    cell.border = thin
                    cell.alignment = Alignment(vertical="center")

            worksheet.freeze_panes = "A2"
            worksheet.row_dimensions[1].height = 30

    buffer.seek(0)
    return buffer.read()


def pdf_to_excel(source, password: str = None) -> bytes:
    """
    Convert PDF (file or bytes) → Excel bytes.

    Strategy:
        1. Extract tables  → each table gets its own sheet
        2. Extract text    → one 'Text' sheet with all plain text
        3. Style headers   → bold white on blue, alternating rows, auto-width

    Args:
        source   : uploaded file object OR raw bytes
        password : PDF password if encrypted (default: None)

    Returns:
        Raw bytes of the generated .xlsx file
    """
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl import Workbook

    # ── Read PDF bytes ────────────────────────────
    if hasattr(source, "read"):
        pdf_bytes = source.read()
    else:
        pdf_bytes = source

    # ── Styles ────────────────────────────────────
    def make_header_style():
        return {
            "font": Font(bold=True, color="FFFFFF", size=11),
            "fill": PatternFill(fill_type="solid", fgColor="2E75B6"),
            "align": Alignment(horizontal="center", vertical="center", wrap_text=True),
            "border": Border(
                left=Side(style="thin"),
                right=Side(style="thin"),
                top=Side(style="thin"),
                bottom=Side(style="thin"),
            ),
        }

    def apply_header(ws):
        style = make_header_style()
        for cell in ws[1]:
            cell.font = style["font"]
            cell.fill = style["fill"]
            cell.alignment = style["align"]
            cell.border = style["border"]
        ws.row_dimensions[1].height = 25

    def apply_rows(ws):
        thin = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin"),
        )
        for row_idx, row in enumerate(ws.iter_rows(min_row=2), start=2):
            color = "DCE6F1" if row_idx % 2 == 0 else "FFFFFF"
            for cell in row:
                cell.fill = PatternFill(fill_type="solid", fgColor=color)
                cell.border = thin
                cell.alignment = Alignment(vertical="center")

    def auto_fit(ws):
        for col_cells in ws.columns:
            max_len = max(
                len(str(cell.value)) if cell.value is not None else 0
                for cell in col_cells
            )
            ws.column_dimensions[col_cells[0].column_letter].width = min(
                max_len + 4, 60
            )

    # ── Parse PDF ─────────────────────────────────
    wb = Workbook()
    wb.remove(wb.active)  # remove default empty sheet

    table_count = 0
    all_text = []

    open_kwargs = {"stream": io.BytesIO(pdf_bytes)}
    pdf_stream = io.BytesIO(pdf_bytes)
    if password:
        open_kwargs["password"] = password

    with pdfplumber.open(pdf_stream, password=password) as pdf:
        total_pages = len(pdf.pages)

        for page_num, page in enumerate(pdf.pages, start=1):

            # ── Extract tables ─────────────────────
            tables = page.extract_tables() or []
            for table in tables:
                if not table:
                    continue

                # clean None cells
                cleaned = [
                    [str(cell).strip() if cell is not None else "" for cell in row]
                    for row in table
                ]

                table_count += 1
                sheet_label = f"Table_{table_count}_P{page_num}"[:31]
                ws = wb.create_sheet(title=sheet_label)

                for row in cleaned:
                    ws.append(row)

                apply_header(ws)
                apply_rows(ws)
                auto_fit(ws)
                ws.freeze_panes = "A2"

            # ── Extract plain text ─────────────────
            # Exclude table regions to avoid duplication
            cropped = page
            for tbl in page.find_tables():
                try:
                    cropped = cropped.outside_bbox(tbl.bbox)
                except Exception:
                    pass

            text = cropped.extract_text(x_tolerance=3, y_tolerance=3) or ""
            if text.strip():
                all_text.append(f"--- Page {page_num} ---")
                all_text.append(text.strip())
                all_text.append("")

    # ── Text sheet ────────────────────────────────
    if all_text:
        ws_text = wb.create_sheet(title="Text")
        ws_text.column_dimensions["A"].width = 100

        # Header
        ws_text.append(["Extracted Text"])
        apply_header(ws_text)

        for line in all_text:
            ws_text.append([line])

        # Light style for text rows
        thin = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin"),
        )
        for row in ws_text.iter_rows(min_row=2):
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical="top")
                cell.border = thin

    # ── Summary sheet (first sheet) ───────────────
    ws_summary = wb.create_sheet(title="Summary", index=0)
    ws_summary.column_dimensions["A"].width = 25
    ws_summary.column_dimensions["B"].width = 40

    summary_data = [
        ["Property", "Value"],
        ["Total Pages", total_pages],
        ["Tables Found", table_count],
        ["Has Text", "Yes" if all_text else "No"],
        ["Sheets", wb.sheetnames.__len__()],
    ]

    for row in summary_data:
        ws_summary.append(row)

    apply_header(ws_summary)
    apply_rows(ws_summary)
    auto_fit(ws_summary)

    # ── Serialize ─────────────────────────────────
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.read()


def pdf_to_jpg(source, dpi: int = 200, quality: int = 85, password: str = None) -> list[dict]:
    """
    Convert PDF → JPG using pymupdf (no poppler / no system deps needed).

    Args:
        source   : uploaded file object OR raw bytes
        dpi      : image resolution (default: 200)
        quality  : JPG quality 1-95 (default: 85)
        password : PDF password if encrypted (default: None)

    Returns:
        list of {
            'page'    : int,
            'bytes'   : bytes,
            'width'   : int,
            'height'  : int,
            'filename': str,
        }
    """
    import fitz          # pymupdf
    from PIL import Image

    # ── Read PDF bytes ────────────────────────────
    if hasattr(source, 'read'):
        pdf_bytes = source.read()
    else:
        pdf_bytes = source

    # ── Open PDF ──────────────────────────────────
    doc = fitz.open(stream=pdf_bytes, filetype='pdf')

    if password:
        if not doc.authenticate(password):
            raise ValueError('Invalid PDF password.')

    # ── Convert each page → JPG ───────────────────
    zoom   = dpi / 72        # pymupdf base DPI is 72
    matrix = fitz.Matrix(zoom, zoom)

    results = []

    for page_num in range(len(doc)):
        page    = doc[page_num]
        pixmap  = page.get_pixmap(matrix=matrix, alpha=False)

        # Pixmap → PIL Image → JPG bytes
        image = Image.frombytes(
            'RGB',
            [pixmap.width, pixmap.height],
            pixmap.samples,
        )

        buffer = io.BytesIO()
        image.save(
            buffer,
            format  ='JPEG',
            quality =quality,
            optimize=True,
        )
        buffer.seek(0)
        jpg_bytes = buffer.read()

        results.append({
            'page'    : page_num + 1,
            'bytes'   : jpg_bytes,
            'width'   : pixmap.width,
            'height'  : pixmap.height,
            'filename': f'page_{page_num + 1}.jpg',
        })

    doc.close()
    return results


# ── Arabic support ─────────────────────────────────
try:
    from bidi.algorithm import get_display
    import arabic_reshaper

    ARABIC_SUPPORT = True
except ImportError:
    ARABIC_SUPPORT = False


def is_arabic(text: str) -> bool:
    """Check if text contains Arabic characters."""
    return bool(re.search(r"[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF]", text))


def process_arabic_text(text: str) -> str:
    """
    Reshape and reorder Arabic text for correct PDF rendering.
    Arabic needs two steps:
        1. Reshape  → connect letters correctly
        2. Bidi     → reverse for RTL display
    """
    if not ARABIC_SUPPORT or not text.strip():
        return text

    try:
        if is_arabic(text):
            reshaped = arabic_reshaper.reshape(text)
            return get_display(reshaped)
        return text
    except Exception:
        return text


def register_arabic_fonts(fonts_dir: str = None) -> dict:
    """
    Register Arabic fonts with reportlab.
    Returns dict of available font names.
    """
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    if fonts_dir is None:
        # Default: fonts/ folder next to manage.py
        fonts_dir = os.path.join(
            os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "fonts"
        )

    registered = {}
    font_files = {
        "Amiri": "Amiri-Regular.ttf",
        "Amiri-Bold": "Amiri-Bold.ttf",
        "Amiri-Italic": "Amiri-Italic.ttf",
        "Amiri-BoldItalic": "Amiri-BoldItalic.ttf",
    }

    for font_name, font_file in font_files.items():
        font_path = os.path.join(fonts_dir, font_file)
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont(font_name, font_path))
                registered[font_name] = font_path
            except Exception:
                pass

    return registered


def word_to_pdf(source, filename: str = "document.docx") -> bytes:
    """
    Convert Word (.docx) → PDF with full Arabic / RTL support.
    """
    from docx import Document as DocxDocument
    from docx.text.paragraph import Paragraph as DocxPara
    from docx.table import Table as DocxTable
    from docx.oxml.ns import qn
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
        PageBreak,
    )

    # ── Read DOCX ─────────────────────────────────
    if hasattr(source, "read"):
        file_bytes = source.read()
    else:
        file_bytes = source

    docx = DocxDocument(io.BytesIO(file_bytes))

    # ── Register Arabic fonts ──────────────────────
    registered_fonts = register_arabic_fonts()
    has_arabic_font = bool(registered_fonts)

    arabic_font = "Amiri" if "Amiri" in registered_fonts else "Helvetica"
    arabic_font_bold = (
        "Amiri-Bold" if "Amiri-Bold" in registered_fonts else "Helvetica-Bold"
    )
    latin_font = "Helvetica"
    latin_font_bold = "Helvetica-Bold"

    # ── Helper: pick font based on content ─────────
    def pick_font(text: str, bold: bool = False) -> str:
        if is_arabic(text):
            return arabic_font_bold if bold else arabic_font
        return latin_font_bold if bold else latin_font

    # ── Helper: process text (Arabic or Latin) ─────
    def prepare_text(text: str) -> str:
        if is_arabic(text):
            return process_arabic_text(text)
        return text

    # ── Helper: escape XML special chars ───────────
    def escape(text: str) -> str:
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    # ── Custom styles ──────────────────────────────
    def make_style(name, font, font_bold, size, color, before, after, align=TA_LEFT):
        return ParagraphStyle(
            name,
            fontName=font_bold,
            fontSize=size,
            textColor=color,
            spaceBefore=before,
            spaceAfter=after,
            leading=size * 1.4,
            alignment=align,
            wordWrap="RTL" if font == arabic_font else "LTR",
        )

    # ── Process paragraph runs → markup ────────────
    def runs_to_markup(para) -> tuple[str, bool]:
        """
        Returns (markup_string, has_arabic).
        Uses correct font per run.
        """
        parts = []
        has_arabic = False

        for run in para.runs:
            raw = run.text
            if not raw:
                continue

            arabic = is_arabic(raw)
            if arabic:
                has_arabic = True
                raw = process_arabic_text(raw)

            text = escape(raw)
            font = pick_font(run.text, bold=run.bold)

            # Wrap in font tag
            text = f'<font name="{font}">{text}</font>'

            if run.bold:
                text = f"<b>{text}</b>"
            if run.italic:
                text = f"<i>{text}</i>"
            if run.underline:
                text = f"<u>{text}</u>"

            parts.append(text)

        return "".join(parts), has_arabic

    # ── Pick paragraph style ───────────────────────
    def pick_style(para, has_arabic: bool) -> ParagraphStyle:
        name = (para.style.name or "Normal").lower()
        align = TA_RIGHT if has_arabic else TA_JUSTIFY
        font = arabic_font if has_arabic else latin_font
        bold = arabic_font_bold if has_arabic else latin_font_bold
        wrap = "RTL" if has_arabic else "LTR"

        base = {
            "fontName": font,
            "leading": 20,
            "wordWrap": wrap,
            "alignment": align,
        }

        if "heading 1" in name:
            return ParagraphStyle(
                "H1",
                **base,
                fontSize=22,
                textColor=colors.HexColor("#2E75B6"),
                fontName=bold,
                spaceBefore=18,
                spaceAfter=14,
            )

        if "heading 2" in name:
            return ParagraphStyle(
                "H2",
                **base,
                fontSize=16,
                textColor=colors.HexColor("#2E75B6"),
                fontName=bold,
                spaceBefore=14,
                spaceAfter=10,
            )

        if "heading 3" in name:
            return ParagraphStyle(
                "H3",
                **base,
                fontSize=13,
                textColor=colors.HexColor("#1F4E79"),
                fontName=bold,
                spaceBefore=10,
                spaceAfter=8,
            )

        if "list bullet" in name:
            return ParagraphStyle(
                "Bullet", **base, fontSize=11, leftIndent=20, spaceAfter=4
            )

        if "list number" in name:
            return ParagraphStyle(
                "Number", **base, fontSize=11, leftIndent=20, spaceAfter=4
            )

        return ParagraphStyle("Body", **base, fontSize=11, spaceAfter=6)

    # ── Process table ──────────────────────────────
    def process_table(tbl) -> Table:
        data = []
        col_cnt = max(len(row.cells) for row in tbl.rows)
        col_w = (A4[0] - 2 * inch) / col_cnt

        for r_idx, row in enumerate(tbl.rows):
            row_data = []
            for cell in row.cells:
                cell_text = "\n".join(p.text for p in cell.paragraphs)
                arabic = is_arabic(cell_text)

                if arabic:
                    cell_text = process_arabic_text(cell_text)

                font = (
                    arabic_font_bold
                    if (r_idx == 0 and arabic)
                    else (
                        arabic_font
                        if arabic
                        else (latin_font_bold if r_idx == 0 else latin_font)
                    )
                )

                align = TA_RIGHT if arabic else TA_LEFT
                wrap = "RTL" if arabic else "LTR"

                style = ParagraphStyle(
                    f"cell_{r_idx}",
                    fontName=font,
                    fontSize=10,
                    textColor=colors.white if r_idx == 0 else colors.black,
                    alignment=align,
                    wordWrap=wrap,
                    leading=14,
                )
                row_data.append(Paragraph(escape(cell_text), style))
            data.append(row_data)

        pdf_tbl = Table(data, colWidths=[col_w] * col_cnt, repeatRows=1)
        pdf_tbl.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2E75B6")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    (
                        "ROWBACKGROUNDS",
                        (0, 1),
                        (-1, -1),
                        [colors.HexColor("#DCE6F1"), colors.white],
                    ),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#AAAAAA")),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("PADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        return pdf_tbl

    # ── Build story ────────────────────────────────
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
        title=filename.replace(".docx", "").replace(".doc", ""),
    )

    story = []

    for element in docx.element.body:
        tag = element.tag.split("}")[-1]

        # ── Paragraph ─────────────────────────────
        if tag == "p":
            para = DocxPara(element, docx)

            # Page break
            if any(
                br.get(qn("w:type")) == "page"
                for br in element.findall(".//" + qn("w:br"))
            ):
                story.append(PageBreak())
                continue

            # Empty
            if not para.text.strip():
                story.append(Spacer(1, 8))
                continue

            markup, has_arabic = runs_to_markup(para)
            style = pick_style(para, has_arabic)

            # Bullet
            name = (para.style.name or "").lower()
            if "list bullet" in name:
                prefix = "• " if not has_arabic else " •"
                try:
                    story.append(Paragraph(prefix + markup, style))
                except Exception:
                    story.append(
                        Paragraph(prefix + escape(prepare_text(para.text)), style)
                    )
                continue

            # Numbered
            if "list number" in name:
                try:
                    story.append(Paragraph("1. " + markup, style))
                except Exception:
                    story.append(
                        Paragraph("1. " + escape(prepare_text(para.text)), style)
                    )
                continue

            try:
                story.append(Paragraph(markup, style))
            except Exception:
                story.append(Paragraph(escape(prepare_text(para.text)), style))

        # ── Table ──────────────────────────────────
        elif tag == "tbl":
            try:
                tbl = DocxTable(element, docx)
                story.append(Spacer(1, 8))
                story.append(process_table(tbl))
                story.append(Spacer(1, 12))
            except Exception as e:
                story.append(
                    Paragraph(
                        f"[Table error: {e}]",
                        ParagraphStyle("err", fontName="Helvetica", fontSize=10),
                    )
                )

        elif tag == "sectPr":
            pass

    doc.build(story)
    buffer.seek(0)
    return buffer.read()


def encrypt_pdf(
    source,
    user_password: str = "",
    owner_password: str = None,
    allow_printing: bool = True,
    allow_copying: bool = True,
    allow_editing: bool = True,
    allow_annotations: bool = True,
) -> bytes:
    """
    Encrypt a PDF with password protection and permissions.

    Args:
        source           : uploaded file object OR raw bytes
        user_password    : password required to open the PDF
        owner_password   : password for full access (default: same as user)
        allow_printing   : allow printing                (default: True)
        allow_copying    : allow copying text            (default: True)
        allow_editing    : allow editing                 (default: True)
        allow_annotations: allow adding annotations      (default: True)

    Returns:
        Raw bytes of the encrypted PDF
    """
    from pypdf import PdfReader, PdfWriter
    from pypdf.generic import NameObject
    import pypdf.constants as pdfconst

    # ── Read PDF bytes ────────────────────────────
    if hasattr(source, "read"):
        pdf_bytes = source.read()
    else:
        pdf_bytes = source

    # ── Validate it is a real PDF ──────────────────
    if not pdf_bytes.startswith(b"%PDF"):
        raise ValueError("Invalid PDF file.")

    # ── Read + Write ──────────────────────────────
    reader = PdfReader(io.BytesIO(pdf_bytes))

    # Check if already encrypted
    if reader.is_encrypted:
        raise ValueError(
            "PDF is already encrypted. " "Please decrypt it first before re-encrypting."
        )

    writer = PdfWriter()

    # ── Copy all pages ─────────────────────────────
    for page in reader.pages:
        writer.add_page(page)

    # ── Copy metadata ──────────────────────────────
    if reader.metadata:
        writer.add_metadata(reader.metadata)

    # ── Set owner password ─────────────────────────
    if owner_password is None:
        owner_password = user_password

    # ── Build permissions ─────────────────────────
    permissions = build_permissions(
        allow_printing=allow_printing,
        allow_copying=allow_copying,
        allow_editing=allow_editing,
        allow_annotations=allow_annotations,
    )

    # ── Encrypt ────────────────────────────────────
    writer.encrypt(
        user_password=user_password,
        owner_password=owner_password,
        use_128bit=True,
        permissions_flag=permissions,
    )

    # ── Serialize ─────────────────────────────────
    buffer = io.BytesIO()
    writer.write(buffer)
    buffer.seek(0)
    return buffer.read()


def build_permissions(
    allow_printing: bool = True,
    allow_copying: bool = True,
    allow_editing: bool = True,
    allow_annotations: bool = True,
) -> int:
    """
    Build PDF permissions flag.

    PDF permission bits (128-bit RC4):
        Bit 3  → Print
        Bit 4  → Modify contents
        Bit 5  → Copy / extract text
        Bit 6  → Annotations
        Bit 9  → Fill forms
        Bit 10 → Extract for accessibility
        Bit 11 → Assemble document
        Bit 12 → Print high quality
    """
    # Start with base permissions (bits 1,2 always 0; rest 1)
    permissions = 0b11111111111111111111111100000000

    if not allow_printing:
        permissions &= ~(1 << 2)  # clear bit 3
        permissions &= ~(1 << 11)  # clear bit 12 (high quality print)

    if not allow_editing:
        permissions &= ~(1 << 3)  # clear bit 4
        permissions &= ~(1 << 8)  # clear bit 9  (forms)
        permissions &= ~(1 << 10)  # clear bit 11 (assemble)

    if not allow_copying:
        permissions &= ~(1 << 4)  # clear bit 5
        permissions &= ~(1 << 9)  # clear bit 10 (accessibility)

    if not allow_annotations:
        permissions &= ~(1 << 5)  # clear bit 6

    return permissions


def decrypt_pdf(source, password: str) -> bytes:
    """
    Decrypt / remove password from a PDF.

    Args:
        source   : uploaded file object OR raw bytes
        password : password to decrypt the PDF

    Returns:
        Raw bytes of the decrypted PDF
    """
    from pypdf import PdfReader, PdfWriter

    # ── Read PDF bytes ────────────────────────────
    if hasattr(source, "read"):
        pdf_bytes = source.read()
    else:
        pdf_bytes = source

    reader = PdfReader(io.BytesIO(pdf_bytes))

    # ── Check if encrypted ─────────────────────────
    if not reader.is_encrypted:
        raise ValueError("PDF is not encrypted.")

    # ── Decrypt ────────────────────────────────────
    result = reader.decrypt(password)
    if result == 0:
        raise ValueError("Wrong password. Could not decrypt PDF.")

    # ── Copy to new writer (removes encryption) ────
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)

    if reader.metadata:
        writer.add_metadata(reader.metadata)

    # ── Serialize ─────────────────────────────────
    buffer = io.BytesIO()
    writer.write(buffer)
    buffer.seek(0)
    return buffer.read()


def get_pdf_info(source, password: str = None) -> dict:
    """
    Get PDF metadata and encryption info.

    Args:
        source   : uploaded file object OR raw bytes
        password : password if encrypted (optional)

    Returns:
        dict with PDF info
    """
    from pypdf import PdfReader

    if hasattr(source, "read"):
        pdf_bytes = source.read()
    else:
        pdf_bytes = source

    reader = PdfReader(io.BytesIO(pdf_bytes))

    # ── Decrypt if needed ──────────────────────────
    if reader.is_encrypted:
        if not password:
            return {
                "encrypted": True,
                "pages": None,
                "metadata": None,
                "permissions": None,
                "message": "PDF is encrypted. Provide password to get full info.",
            }
        result = reader.decrypt(password)
        if result == 0:
            raise ValueError("Wrong password.")

    # ── Collect metadata ───────────────────────────
    meta = {}
    if reader.metadata:
        meta = {
            "title": reader.metadata.get("/Title", ""),
            "author": reader.metadata.get("/Author", ""),
            "subject": reader.metadata.get("/Subject", ""),
            "creator": reader.metadata.get("/Creator", ""),
            "producer": reader.metadata.get("/Producer", ""),
            "created": str(reader.metadata.get("/CreationDate", "")),
            "modified": str(reader.metadata.get("/ModDate", "")),
        }

    return {
        "encrypted": reader.is_encrypted,
        "pages": len(reader.pages),
        "metadata": meta,
        "file_size": len(pdf_bytes),
        "pdf_version": reader.pdf_header,
    }


def compress_pdf(
    source,
    compression_level: str = "medium",
    password: str = None,
) -> dict:
    """
    Compress a PDF file by reducing image quality and removing redundant data.

    Args:
        source            : uploaded file object OR raw bytes
        compression_level : 'low' | 'medium' | 'high' | 'extreme'
        password          : PDF password if encrypted

    Compression levels:
        low     → light compression, best quality  (image DPI: 150, quality: 85)
        medium  → balanced compression             (image DPI: 120, quality: 72)
        high    → aggressive compression           (image DPI: 96,  quality: 60)
        extreme → maximum compression, low quality (image DPI: 72,  quality: 40)

    Returns:
        {
            'bytes'            : bytes,
            'original_size'    : int,
            'compressed_size'  : int,
            'reduction_percent': float,
            'compression_level': str,
        }
    """
    import fitz  # pymupdf
    from PIL import Image

    # ── Read PDF bytes ────────────────────────────
    if hasattr(source, "read"):
        pdf_bytes = source.read()
    else:
        pdf_bytes = source

    original_size = len(pdf_bytes)

    # ── Validate PDF ──────────────────────────────
    if not pdf_bytes.startswith(b"%PDF"):
        raise ValueError("Invalid PDF file.")

    # ── Compression level config ───────────────────
    levels = {
        "low": {"dpi": 150, "quality": 85, "deflate": 3},
        "medium": {"dpi": 120, "quality": 72, "deflate": 6},
        "high": {"dpi": 96, "quality": 60, "deflate": 7},
        "extreme": {"dpi": 72, "quality": 40, "deflate": 9},
    }

    if compression_level not in levels:
        raise ValueError(
            f'Invalid compression_level: "{compression_level}". '
            f"Must be one of: {list(levels.keys())}"
        )

    config = levels[compression_level]
    dpi = config["dpi"]
    quality = config["quality"]

    # ── Open PDF ──────────────────────────────────
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    if doc.is_encrypted:
        if not password:
            raise ValueError("PDF is encrypted. Provide a password.")
        if not doc.authenticate(password):
            raise ValueError("Wrong password.")

    # ── Step 1: Compress images on each page ───────
    for page_num in range(len(doc)):
        page = doc[page_num]

        # Get all images on this page
        image_list = page.get_images(full=True)

        for img_info in image_list:
            xref = img_info[0]  # image reference number
            base_image = doc.extract_image(xref)

            if not base_image:
                continue

            img_bytes = base_image["image"]
            img_ext = base_image["ext"]

            # Skip non-raster images (masks, etc.)
            if img_ext not in ("jpeg", "jpg", "png", "bmp", "tiff"):
                continue

            try:
                # ── Open image with PIL ───────────
                pil_img = Image.open(io.BytesIO(img_bytes))

                # Convert to RGB if needed
                if pil_img.mode in ("RGBA", "P", "LA"):
                    pil_img = pil_img.convert("RGB")
                elif pil_img.mode == "L":
                    pass  # keep grayscale as is

                # ── Resize if DPI is too high ──────
                orig_w, orig_h = pil_img.size
                if orig_w > dpi * 8 or orig_h > dpi * 11:
                    scale = min(dpi * 8 / orig_w, dpi * 11 / orig_h)
                    new_w = max(1, int(orig_w * scale))
                    new_h = max(1, int(orig_h * scale))
                    pil_img = pil_img.resize((new_w, new_h), Image.LANCZOS)

                # ── Save as JPEG with compression ──
                img_buffer = io.BytesIO()
                pil_img.save(
                    img_buffer,
                    format="JPEG",
                    quality=quality,
                    optimize=True,
                )
                img_buffer.seek(0)
                new_img_bytes = img_buffer.read()

                # Only replace if new image is smaller
                if len(new_img_bytes) < len(img_bytes):
                    doc.update_stream(xref, new_img_bytes)

            except Exception:
                continue  # skip problematic images silently

    # ── Step 2: Clean and compress PDF structure ───
    buffer = io.BytesIO()
    doc.save(
        buffer,
        garbage=4,  # remove unused objects (0-4, 4=most aggressive)
        deflate=True,  # compress streams
        deflate_images=True,  # compress images
        deflate_fonts=True,  # compress fonts
        clean=True,  # clean content streams
        pretty=False,  # no pretty printing (saves space)
        linear=False,  # no linearization (saves space)
    )
    doc.close()

    buffer.seek(0)
    compressed_bytes = buffer.read()
    compressed_size = len(compressed_bytes)

    # ── Calculate reduction ────────────────────────
    reduction = (
        (original_size - compressed_size) / original_size * 100
        if original_size > 0
        else 0
    )

    return {
        "bytes": compressed_bytes,
        "original_size": original_size,
        "compressed_size": compressed_size,
        "reduction_percent": round(reduction, 2),
        "compression_level": compression_level,
        "original_size_kb": round(original_size / 1024, 2),
        "compressed_size_kb": round(compressed_size / 1024, 2),
        "original_size_mb": round(original_size / (1024 * 1024), 2),
        "compressed_size_mb": round(compressed_size / (1024 * 1024), 2),
    }


def pdf_to_png(
    source,
    dpi: int = 200,
    pages: list = None,
    password: str = None,
) -> list[dict]:
    """
    Convert PDF (file or bytes) → list of PNG images (one per page).

    Args:
        source   : uploaded file object OR raw bytes
        dpi      : image resolution (default: 200)
        pages    : list of page numbers to convert (1-based)
                   None = convert all pages
        password : PDF password if encrypted

    Returns:
        list of {
            'page'    : int,
            'bytes'   : bytes,
            'width'   : int,
            'height'  : int,
            'filename': str,
        }
    """
    import fitz  # pymupdf
    from PIL import Image

    # ── Read PDF bytes ────────────────────────────
    if hasattr(source, "read"):
        pdf_bytes = source.read()
    else:
        pdf_bytes = source

    # ── Validate PDF ──────────────────────────────
    if not pdf_bytes.startswith(b"%PDF"):
        raise ValueError("Invalid PDF file.")

    # ── Open PDF ──────────────────────────────────
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    if doc.is_encrypted:
        if not password:
            raise ValueError("PDF is encrypted. Provide a password.")
        if not doc.authenticate(password):
            raise ValueError("Wrong password.")

    total_pages = len(doc)

    # ── Validate pages ────────────────────────────
    if pages is not None:
        invalid = [p for p in pages if not (1 <= p <= total_pages)]
        if invalid:
            raise ValueError(
                f"Invalid page numbers: {invalid}. "
                f"PDF has {total_pages} pages (1-{total_pages})."
            )
        pages_to_convert = pages
    else:
        pages_to_convert = list(range(1, total_pages + 1))

    # ── Convert each page → PNG ───────────────────
    zoom = dpi / 72  # pymupdf base DPI is 72
    matrix = fitz.Matrix(zoom, zoom)
    results = []

    for page_num in pages_to_convert:
        page = doc[page_num - 1]  # 0-based index

        # ── Render page → pixmap ──────────────────
        pixmap = page.get_pixmap(
            matrix=matrix,
            alpha=True,  # PNG supports transparency
        )

        # ── Pixmap → PIL Image → PNG bytes ────────
        image = Image.frombytes(
            "RGBA",
            [pixmap.width, pixmap.height],
            pixmap.samples,
        )

        buffer = io.BytesIO()
        image.save(
            buffer,
            format="PNG",
            optimize=True,
            compress_level=6,  # 0-9, 6 is balanced
        )
        buffer.seek(0)
        png_bytes = buffer.read()

        results.append(
            {
                "page": page_num,
                "bytes": png_bytes,
                "width": pixmap.width,
                "height": pixmap.height,
                "filename": f"page_{page_num}.png",
                "size_kb": round(len(png_bytes) / 1024, 2),
            }
        )

    doc.close()
    return results


def rotate_pdf(
    source,
    rotation : int  = 90,
    pages    : list = None,
    password : str  = None,
) -> tuple:
    """
    Rotate pages in a PDF.

    Args:
        source   : uploaded file object OR raw bytes
        rotation : degrees to rotate (90, 180, 270, -90)
        pages    : list of page numbers to rotate (1-based)
                   None = rotate all pages
        password : PDF password if encrypted

    Returns:
        (rotated_bytes, total_pages, rotated_count)
    """
    from pypdf import PdfReader, PdfWriter

    # ── Read PDF bytes ────────────────────────────
    if hasattr(source, 'read'):
        pdf_bytes = source.read()
    else:
        pdf_bytes = source

    # ── Validate PDF ──────────────────────────────
    if not pdf_bytes.startswith(b'%PDF'):
        raise ValueError('Invalid PDF file.')

    # ── Normalize rotation ─────────────────────────
    valid_rotations = [90, 180, 270, -90, -180, -270]
    if rotation not in valid_rotations:
        raise ValueError(
            f'Invalid rotation: {rotation}. '
            f'Must be one of: {valid_rotations}'
        )
    rotation = rotation % 360

    # ── Read PDF ──────────────────────────────────
    reader = PdfReader(io.BytesIO(pdf_bytes))

    # ── Decrypt if needed ──────────────────────────
    if reader.is_encrypted:
        if not password:
            raise ValueError('PDF is encrypted. Provide a password.')
        if reader.decrypt(password) == 0:
            raise ValueError('Wrong password.')

    total_pages = len(reader.pages)

    # ── Validate page numbers ──────────────────────
    if pages is not None:
        invalid = [p for p in pages if not (1 <= p <= total_pages)]
        if invalid:
            raise ValueError(
                f'Invalid page numbers: {invalid}. '
                f'PDF has {total_pages} pages (1-{total_pages}).'
            )
        pages_to_rotate = set(pages)
    else:
        pages_to_rotate = set(range(1, total_pages + 1))

    # ── Rotate + write ─────────────────────────────
    writer = PdfWriter()

    for i, page in enumerate(reader.pages, start=1):
        if i in pages_to_rotate:
            page.rotate(rotation)
        writer.add_page(page)

    # ── Copy metadata ──────────────────────────────
    if reader.metadata:
        writer.add_metadata(reader.metadata)

    # ── Serialize ─────────────────────────────────
    buffer = io.BytesIO()
    writer.write(buffer)
    buffer.seek(0)

    return buffer.read(), total_pages, len(pages_to_rotate)


def get_pdf_page_info(source, password: str = None) -> list:
    """
    Get rotation and dimension info for each page.

    Args:
        source   : uploaded file object OR raw bytes
        password : PDF password if encrypted

    Returns:
        list of { page, width, height, rotation }
    """
    from pypdf import PdfReader

    if hasattr(source, 'read'):
        pdf_bytes = source.read()
    else:
        pdf_bytes = source

    reader = PdfReader(io.BytesIO(pdf_bytes))

    if reader.is_encrypted:
        if not password:
            raise ValueError('PDF is encrypted. Provide a password.')
        if reader.decrypt(password) == 0:
            raise ValueError('Wrong password.')

    result = []
    for i, page in enumerate(reader.pages, start=1):
        result.append({
            'page'    : i,
            'width'   : float(page.mediabox.width),
            'height'  : float(page.mediabox.height),
            'rotation': page.rotation,
        })

    return result


def watermark_pdf(
    source,
    watermark_text: str = "CONFIDENTIAL",
    watermark_type: str = "text",
    watermark_image: bytes = None,
    opacity: float = 0.3,
    font_size: int = 60,
    color: str = "red",
    angle: int = 45,
    position: str = "center",
    pages: list = None,
    password: str = None,
) -> bytes:
    """
    Add text or image watermark to a PDF.
    Returns raw bytes of the watermarked PDF.
    """
    from pypdf import PdfReader, PdfWriter
    from reportlab.lib import colors as rl_colors
    from reportlab.pdfgen import canvas as rl_canvas

    # ── Read PDF bytes ────────────────────────────
    if hasattr(source, "read"):
        pdf_bytes = source.read()
    else:
        pdf_bytes = source

    if not pdf_bytes.startswith(b"%PDF"):
        raise ValueError("Invalid PDF file.")

    # ── Validate ──────────────────────────────────
    valid_positions = (
        "center",
        "top",
        "bottom",
        "top-left",
        "top-right",
        "bottom-left",
        "bottom-right",
    )
    if position not in valid_positions:
        raise ValueError(f"Invalid position. Must be one of: {valid_positions}")
    if not (0.0 <= opacity <= 1.0):
        raise ValueError("opacity must be between 0.0 and 1.0.")
    if watermark_type not in ("text", "image"):
        raise ValueError('watermark_type must be "text" or "image".')
    if watermark_type == "image" and not watermark_image:
        raise ValueError("watermark_image bytes are required for image watermark.")

    # ── Color map ─────────────────────────────────
    color_map = {
        "red": rl_colors.red,
        "blue": rl_colors.blue,
        "grey": rl_colors.grey,
        "gray": rl_colors.grey,
        "black": rl_colors.black,
        "green": rl_colors.green,
        "yellow": rl_colors.yellow,
        "white": rl_colors.white,
    }
    wm_color = color_map.get(color.lower(), rl_colors.red)

    # ── Read PDF ──────────────────────────────────
    reader = PdfReader(io.BytesIO(pdf_bytes))

    if reader.is_encrypted:
        if not password:
            raise ValueError("PDF is encrypted. Provide a password.")
        if reader.decrypt(password) == 0:
            raise ValueError("Wrong password.")

    total_pages = len(reader.pages)

    # ── Pages to watermark ────────────────────────
    if pages is not None:
        invalid = [p for p in pages if not (1 <= p <= total_pages)]
        if invalid:
            raise ValueError(
                f"Invalid page numbers: {invalid}. "
                f"PDF has {total_pages} pages (1-{total_pages})."
            )
        pages_to_watermark = set(pages)
    else:
        pages_to_watermark = set(range(1, total_pages + 1))

    # ── Position calculator ────────────────────────
    def get_position(pos, pw, ph):
        padding = 60
        return {
            "center": (pw / 2, ph / 2),
            "top": (pw / 2, ph - padding),
            "bottom": (pw / 2, padding),
            "top-left": (padding, ph - padding),
            "top-right": (pw - padding, ph - padding),
            "bottom-left": (padding, padding),
            "bottom-right": (pw - padding, padding),
        }.get(pos, (pw / 2, ph / 2))

    # ── Build text watermark page ──────────────────
    def make_text_watermark(pw, ph) -> bytes:
        buf = io.BytesIO()
        c = rl_canvas.Canvas(buf, pagesize=(pw, ph))

        x, y = get_position(position, pw, ph)

        c.saveState()
        c.setFillColor(wm_color, alpha=opacity)
        c.setFont("Helvetica-Bold", font_size)
        c.translate(x, y)
        c.rotate(angle)
        c.drawCentredString(0, 0, watermark_text)
        c.restoreState()

        c.save()
        buf.seek(0)
        return buf.read()

    # ── Build image watermark page ─────────────────
    def make_image_watermark(pw, ph) -> bytes:
        from PIL import Image as PILImage
        from reportlab.lib.utils import ImageReader

        buf = io.BytesIO()
        c = rl_canvas.Canvas(buf, pagesize=(pw, ph))

        # ── Open + convert image ───────────────────
        pil_img = PILImage.open(io.BytesIO(watermark_image))

        # Convert palette/RGBA → RGBA for consistent handling
        if pil_img.mode == "P":
            pil_img = pil_img.convert("RGBA")
        if pil_img.mode not in ("RGB", "RGBA", "L"):
            pil_img = pil_img.convert("RGBA")

        # ── Apply opacity to image ─────────────────
        if pil_img.mode == "RGBA":
            r, g, b, a = pil_img.split()
            # Scale alpha channel by opacity
            a = a.point(lambda px: int(px * opacity))
            pil_img = PILImage.merge("RGBA", (r, g, b, a))
        else:
            # For RGB images, convert to RGBA and apply opacity
            pil_img = pil_img.convert("RGBA")
            r, g, b, a = pil_img.split()
            a = a.point(lambda px: int(px * opacity))
            pil_img = PILImage.merge("RGBA", (r, g, b, a))

        # ── Save processed image to buffer ─────────
        img_buf = io.BytesIO()
        pil_img.save(img_buf, format="PNG")
        img_buf.seek(0)

        # ── Scale image to fit page (max 40%) ──────
        orig_w, orig_h = pil_img.size
        max_w = pw * 0.4
        max_h = ph * 0.4
        scale = min(max_w / orig_w, max_h / orig_h)
        draw_w = orig_w * scale
        draw_h = orig_h * scale

        # ── Get anchor position ────────────────────
        x, y = get_position(position, pw, ph)

        # ── Draw image ─────────────────────────────
        c.saveState()
        c.translate(x, y)
        c.rotate(angle)

        # Draw centered on anchor point
        c.drawImage(
            ImageReader(img_buf),
            x=-draw_w / 2,
            y=-draw_h / 2,
            width=draw_w,
            height=draw_h,
            mask="auto",
        )
        c.restoreState()

        c.save()
        buf.seek(0)
        return buf.read()

    # ── Apply watermark to each page ──────────────
    writer = PdfWriter()

    for i, page in enumerate(reader.pages, start=1):
        if i in pages_to_watermark:
            pw = float(page.mediabox.width)
            ph = float(page.mediabox.height)

            try:
                if watermark_type == "image":
                    wm_bytes = make_image_watermark(pw, ph)
                else:
                    wm_bytes = make_text_watermark(pw, ph)

                wm_page = PdfReader(io.BytesIO(wm_bytes)).pages[0]
                page.merge_page(wm_page)

            except Exception as e:
                # If watermark fails on a page, keep original page
                raise RuntimeError(f"Watermark failed on page {i}: {e}")

        writer.add_page(page)

    # ── Copy metadata ──────────────────────────────
    if reader.metadata:
        writer.add_metadata(reader.metadata)

    # ── Serialize → always return bytes ───────────
    buffer = io.BytesIO()
    writer.write(buffer)
    buffer.seek(0)
    return buffer.read()  # ← always bytes, never None


def format_json(
    source,
    indent: int = 4,
    sort_keys: bool = False,
    minify: bool = False,
    ensure_ascii: bool = False,
) -> dict:
    """
    Format / validate / minify JSON from file or raw text.

    Args:
        source       : uploaded file object OR raw JSON string
        indent       : indentation spaces (default: 4)
        sort_keys    : sort keys alphabetically (default: False)
        minify       : minify JSON (overrides indent) (default: False)
        ensure_ascii : escape non-ASCII chars (default: False)

    Returns:
        {
            'formatted'  : str,
            'is_valid'   : bool,
            'key_count'  : int,
            'size_original' : int,
            'size_formatted': int,
            'type'       : 'object' | 'array' | 'other',
            'depth'      : int,
        }
    """
    # ── Read source ───────────────────────────────
    if hasattr(source, "read"):
        raw = source.read()
        if isinstance(raw, bytes):
            raw = raw.decode("utf-8")
    else:
        raw = source

    raw = raw.strip()

    if not raw:
        raise ValueError("Empty input. Provide a valid JSON string or file.")

    # ── Parse JSON ────────────────────────────────
    try:
        parsed = json.loads(raw)
    except json.JSONDecodeError as e:
        return {
            "formatted": None,
            "is_valid": False,
            "error": str(e),
            "error_line": e.lineno,
            "error_column": e.colno,
            "error_position": e.pos,
            "size_original": len(raw),
        }

    # ── Format ────────────────────────────────────
    if minify:
        formatted = json.dumps(
            parsed,
            separators=(",", ":"),
            ensure_ascii=ensure_ascii,
        )
    else:
        formatted = json.dumps(
            parsed,
            indent=indent,
            sort_keys=sort_keys,
            ensure_ascii=ensure_ascii,
        )

    # ── Analyze structure ─────────────────────────
    def get_depth(obj, level=0):
        if isinstance(obj, dict):
            if not obj:
                return level
            return max(get_depth(v, level + 1) for v in obj.values())
        if isinstance(obj, list):
            if not obj:
                return level
            return max(get_depth(i, level + 1) for i in obj)
        return level

    def count_keys(obj):
        if isinstance(obj, dict):
            return len(obj) + sum(count_keys(v) for v in obj.values())
        if isinstance(obj, list):
            return sum(count_keys(i) for i in obj)
        return 0

    json_type = (
        "object"
        if isinstance(parsed, dict)
        else "array" if isinstance(parsed, list) else "other"
    )

    return {
        "formatted": formatted,
        "is_valid": True,
        "type": json_type,
        "depth": get_depth(parsed),
        "key_count": count_keys(parsed),
        "item_count": len(parsed) if isinstance(parsed, (dict, list)) else 1,
        "size_original": len(raw),
        "size_formatted": len(formatted),
        "size_original_kb": round(len(raw) / 1024, 2),
        "size_formatted_kb": round(len(formatted) / 1024, 2),
        "minified": minify,
        "sorted_keys": sort_keys,
        "indent": indent,
    }


def ocr_pdf(
    source,
    language: str = "eng",
    dpi: int = 300,
    pages: list = None,
    password: str = None,
) -> dict:
    """
    Extract text from a PDF using pymupdf only.
    No Tesseract or system dependencies needed.

    Strategy:
        1. Try direct text extraction (for digital/selectable PDFs)
        2. If page has no text → render to image → extract via pymupdf OCR
        3. Combine results per page

    Args:
        source   : uploaded file object OR raw bytes
        language : language hint (eng, ara, fra, etc.)  (default: eng)
        dpi      : render resolution for image pages    (default: 300)
        pages    : list of page numbers (1-based)       (default: all)
        password : PDF password if encrypted

    Returns:
        {
            'full_text'   : str,
            'pages'       : [ { page, text, word_count, char_count, method } ],
            'total_pages' : int,
            'word_count'  : int,
            'char_count'  : int,
            'language'    : str,
            'dpi'         : int,
        }
    """
    import fitz  # pymupdf
    from PIL import Image, ImageEnhance, ImageFilter

    # ── Read PDF bytes ────────────────────────────
    if hasattr(source, "read"):
        pdf_bytes = source.read()
    else:
        pdf_bytes = source

    if not pdf_bytes.startswith(b"%PDF"):
        raise ValueError("Invalid PDF file.")

    # ── Open PDF ──────────────────────────────────
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    if doc.is_encrypted:
        if not password:
            raise ValueError("PDF is encrypted. Provide a password.")
        if not doc.authenticate(password):
            raise ValueError("Wrong password.")

    total_pages = len(doc)

    # ── Validate pages ────────────────────────────
    if pages is not None:
        invalid = [p for p in pages if not (1 <= p <= total_pages)]
        if invalid:
            raise ValueError(
                f"Invalid page numbers: {invalid}. "
                f"PDF has {total_pages} pages (1-{total_pages})."
            )
        pages_to_ocr = pages
    else:
        pages_to_ocr = list(range(1, total_pages + 1))

    # ── Render config ─────────────────────────────
    zoom = dpi / 72
    matrix = fitz.Matrix(zoom, zoom)

    # ── Process each page ─────────────────────────
    page_results = []
    full_texts = []

    for page_num in pages_to_ocr:
        page = doc[page_num - 1]

        # ── Step 1: Try direct text extraction ────
        direct_text = page.get_text("text").strip()

        if direct_text and len(direct_text) > 20:
            # Good digital text found — use directly
            text = direct_text
            method = "digital"

        else:
            # ── Step 2: Image-based extraction ────
            # Render page to high-res image
            pixmap = page.get_pixmap(matrix=matrix, alpha=False)

            # Pixmap → PIL Image
            pil_img = Image.frombytes(
                "RGB",
                [pixmap.width, pixmap.height],
                pixmap.samples,
            )

            # ── Preprocess image ──────────────────
            pil_img = _preprocess_image_for_ocr(pil_img)

            # ── Save processed image to bytes ──────
            img_buffer = io.BytesIO()
            pil_img.save(img_buffer, format="PNG", optimize=True)
            img_buffer.seek(0)

            # ── Re-open with pymupdf for text ──────
            # Create a temporary single-page PDF from image
            img_doc = fitz.open(stream=img_buffer.read(), filetype="png")
            img_page = img_doc[0]

            # Extract text blocks from rendered image
            # pymupdf can extract text from image-based pages too
            blocks = img_page.get_text("blocks")
            text = "\n".join(b[4].strip() for b in blocks if b[4].strip())

            if not text:
                # Final fallback: get all text from original page
                text = page.get_text("words")
                text = " ".join(w[4] for w in text).strip()

            img_doc.close()
            method = "image"

        word_count = len(text.split()) if text else 0
        char_count = len(text)

        page_results.append(
            {
                "page": page_num,
                "text": text,
                "word_count": word_count,
                "char_count": char_count,
                "method": method,
            }
        )

        full_texts.append(f"--- Page {page_num} ---\n{text}")

    doc.close()

    full_text = "\n\n".join(full_texts)
    word_count = sum(p["word_count"] for p in page_results)
    char_count = sum(p["char_count"] for p in page_results)

    return {
        "full_text": full_text,
        "pages": page_results,
        "total_pages": len(page_results),
        "word_count": word_count,
        "char_count": char_count,
        "language": language,
        "dpi": dpi,
    }


def _preprocess_image_for_ocr(image) -> "Image":
    """
    Preprocess PIL image for better text extraction accuracy.

    Steps:
        1. Convert to grayscale
        2. Enhance contrast
        3. Sharpen edges
        4. Denoise
    """
    from PIL import Image, ImageEnhance, ImageFilter

    # ── Upscale small images ──────────────────────
    w, h = image.size
    if w < 1000 or h < 1000:
        scale = max(1000 / w, 1000 / h)
        new_w = int(w * scale)
        new_h = int(h * scale)
        image = image.resize((new_w, new_h), Image.LANCZOS)

    # ── Convert to grayscale ──────────────────────
    gray = image.convert("L")

    # ── Enhance contrast ──────────────────────────
    contrast = ImageEnhance.Contrast(gray)
    enhanced = contrast.enhance(2.0)

    # ── Sharpen ───────────────────────────────────
    sharpener = ImageEnhance.Sharpness(enhanced)
    sharpened = sharpener.enhance(2.0)

    # ── Remove noise ──────────────────────────────
    denoised = sharpened.filter(ImageFilter.MedianFilter(size=3))

    return denoised


def convert_timestamp(
    value: str,
    from_tz: str = "UTC",
    to_tz: str = "UTC",
    from_format: str = None,
    to_format: str = None,
) -> dict:
    """
    Convert timestamps between formats and timezones.

    Args:
        value       : timestamp value to convert
                      - Unix timestamp  : '1704067200'
                      - ISO 8601        : '2024-01-01T00:00:00Z'
                      - Human readable  : '2024-01-01 12:00:00'
                      - Date only       : '2024-01-01'
                      - Relative        : 'now', 'today', 'yesterday'
        from_tz     : source timezone               (default: UTC)
        to_tz       : target timezone               (default: UTC)
        from_format : input format (strptime)       (default: auto-detect)
        to_format   : output format (strftime)      (default: all formats)

    Returns:
        dict with all converted formats
    """
    from datetime import datetime, timezone, timedelta
    import time
    import calendar

    try:
        import zoneinfo

        def get_tz(name):
            if name.upper() == "UTC":
                return timezone.utc
            return zoneinfo.ZoneInfo(name)

    except ImportError:
        try:
            import pytz

            def get_tz(name):
                if name.upper() == "UTC":
                    return pytz.utc
                return pytz.timezone(name)

        except ImportError:

            def get_tz(name):
                return timezone.utc

    # ── Parse relative values ─────────────────────
    now = datetime.now(timezone.utc)

    relative_map = {
        "now": now,
        "today": now.replace(hour=0, minute=0, second=0, microsecond=0),
        "yesterday": (now - timedelta(days=1)).replace(
            hour=0, minute=0, second=0, microsecond=0
        ),
        "tomorrow": (now + timedelta(days=1)).replace(
            hour=0, minute=0, second=0, microsecond=0
        ),
    }

    value = str(value).strip()

    if value.lower() in relative_map:
        dt = relative_map[value.lower()]

    # ── Parse Unix timestamp ───────────────────────
    elif value.lstrip("-").replace(".", "").isdigit():
        ts = float(value)

        # Detect milliseconds vs seconds
        if abs(ts) > 1e10:
            ts = ts / 1000  # convert ms → seconds

        dt = datetime.fromtimestamp(ts, tz=timezone.utc)

    # ── Parse with custom format ───────────────────
    elif from_format:
        try:
            dt = datetime.strptime(value, from_format)
            if dt.tzinfo is None:
                src_tz = get_tz(from_tz)
                dt = dt.replace(tzinfo=src_tz)
        except ValueError as e:
            raise ValueError(f'Cannot parse "{value}" with format "{from_format}": {e}')

    # ── Auto-detect format ─────────────────────────
    else:
        dt = _parse_datetime_auto(value, from_tz, get_tz)

    # ── Convert to target timezone ─────────────────
    try:
        target_tz = get_tz(to_tz)
        dt_target = dt.astimezone(target_tz)
    except Exception:
        dt_target = dt

    # ── Unix timestamps ────────────────────────────
    unix_seconds = int(dt_target.timestamp())
    unix_ms = int(dt_target.timestamp() * 1000)
    unix_ns = int(dt_target.timestamp() * 1_000_000_000)

    # ── Format output ─────────────────────────────
    formats = {
        "unix_seconds": unix_seconds,
        "unix_ms": unix_ms,
        "unix_ns": unix_ns,
        "iso_8601": dt_target.strftime("%Y-%m-%dT%H:%M:%S") + _tz_offset(dt_target),
        "iso_8601_utc": dt.astimezone(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
        "rfc_2822": dt_target.strftime("%a, %d %b %Y %H:%M:%S ")
        + _tz_offset(dt_target),
        "rfc_3339": dt_target.strftime("%Y-%m-%dT%H:%M:%S") + _tz_offset(dt_target),
        "human_readable": dt_target.strftime("%B %d, %Y %I:%M:%S %p"),
        "date_only": dt_target.strftime("%Y-%m-%d"),
        "time_only": dt_target.strftime("%H:%M:%S"),
        "datetime_local": dt_target.strftime("%Y-%m-%d %H:%M:%S"),
        "day_of_week": dt_target.strftime("%A"),
        "day_of_year": dt_target.timetuple().tm_yday,
        "week_number": dt_target.strftime("%W"),
        "quarter": f"Q{(dt_target.month - 1) // 3 + 1}",
        "relative": _get_relative_time(dt_target),
        "utc_offset": _tz_offset(dt_target),
        "timezone": to_tz,
    }

    if to_format:
        try:
            formats["custom"] = dt_target.strftime(to_format)
        except Exception as e:
            formats["custom_error"] = str(e)

    return {
        "input": value,
        "from_timezone": from_tz,
        "to_timezone": to_tz,
        "is_dst": _is_dst(dt_target),
        "formats": formats,
    }


def _parse_datetime_auto(value: str, from_tz: str, get_tz) -> "datetime":
    """Auto-detect and parse datetime string."""
    from datetime import datetime, timezone

    # Common formats to try in order
    format_attempts = [
        "%Y-%m-%dT%H:%M:%SZ",  # ISO 8601 UTC
        "%Y-%m-%dT%H:%M:%S%z",  # ISO 8601 with offset
        "%Y-%m-%dT%H:%M:%S",  # ISO 8601 no tz
        "%Y-%m-%dT%H:%M:%S.%fZ",  # ISO 8601 with microseconds UTC
        "%Y-%m-%dT%H:%M:%S.%f%z",  # ISO 8601 with microseconds
        "%Y-%m-%d %H:%M:%S",  # Common datetime
        "%Y-%m-%d %H:%M",  # No seconds
        "%Y-%m-%d",  # Date only
        "%d/%m/%Y %H:%M:%S",  # DD/MM/YYYY
        "%d/%m/%Y %H:%M",  # DD/MM/YYYY no seconds
        "%d/%m/%Y",  # DD/MM/YYYY date only
        "%m/%d/%Y %H:%M:%S",  # US format
        "%m/%d/%Y %H:%M",  # US format no seconds
        "%m/%d/%Y",  # US date only
        "%d-%m-%Y %H:%M:%S",  # DD-MM-YYYY
        "%d-%m-%Y",  # DD-MM-YYYY date only
        "%B %d, %Y %I:%M:%S %p",  # Human readable
        "%B %d, %Y",  # Month Day Year
        "%b %d, %Y %H:%M:%S",  # Short month
        "%b %d, %Y",  # Short month date only
        "%a, %d %b %Y %H:%M:%S %z",  # RFC 2822
        "%Y%m%d",  # Compact date
        "%Y%m%dT%H%M%S",  # Compact datetime
    ]

    for fmt in format_attempts:
        try:
            dt = datetime.strptime(value, fmt)
            if dt.tzinfo is None:
                src_tz = get_tz(from_tz)
                dt = dt.replace(tzinfo=src_tz)
            return dt
        except ValueError:
            continue

    raise ValueError(
        f'Cannot parse timestamp: "{value}". ' f"Try providing from_format explicitly."
    )


def _tz_offset(dt) -> str:
    """Get timezone offset string like +05:00 or Z."""
    from datetime import timezone

    if dt.tzinfo is None:
        return ""
    offset = dt.utcoffset()
    if offset is None:
        return ""
    total_seconds = int(offset.total_seconds())
    if total_seconds == 0:
        return "+00:00"
    sign = "+" if total_seconds >= 0 else "-"
    abs_sec = abs(total_seconds)
    hours = abs_sec // 3600
    minutes = (abs_sec % 3600) // 60
    return f"{sign}{hours:02d}:{minutes:02d}"


def _get_relative_time(dt) -> str:
    """Get human-readable relative time like '2 hours ago'."""
    from datetime import datetime, timezone

    now = datetime.now(timezone.utc)
    delta = now - dt.astimezone(timezone.utc)
    secs = int(delta.total_seconds())
    abs_s = abs(secs)
    future = secs < 0

    def fmt(n, unit):
        label = f'{n} {unit}{"s" if n != 1 else ""}'
        return f"in {label}" if future else f"{label} ago"

    if abs_s < 5:
        return "just now"
    if abs_s < 60:
        return fmt(abs_s, "second")
    if abs_s < 3600:
        return fmt(abs_s // 60, "minute")
    if abs_s < 86400:
        return fmt(abs_s // 3600, "hour")
    if abs_s < 604800:
        return fmt(abs_s // 86400, "day")
    if abs_s < 2592000:
        return fmt(abs_s // 604800, "week")
    if abs_s < 31536000:
        return fmt(abs_s // 2592000, "month")
    return fmt(abs_s // 31536000, "year")


def _is_dst(dt) -> bool:
    """Check if datetime is in daylight saving time."""
    try:
        import time

        ts = dt.timestamp()
        local_tm = time.localtime(ts)
        return bool(local_tm.tm_isdst)
    except Exception:
        return False


def batch_convert_timestamps(
    values: list,
    from_tz: str = "UTC",
    to_tz: str = "UTC",
    to_format: str = None,
) -> list:
    """
    Convert multiple timestamps at once.

    Args:
        values    : list of timestamp strings
        from_tz   : source timezone
        to_tz     : target timezone
        to_format : output format (strftime)

    Returns:
        list of conversion results
    """
    results = []
    for value in values:
        try:
            result = convert_timestamp(
                value,
                from_tz=from_tz,
                to_tz=to_tz,
                to_format=to_format,
            )
            results.append({"input": value, "success": True, **result})
        except Exception as e:
            results.append(
                {
                    "input": value,
                    "success": False,
                    "error": str(e),
                }
            )
    return results


def base64_encode(
    source,
    encoding: str = "standard",
    chunk_size: int = 0,
    filename: str = None,
) -> dict:
    """
    Encode text, file, or bytes to Base64.

    Args:
        source     : text string | file object | bytes
        encoding   : 'standard' | 'url_safe' | 'mime'  (default: standard)
        chunk_size : split output into chunks of N chars (default: 0 = no split)
        filename   : original filename for file inputs

    Returns:
        {
            'encoded'        : str,
            'encoding'       : str,
            'original_size'  : int,
            'encoded_size'   : int,
            'is_file'        : bool,
            'filename'       : str,
            'mime_type'      : str,
            'data_uri'       : str,
        }
    """
    import base64
    import mimetypes

    # ── Read source ───────────────────────────────
    is_file = False
    mime_type = "text/plain"

    if hasattr(source, "read"):
        # File object
        raw_bytes = source.read()
        is_file = True
        if filename:
            mime_type = mimetypes.guess_type(filename)[0] or "application/octet-stream"
    elif isinstance(source, bytes):
        raw_bytes = source
    elif isinstance(source, str):
        raw_bytes = source.encode("utf-8")
    else:
        raise ValueError("source must be a string, bytes, or file object.")

    if not raw_bytes:
        raise ValueError("Empty input. Nothing to encode.")

    original_size = len(raw_bytes)

    # ── Encode ────────────────────────────────────
    if encoding == "url_safe":
        encoded_bytes = base64.urlsafe_b64encode(raw_bytes)
    elif encoding == "mime":
        # MIME encoding adds line breaks every 76 chars
        import base64 as b64

        encoded_bytes = b64.encodebytes(raw_bytes)
        encoded_str = encoded_bytes.decode("utf-8").strip()
    else:
        encoded_bytes = base64.b64encode(raw_bytes)

    if encoding != "mime":
        encoded_str = encoded_bytes.decode("utf-8")

    # ── Chunk output ──────────────────────────────
    if chunk_size > 0 and encoding != "mime":
        encoded_str = "\n".join(
            encoded_str[i : i + chunk_size]
            for i in range(0, len(encoded_str), chunk_size)
        )

    # ── Data URI ──────────────────────────────────
    clean_encoded = encoded_str.replace("\n", "")
    data_uri = f"data:{mime_type};base64,{clean_encoded}"

    return {
        "encoded": encoded_str,
        "encoding": encoding,
        "original_size": original_size,
        "encoded_size": len(encoded_str),
        "original_size_kb": round(original_size / 1024, 2),
        "encoded_size_kb": round(len(encoded_str) / 1024, 2),
        "is_file": is_file,
        "filename": filename or "",
        "mime_type": mime_type,
        "data_uri": data_uri if is_file else "",
        "overhead_percent": round(
            (len(encoded_str) - original_size) / original_size * 100, 2
        ),
    }


def base64_decode(
    source,
    encoding: str = "standard",
    as_text: bool = True,
    text_encoding: str = "utf-8",
) -> dict:
    """
    Decode Base64 string to text or bytes.

    Args:
        source        : Base64 encoded string | file object
        encoding      : 'standard' | 'url_safe'    (default: standard)
        as_text       : try to decode as text       (default: True)
        text_encoding : text encoding to use        (default: utf-8)

    Returns:
        {
            'decoded_text'  : str | None,
            'decoded_bytes' : bytes,
            'is_text'       : bool,
            'original_size' : int,
            'decoded_size'  : int,
            'encoding'      : str,
        }
    """
    import base64

    # ── Read source ───────────────────────────────
    if hasattr(source, "read"):
        raw = source.read()
        if isinstance(raw, bytes):
            raw = raw.decode("utf-8")
    elif isinstance(source, str):
        raw = source
    else:
        raise ValueError("source must be a string or file object.")

    # ── Clean input ───────────────────────────────
    # Remove whitespace, newlines, data URI prefix
    raw = raw.strip()
    if raw.startswith("data:"):
        # Strip data URI prefix: data:image/png;base64,xxxxx
        raw = raw.split(",", 1)[-1]

    # Remove all whitespace (handles MIME line breaks)
    raw = "".join(raw.split())

    if not raw:
        raise ValueError("Empty input. Nothing to decode.")

    # ── Add padding if needed ─────────────────────
    missing_padding = len(raw) % 4
    if missing_padding:
        raw += "=" * (4 - missing_padding)

    # ── Decode ────────────────────────────────────
    try:
        if encoding == "url_safe":
            decoded_bytes = base64.urlsafe_b64decode(raw)
        else:
            decoded_bytes = base64.b64decode(raw)
    except Exception as e:
        raise ValueError(f"Invalid Base64 input: {e}")

    decoded_size = len(decoded_bytes)
    original_size = len(raw)

    # ── Try decode as text ────────────────────────
    decoded_text = None
    is_text = False

    if as_text:
        try:
            decoded_text = decoded_bytes.decode(text_encoding)
            is_text = True
        except (UnicodeDecodeError, LookupError):
            # Binary content — not text
            is_text = False
            decoded_text = None

    return {
        "decoded_text": decoded_text,
        "decoded_bytes": decoded_bytes,
        "is_text": is_text,
        "encoding": encoding,
        "original_size": original_size,
        "decoded_size": decoded_size,
        "original_size_kb": round(original_size / 1024, 2),
        "decoded_size_kb": round(decoded_size / 1024, 2),
        "text_encoding": text_encoding if is_text else None,
    }


def base64_validate(source) -> dict:
    """
    Validate whether a string is valid Base64.

    Args:
        source : string to validate

    Returns:
        { 'is_valid', 'encoding', 'decoded_size', 'message' }
    """
    import base64 as b64

    if hasattr(source, "read"):
        raw = source.read().decode("utf-8")
    else:
        raw = str(source).strip()

    # Strip data URI prefix
    if raw.startswith("data:"):
        raw = raw.split(",", 1)[-1]

    raw = "".join(raw.split())

    # Check character set
    standard_chars = set(
        "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789+/="
    )
    url_safe_chars = set(
        "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789-_="
    )

    is_url_safe = all(c in url_safe_chars for c in raw) and ("-" in raw or "_" in raw)

    # Add padding
    padded = raw + "=" * (-len(raw) % 4)

    for enc_name, decode_fn in [
        ("standard", b64.b64decode),
        ("url_safe", b64.urlsafe_b64decode),
    ]:
        try:
            decoded = decode_fn(padded)
            decoded_size = len(decoded)
            return {
                "is_valid": True,
                "encoding": enc_name,
                "decoded_size": decoded_size,
                "decoded_size_kb": round(decoded_size / 1024, 2),
                "input_size": len(raw),
                "message": f"Valid {enc_name} Base64 string.",
            }
        except Exception:
            continue

    return {
        "is_valid": False,
        "encoding": None,
        "message": "Invalid Base64 string.",
        "input_size": len(raw),
    }


def generate_uuid(
    version: int = 4,
    count: int = 1,
    uppercase: bool = False,
    hyphens: bool = True,
    braces: bool = False,
    prefix: str = "",
    suffix: str = "",
    namespace: str = None,
    name: str = None,
    seed: int = None,
) -> dict:
    """
    Generate UUID(s) with multiple options.

    Args:
        version   : UUID version 1|3|4|5|6|7    (default: 4)
        count     : number of UUIDs to generate  (default: 1)
        uppercase : output in uppercase           (default: False)
        hyphens   : include hyphens              (default: True)
        braces    : wrap in curly braces {}       (default: False)
        prefix    : add prefix to each UUID       (default: '')
        suffix    : add suffix to each UUID       (default: '')
        namespace : namespace for v3/v5           (default: DNS)
                    dns | url | oid | x500 | custom UUID
        name      : name for v3/v5               (required for v3/v5)
        seed      : random seed for reproducible  (default: None)

    UUID Versions:
        v1  → time-based + MAC address
        v3  → MD5 hash of namespace + name
        v4  → random (most common)
        v5  → SHA-1 hash of namespace + name
        v6  → reordered time-based (sortable)
        v7  → Unix timestamp-based (sortable, modern)

    Returns:
        {
            'uuids'    : list,
            'version'  : int,
            'count'    : int,
            'options'  : dict,
        }
    """
    import uuid
    import random
    import time
    import os

    # ── Validate ──────────────────────────────────
    if version not in (1, 3, 4, 5, 6, 7):
        raise ValueError(f"Invalid version: {version}. " f"Supported: 1, 3, 4, 5, 6, 7")

    if not (1 <= count <= 1000):
        raise ValueError("count must be between 1 and 1000.")

    if version in (3, 5) and not name:
        raise ValueError(f'UUID v{version} requires a "name" parameter.')

    # ── Namespace for v3/v5 ───────────────────────
    namespace_map = {
        "dns": uuid.NAMESPACE_DNS,
        "url": uuid.NAMESPACE_URL,
        "oid": uuid.NAMESPACE_OID,
        "x500": uuid.NAMESPACE_X500,
    }

    if version in (3, 5):
        if namespace is None or namespace.lower() in namespace_map:
            ns = namespace_map.get(
                (namespace or "dns").lower(),
                uuid.NAMESPACE_DNS,
            )
        else:
            # Try custom UUID namespace
            try:
                ns = uuid.UUID(namespace)
            except ValueError:
                raise ValueError(
                    f'Invalid namespace: "{namespace}". '
                    f"Use: dns, url, oid, x500, or a valid UUID string."
                )

    # ── Seed for reproducible results ─────────────
    if seed is not None:
        random.seed(seed)

    # ── Generate UUIDs ────────────────────────────
    generated = []

    for _ in range(count):
        if version == 1:
            uid = uuid.uuid1()

        elif version == 3:
            uid = uuid.uuid3(ns, name)

        elif version == 4:
            if seed is not None:
                # Seeded random UUID
                rand_bytes = bytes(random.randint(0, 255) for _ in range(16))
                uid = uuid.UUID(bytes=rand_bytes, version=4)
            else:
                uid = uuid.uuid4()

        elif version == 5:
            uid = uuid.uuid5(ns, name)

        elif version == 6:
            uid = _generate_uuid6()

        elif version == 7:
            uid = _generate_uuid7()

        # ── Format output ──────────────────────────
        uid_str = str(uid)

        if not hyphens:
            uid_str = uid_str.replace("-", "")

        if uppercase:
            uid_str = uid_str.upper()

        if braces:
            uid_str = "{" + uid_str + "}"

        if prefix:
            uid_str = prefix + uid_str

        if suffix:
            uid_str = uid_str + suffix

        generated.append(uid_str)

    # ── Build info ────────────────────────────────
    version_info = {
        1: "Time-based + MAC address",
        3: "MD5 hash (namespace + name)",
        4: "Random (cryptographically secure)",
        5: "SHA-1 hash (namespace + name)",
        6: "Reordered time-based (sortable)",
        7: "Unix timestamp-based (sortable, modern)",
    }

    return {
        "uuids": generated,
        "version": version,
        "count": len(generated),
        "description": version_info.get(version, ""),
        "options": {
            "uppercase": uppercase,
            "hyphens": hyphens,
            "braces": braces,
            "prefix": prefix,
            "suffix": suffix,
            "namespace": namespace,
            "name": name,
            "seed": seed,
        },
    }


def _generate_uuid6() -> "uuid.UUID":
    """
    Generate UUID version 6 (reordered time-based, sortable).
    Reorders UUID v1 timestamp for lexicographic sorting.
    """
    import uuid
    import time

    # Get UUID v1 and reorder timestamp
    uid1 = uuid.uuid1()
    uid1_int = uid1.int

    # Extract time fields from v1
    time_low = (uid1_int >> 96) & 0xFFFFFFFF
    time_mid = (uid1_int >> 80) & 0xFFFF
    time_hi = (uid1_int >> 64) & 0x0FFF

    # Reorder: time_hi + time_mid + time_low (sortable)
    time_v6 = (time_hi << 48) | (time_mid << 32) | time_low

    # Rebuild UUID with v6
    clock_seq = (uid1_int >> 48) & 0x3FFF
    node = uid1_int & 0xFFFFFFFFFFFF

    uid6_int = (
        (time_v6 & 0x0FFFFFFFFFFFFFFF) << 64
        | 0x6000_0000_0000_0000
        | clock_seq << 48
        | node
    )

    return uuid.UUID(int=uid6_int)


def _generate_uuid7() -> "uuid.UUID":
    """
    Generate UUID version 7 (Unix timestamp milliseconds, sortable).
    Modern replacement for v1/v6 — monotonically increasing.
    """
    import uuid
    import time
    import os

    # 48-bit Unix timestamp in milliseconds
    ms = int(time.time() * 1000)
    rand_a = int.from_bytes(os.urandom(2), "big") & 0x0FFF
    rand_b = int.from_bytes(os.urandom(8), "big") & 0x3FFFFFFFFFFFFFFF

    uid7_int = (
        (ms & 0xFFFFFFFFFFFF) << 80
        | 0x7000_0000_0000_0000_0000
        | (rand_a & 0x0FFF) << 64
        | 0x8000_0000_0000_0000
        | (rand_b & 0x3FFFFFFFFFFFFFFF)
    )

    return uuid.UUID(int=uid7_int)


def validate_uuid(value: str) -> dict:
    """
    Validate a UUID string and return its info.

    Args:
        value : UUID string to validate

    Returns:
        { 'is_valid', 'version', 'variant', 'uuid', 'formatted' }
    """
    import uuid

    value = str(value).strip()

    # Strip braces and formatting
    clean = value.strip("{}").replace("-", "").strip()

    try:
        # Try parsing with hyphens first
        try:
            uid = uuid.UUID(value.strip("{}"))
        except ValueError:
            # Try without hyphens
            uid = uuid.UUID(clean)

        # ── Format all variants ────────────────────
        uid_str = str(uid)
        formatted = {
            "standard": uid_str,
            "uppercase": uid_str.upper(),
            "no_hyphens": uid_str.replace("-", ""),
            "braces": "{" + uid_str + "}",
            "urn": f"urn:uuid:{uid_str}",
            "int": uid.int,
            "hex": uid.hex,
            "bytes": list(uid.bytes),
        }

        version_info = {
            1: "Time-based + MAC address",
            3: "MD5 hash (namespace + name)",
            4: "Random",
            5: "SHA-1 hash (namespace + name)",
        }

        return {
            "is_valid": True,
            "uuid": uid_str,
            "version": uid.version,
            "variant": str(uid.variant),
            "description": version_info.get(uid.version, "Unknown"),
            "formatted": formatted,
        }

    except ValueError:
        return {
            "is_valid": False,
            "uuid": value,
            "error": f'"{value}" is not a valid UUID.',
        }


def bulk_generate_uuids(
    version: int = 4,
    count: int = 10,
    format: str = "standard",
) -> dict:
    """
    Generate bulk UUIDs in multiple export formats.

    Args:
        version : UUID version                (default: 4)
        count   : number of UUIDs 1-1000     (default: 10)
        format  : standard | csv | json | sql (default: standard)

    Returns:
        { 'uuids', 'export', 'format', 'count' }
    """
    import uuid as _uuid
    import json as _json

    if not (1 <= count <= 1000):
        raise ValueError("count must be between 1 and 1000.")

    result = generate_uuid(version=version, count=count)
    uuids = result["uuids"]

    if format == "csv":
        export = "uuid\n" + "\n".join(uuids)

    elif format == "json":
        export = _json.dumps(uuids, indent=2)

    elif format == "sql":
        rows = ",\n".join(f"  ('{u}')" for u in uuids)
        export = f"INSERT INTO uuids (id) VALUES\n{rows};"

    elif format == "array":
        items = ", ".join(f'"{u}"' for u in uuids)
        export = f"[{items}]"

    else:
        export = "\n".join(uuids)

    return {
        "uuids": uuids,
        "export": export,
        "format": format,
        "count": len(uuids),
        "version": version,
    }


def encrypt_file(
    source,
    password: str,
    algorithm: str = "AES-256-GCM",
    filename: str = None,
) -> dict:
    """
    Encrypt any file using symmetric encryption.

    Args:
        source    : uploaded file object OR raw bytes
        password  : encryption password
        algorithm : AES-256-GCM | AES-256-CBC | ChaCha20   (default: AES-256-GCM)
        filename  : original filename

    Encryption details:
        AES-256-GCM   → authenticated encryption, best security (recommended)
        AES-256-CBC   → classic encryption, widely compatible
        ChaCha20      → fast, secure, good for mobile/embedded

    File format (encrypted output):
        [4 bytes]  magic header  "ENC1"
        [1 byte]   algorithm ID  (1=GCM, 2=CBC, 3=ChaCha20)
        [1 byte]   filename_len
        [N bytes]  filename (utf-8)
        [16 bytes] salt
        [12/16 bytes] IV / nonce
        [16 bytes] tag (GCM only, else empty)
        [N bytes]  encrypted data

    Returns:
        {
            'encrypted_bytes' : bytes,
            'algorithm'       : str,
            'original_size'   : int,
            'encrypted_size'  : int,
            'filename'        : str,
        }
    """
    import os
    import struct
    from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
    from cryptography.hazmat.primitives import hashes, padding
    from cryptography.hazmat.primitives.ciphers import Cipher, algorithms, modes
    from cryptography.hazmat.backends import default_backend

    # ── Read source ───────────────────────────────
    if hasattr(source, "read"):
        raw_bytes = source.read()
        if filename is None and hasattr(source, "name"):
            filename = source.name
    elif isinstance(source, bytes):
        raw_bytes = source
    else:
        raw_bytes = source.encode("utf-8")

    if not raw_bytes:
        raise ValueError("Empty file. Nothing to encrypt.")

    if not password:
        raise ValueError("Password is required.")

    filename = filename or "encrypted_file"
    original_size = len(raw_bytes)

    # ── Validate algorithm ─────────────────────────
    valid_algorithms = ("AES-256-GCM", "AES-256-CBC", "ChaCha20")
    if algorithm not in valid_algorithms:
        raise ValueError(
            f'Invalid algorithm: "{algorithm}". ' f"Must be one of: {valid_algorithms}"
        )

    # ── Derive key from password ───────────────────
    salt = os.urandom(16)
    kdf = PBKDF2HMAC(
        algorithm=hashes.SHA256(),
        length=32,
        salt=salt,
        iterations=600_000,
        backend=default_backend(),
    )
    key = kdf.derive(password.encode("utf-8"))

    # ── Encrypt ───────────────────────────────────
    tag = b""
    algo_id = 0
    encrypted_data = b""

    if algorithm == "AES-256-GCM":
        algo_id = 1
        iv = os.urandom(12)
        cipher = Cipher(
            algorithms.AES(key),
            modes.GCM(iv),
            backend=default_backend(),
        )
        encryptor = cipher.encryptor()
        encrypted_data = encryptor.update(raw_bytes) + encryptor.finalize()
        tag = encryptor.tag  # 16 bytes

    elif algorithm == "AES-256-CBC":
        algo_id = 2
        iv = os.urandom(16)

        # PKCS7 padding
        padder = padding.PKCS7(128).padder()
        padded_data = padder.update(raw_bytes) + padder.finalize()

        cipher = Cipher(
            algorithms.AES(key),
            modes.CBC(iv),
            backend=default_backend(),
        )
        encryptor = cipher.encryptor()
        encrypted_data = encryptor.update(padded_data) + encryptor.finalize()

    elif algorithm == "ChaCha20":
        algo_id = 3
        iv = os.urandom(16)  # ChaCha20 nonce is 16 bytes
        cipher = Cipher(
            algorithms.ChaCha20(key, iv),
            mode=None,
            backend=default_backend(),
        )
        encryptor = cipher.encryptor()
        encrypted_data = encryptor.update(raw_bytes) + encryptor.finalize()

    # ── Build output file ──────────────────────────
    # Header structure:
    # magic(4) + algo_id(1) + filename_len(1) + filename(N)
    # + salt(16) + iv(len) + tag(16 or 0) + encrypted_data
    fname_bytes = filename.encode("utf-8")[:255]
    fname_len = len(fname_bytes)

    header = (
        b"ENC1"  # magic
        + struct.pack("B", algo_id)  # algorithm ID
        + struct.pack("B", fname_len)  # filename length
        + fname_bytes  # filename
        + salt  # 16 bytes salt
        + iv  # IV/nonce
        + tag  # auth tag (GCM only, else b'')
    )

    output_bytes = header + encrypted_data

    return {
        "encrypted_bytes": output_bytes,
        "algorithm": algorithm,
        "original_size": original_size,
        "encrypted_size": len(output_bytes),
        "original_size_kb": round(original_size / 1024, 2),
        "encrypted_size_kb": round(len(output_bytes) / 1024, 2),
        "filename": filename,
        "output_filename": filename + ".enc",
    }


def decrypt_file(
    source,
    password: str,
) -> dict:
    """
    Decrypt a file encrypted by encrypt_file().

    Args:
        source   : uploaded .enc file object OR raw bytes
        password : decryption password

    Returns:
        {
            'decrypted_bytes' : bytes,
            'algorithm'       : str,
            'original_filename': str,
            'original_size'   : int,
            'encrypted_size'  : int,
        }
    """
    import struct
    from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
    from cryptography.hazmat.primitives import hashes, padding
    from cryptography.hazmat.primitives.ciphers import Cipher, algorithms, modes
    from cryptography.hazmat.backends import default_backend
    from cryptography.exceptions import InvalidTag

    # ── Read source ───────────────────────────────
    if hasattr(source, "read"):
        enc_bytes = source.read()
    elif isinstance(source, bytes):
        enc_bytes = source
    else:
        raise ValueError("Invalid source.")

    if not password:
        raise ValueError("Password is required.")

    encrypted_size = len(enc_bytes)

    # ── Validate magic header ──────────────────────
    if len(enc_bytes) < 4 or enc_bytes[:4] != b"ENC1":
        raise ValueError(
            "Invalid encrypted file. " "File was not encrypted with this tool."
        )

    # ── Parse header ──────────────────────────────
    offset = 4

    algo_id = struct.unpack("B", enc_bytes[offset : offset + 1])[0]
    offset += 1

    fname_len = struct.unpack("B", enc_bytes[offset : offset + 1])[0]
    offset += 1

    filename = enc_bytes[offset : offset + fname_len].decode("utf-8", errors="replace")
    offset += fname_len

    salt = enc_bytes[offset : offset + 16]
    offset += 16

    # ── Algorithm map ─────────────────────────────
    algo_map = {
        1: "AES-256-GCM",
        2: "AES-256-CBC",
        3: "ChaCha20",
    }

    if algo_id not in algo_map:
        raise ValueError(f"Unknown algorithm ID: {algo_id}.")

    algorithm = algo_map[algo_id]

    # ── IV size by algorithm ───────────────────────
    iv_sizes = {
        "AES-256-GCM": 12,
        "AES-256-CBC": 16,
        "ChaCha20": 16,
    }
    iv_size = iv_sizes[algorithm]

    iv = enc_bytes[offset : offset + iv_size]
    offset += iv_size

    # ── Auth tag (GCM only) ────────────────────────
    tag = b""
    if algorithm == "AES-256-GCM":
        tag = enc_bytes[offset : offset + 16]
        offset += 16

    encrypted_data = enc_bytes[offset:]

    # ── Derive key ────────────────────────────────
    kdf = PBKDF2HMAC(
        algorithm=hashes.SHA256(),
        length=32,
        salt=salt,
        iterations=600_000,
        backend=default_backend(),
    )
    key = kdf.derive(password.encode("utf-8"))

    # ── Decrypt ───────────────────────────────────
    try:
        if algorithm == "AES-256-GCM":
            cipher = Cipher(
                algorithms.AES(key),
                modes.GCM(iv, tag),
                backend=default_backend(),
            )
            decryptor = cipher.decryptor()
            decrypted_data = decryptor.update(encrypted_data) + decryptor.finalize()

        elif algorithm == "AES-256-CBC":
            cipher = Cipher(
                algorithms.AES(key),
                modes.CBC(iv),
                backend=default_backend(),
            )
            decryptor = cipher.decryptor()
            padded_data = decryptor.update(encrypted_data) + decryptor.finalize()

            # Remove PKCS7 padding
            unpadder = padding.PKCS7(128).unpadder()
            decrypted_data = unpadder.update(padded_data) + unpadder.finalize()

        elif algorithm == "ChaCha20":
            cipher = Cipher(
                algorithms.ChaCha20(key, iv),
                mode=None,
                backend=default_backend(),
            )
            decryptor = cipher.decryptor()
            decrypted_data = decryptor.update(encrypted_data) + decryptor.finalize()

    except InvalidTag:
        raise ValueError("Wrong password or corrupted file. " "Authentication failed.")
    except Exception as e:
        raise ValueError(f"Decryption failed: {e}")

    return {
        "decrypted_bytes": decrypted_data,
        "algorithm": algorithm,
        "original_filename": filename,
        "original_size": len(decrypted_data),
        "encrypted_size": encrypted_size,
        "original_size_kb": round(len(decrypted_data) / 1024, 2),
        "encrypted_size_kb": round(encrypted_size / 1024, 2),
    }


def get_file_hash(
    source,
    algorithms_list: list = None,
) -> dict:
    """
    Calculate hash(es) of a file.

    Args:
        source          : file object OR bytes OR string
        algorithms_list : list of hash algorithms to compute
                          md5 | sha1 | sha256 | sha512 | sha3_256
                          default: all

    Returns:
        { 'hashes': { 'md5': '...', 'sha256': '...' }, 'size': int }
    """
    import hashlib

    if hasattr(source, "read"):
        data = source.read()
    elif isinstance(source, bytes):
        data = source
    elif isinstance(source, str):
        data = source.encode("utf-8")
    else:
        raise ValueError("Invalid source.")

    if not data:
        raise ValueError("Empty input.")

    supported = {
        "md5": hashlib.md5,
        "sha1": hashlib.sha1,
        "sha256": hashlib.sha256,
        "sha512": hashlib.sha512,
        "sha3_256": hashlib.sha3_256,
        "sha3_512": hashlib.sha3_512,
        "blake2b": hashlib.blake2b,
    }

    if algorithms_list is None:
        algorithms_list = list(supported.keys())

    hashes = {}
    for algo in algorithms_list:
        if algo.lower() in supported:
            h = supported[algo.lower()]()
            h.update(data)
            hashes[algo.lower()] = h.hexdigest()
        else:
            hashes[algo.lower()] = f"unsupported algorithm: {algo}"

    return {
        "hashes": hashes,
        "size": len(data),
        "size_kb": round(len(data) / 1024, 2),
        "size_mb": round(len(data) / (1024 * 1024), 2),
    }


def generate_password(
    length: int = 16,
    count: int = 1,
    uppercase: bool = True,
    lowercase: bool = True,
    digits: bool = True,
    symbols: bool = True,
    exclude_chars: str = "",
    exclude_similar: bool = False,
    exclude_ambiguous: bool = False,
    custom_chars: str = "",
    prefix: str = "",
    suffix: str = "",
    no_repeat: bool = False,
) -> dict:
    """
    Generate secure password(s) with full customization.

    Args:
        length           : password length 4-256         (default: 16)
        count            : number of passwords 1-100     (default: 1)
        uppercase        : include A-Z                   (default: True)
        lowercase        : include a-z                   (default: True)
        digits           : include 0-9                   (default: True)
        symbols          : include !@#$%^&*              (default: True)
        exclude_chars    : specific chars to exclude     (default: '')
        exclude_similar  : exclude 0O1lI                 (default: False)
        exclude_ambiguous: exclude {}[]()/'"`~,;:.<>     (default: False)
        custom_chars     : use only these characters     (default: '')
        prefix           : add prefix to password        (default: '')
        suffix           : add suffix to password        (default: '')
        no_repeat        : no repeating characters       (default: False)

    Returns:
        {
            'passwords'  : list,
            'count'      : int,
            'length'     : int,
            'strength'   : str,
            'entropy'    : float,
            'options'    : dict,
        }
    """
    import secrets
    import string
    import math

    # ── Validate ──────────────────────────────────
    if not (4 <= length <= 256):
        raise ValueError("length must be between 4 and 256.")

    if not (1 <= count <= 100):
        raise ValueError("count must be between 1 and 100.")

    # ── Build character pool ───────────────────────
    if custom_chars:
        pool = custom_chars
    else:
        pool = ""
        if uppercase:
            pool += string.ascii_uppercase
        if lowercase:
            pool += string.ascii_lowercase
        if digits:
            pool += string.digits
        if symbols:
            pool += "!@#$%^&*()-_=+[]{}|;:,.<>?"

        if not pool:
            raise ValueError("At least one character type must be selected.")

    # ── Exclude similar characters ─────────────────
    if exclude_similar:
        for ch in "0O1lI":
            pool = pool.replace(ch, "")

    # ── Exclude ambiguous characters ───────────────
    if exclude_ambiguous:
        for ch in "{}[]()\/'\"`~,;:.<>":
            pool = pool.replace(ch, "")

    # ── Exclude specific characters ────────────────
    for ch in exclude_chars:
        pool = pool.replace(ch, "")

    if not pool:
        raise ValueError(
            "Character pool is empty after exclusions. " "Please adjust your settings."
        )

    # ── Check no_repeat feasibility ────────────────
    actual_length = length - len(prefix) - len(suffix)

    if no_repeat and len(pool) < actual_length:
        raise ValueError(
            f"Cannot generate a {actual_length}-char password "
            f"without repeating from a pool of {len(pool)} characters. "
            f"Reduce length or disable no_repeat."
        )

    # ── Generate passwords ─────────────────────────
    passwords = []
    pool_list = list(pool)

    for _ in range(count):
        if no_repeat:
            # Sample without replacement
            chars = secrets.SystemRandom().sample(pool_list, actual_length)
        else:
            chars = [secrets.choice(pool_list) for _ in range(actual_length)]

        # Ensure at least one char from each required set
        if not custom_chars and not no_repeat:
            required = []
            if uppercase and not exclude_similar:
                uc = [
                    c
                    for c in string.ascii_uppercase
                    if c not in exclude_chars and c not in "0O1lI" * exclude_similar
                ]
                if uc:
                    required.append(secrets.choice(uc))
            if lowercase and not exclude_similar:
                lc = [
                    c
                    for c in string.ascii_lowercase
                    if c not in exclude_chars and c not in "0O1lI" * exclude_similar
                ]
                if lc:
                    required.append(secrets.choice(lc))
            if digits:
                dg = [c for c in string.digits if c not in exclude_chars]
                if dg:
                    required.append(secrets.choice(dg))
            if symbols:
                sy = [c for c in "!@#$%^&*()-_=+[]{}|;:,.<>?" if c not in exclude_chars]
                if sy:
                    required.append(secrets.choice(sy))

            # Replace random positions with required chars
            for i, req_char in enumerate(required):
                if i < len(chars):
                    pos = secrets.randbelow(len(chars))
                    chars[pos] = req_char

        password = prefix + "".join(chars) + suffix
        passwords.append(password)

    # ── Calculate entropy ──────────────────────────
    pool_size = len(pool)
    entropy = actual_length * math.log2(pool_size) if pool_size > 1 else 0

    # ── Calculate strength ─────────────────────────
    strength = _calculate_strength(entropy)

    # ── Crack time estimate ───────────────────────
    crack_time = _estimate_crack_time(entropy)

    return {
        "passwords": passwords,
        "count": len(passwords),
        "length": length,
        "pool_size": pool_size,
        "entropy": round(entropy, 2),
        "strength": strength,
        "crack_time": crack_time,
        "options": {
            "uppercase": uppercase,
            "lowercase": lowercase,
            "digits": digits,
            "symbols": symbols,
            "exclude_similar": exclude_similar,
            "exclude_ambiguous": exclude_ambiguous,
            "exclude_chars": exclude_chars,
            "no_repeat": no_repeat,
            "prefix": prefix,
            "suffix": suffix,
            "custom_chars": custom_chars,
        },
    }


def _calculate_strength(entropy: float) -> str:
    """Calculate password strength from entropy bits."""
    if entropy < 28:
        return "very_weak"
    if entropy < 36:
        return "weak"
    if entropy < 60:
        return "moderate"
    if entropy < 128:
        return "strong"
    return "very_strong"


def _estimate_crack_time(entropy: float) -> dict:
    """
    Estimate time to crack password at different attack speeds.

    Attack speeds (guesses/second):
        online_throttled  : 100/s      (online with rate limiting)
        online_unthrottled: 10,000/s   (online without limiting)
        offline_slow      : 1M/s       (bcrypt/scrypt)
        offline_fast      : 100B/s     (MD5/SHA1 GPU)
        massive_attack    : 100T/s     (nation state)
    """
    import math

    combinations = 2**entropy

    speeds = {
        "online_throttled": 100,
        "online_unthrottled": 10_000,
        "offline_slow": 1_000_000,
        "offline_fast": 100_000_000_000,
        "massive_attack": 100_000_000_000_000,
    }

    def format_time(seconds: float) -> str:
        if seconds < 1:
            return "less than a second"
        if seconds < 60:
            return f"{int(seconds)} seconds"
        if seconds < 3600:
            return f"{int(seconds/60)} minutes"
        if seconds < 86400:
            return f"{int(seconds/3600)} hours"
        if seconds < 2592000:
            return f"{int(seconds/86400)} days"
        if seconds < 31536000:
            return f"{int(seconds/2592000)} months"
        if seconds < 3153600000:
            return f"{int(seconds/31536000)} years"
        centuries = seconds / 3153600000
        if centuries < 1000:
            return f"{int(centuries)} centuries"
        return "longer than the age of the universe"

    return {
        name: format_time(combinations / speed / 2) for name, speed in speeds.items()
    }


def check_password_strength(password: str) -> dict:
    """
    Analyze the strength of an existing password.

    Args:
        password : password string to analyze

    Returns:
        {
            'strength'   : str,
            'entropy'    : float,
            'score'      : int,
            'issues'     : list,
            'suggestions': list,
            'crack_time' : dict,
        }
    """
    import math
    import string
    import re

    if not password:
        raise ValueError("Password cannot be empty.")

    length = len(password)

    # ── Build charset size ────────────────────────
    charset = 0
    has_upper = bool(re.search(r"[A-Z]", password))
    has_lower = bool(re.search(r"[a-z]", password))
    has_digit = bool(re.search(r"\d", password))
    has_symbol = bool(re.search(r"[^A-Za-z0-9]", password))
    has_space = " " in password

    if has_upper:
        charset += 26
    if has_lower:
        charset += 26
    if has_digit:
        charset += 10
    if has_symbol:
        charset += 32
    if has_space:
        charset += 1

    if charset == 0:
        charset = 1

    entropy = length * math.log2(charset)

    # ── Score (0-100) ─────────────────────────────
    score = 0
    score += min(length * 4, 40)  # length up to 40pts
    score += 10 if has_upper else 0
    score += 10 if has_lower else 0
    score += 10 if has_digit else 0
    score += 15 if has_symbol else 0
    score += 5 if length >= 20 else 0
    score = min(score, 100)

    # ── Issues ────────────────────────────────────
    issues = []
    suggestions = []

    if length < 8:
        issues.append("Too short (less than 8 characters).")
        suggestions.append("Use at least 8 characters.")

    if length < 12:
        suggestions.append("Consider using 12+ characters for better security.")

    if not has_upper:
        issues.append("No uppercase letters.")
        suggestions.append("Add uppercase letters (A-Z).")

    if not has_lower:
        issues.append("No lowercase letters.")
        suggestions.append("Add lowercase letters (a-z).")

    if not has_digit:
        issues.append("No digits.")
        suggestions.append("Add numbers (0-9).")

    if not has_symbol:
        issues.append("No special characters.")
        suggestions.append("Add symbols (!@#$%^&*).")

    # Check for repeating patterns
    if re.search(r"(.)\1{2,}", password):
        issues.append('Contains repeating characters (e.g. "aaa").')
        suggestions.append("Avoid repeating the same character multiple times.")

    # Check for sequential chars
    if re.search(
        r"(012|123|234|345|456|567|678|789|890|abc|bcd|cde)", password.lower()
    ):
        issues.append('Contains sequential characters (e.g. "123", "abc").')
        suggestions.append("Avoid predictable sequences.")

    # Common patterns
    common_patterns = [
        "password",
        "qwerty",
        "admin",
        "123456",
        "letmein",
        "welcome",
        "monkey",
        "dragon",
    ]
    if any(p in password.lower() for p in common_patterns):
        issues.append("Contains common password patterns.")
        suggestions.append("Avoid common words and patterns.")

    strength = _calculate_strength(entropy)
    crack_time = _estimate_crack_time(entropy)

    return {
        "password": "*" * min(len(password), 4) + "...",  # masked
        "length": length,
        "strength": strength,
        "score": score,
        "entropy": round(entropy, 2),
        "charset_size": charset,
        "has_uppercase": has_upper,
        "has_lowercase": has_lower,
        "has_digits": has_digit,
        "has_symbols": has_symbol,
        "issues": issues,
        "suggestions": suggestions,
        "crack_time": crack_time,
    }


def generate_passphrase(
    words: int = 4,
    count: int = 1,
    separator: str = "-",
    capitalize: bool = True,
    include_digit: bool = True,
) -> dict:
    """
    Generate memorable passphrase using random words.

    Args:
        words        : number of words 3-10     (default: 4)
        count        : number of passphrases    (default: 1)
        separator    : word separator           (default: -)
        capitalize   : capitalize each word     (default: True)
        include_digit: append random digit      (default: True)

    Returns:
        { 'passphrases', 'count', 'words', 'entropy', 'strength' }
    """
    import secrets
    import math

    if not (3 <= words <= 10):
        raise ValueError("words must be between 3 and 10.")
    if not (1 <= count <= 100):
        raise ValueError("count must be between 1 and 100.")

    # ── Word list (EFF large wordlist subset) ──────
    word_list = [
        "apple",
        "brave",
        "cloud",
        "dance",
        "earth",
        "flame",
        "grace",
        "happy",
        "ivory",
        "joker",
        "kneel",
        "lemon",
        "magic",
        "noble",
        "ocean",
        "peace",
        "queen",
        "river",
        "stone",
        "tiger",
        "ultra",
        "vapor",
        "water",
        "xenon",
        "yacht",
        "zebra",
        "amber",
        "blaze",
        "crisp",
        "daisy",
        "eagle",
        "frost",
        "gloom",
        "haste",
        "input",
        "jolly",
        "karma",
        "lunar",
        "marsh",
        "nerve",
        "orbit",
        "pixel",
        "quill",
        "rainy",
        "solar",
        "thorn",
        "umbra",
        "vivid",
        "windy",
        "xeric",
        "young",
        "zesty",
        "agile",
        "blend",
        "coral",
        "drift",
        "elite",
        "fauna",
        "globe",
        "haven",
        "ideal",
        "jazzy",
        "knack",
        "light",
        "mocha",
        "night",
        "olive",
        "prism",
        "quirk",
        "radio",
        "sigma",
        "token",
        "ultra",
        "vague",
        "woken",
        "axiom",
        "birch",
        "cedar",
        "depot",
        "epoch",
        "flair",
        "giant",
        "hyper",
        "index",
        "joint",
        "kraft",
        "lilac",
        "maple",
        "north",
        "oasis",
        "piano",
        "quota",
        "radar",
        "shelf",
        "trove",
        "unity",
        "viola",
        "waltz",
        "exist",
        "years",
        "azure",
        "boxer",
        "cider",
        "delta",
        "ember",
        "field",
        "grand",
        "herbs",
        "inlet",
        "jewel",
        "kapok",
        "largo",
        "mango",
        "niche",
        "onset",
        "perch",
        "quaff",
        "renew",
        "swamp",
        "tryst",
        "upper",
        "verve",
        "wagon",
        "extra",
        "yield",
    ]

    passphrases = []

    for _ in range(count):
        selected = [secrets.choice(word_list) for _ in range(words)]

        if capitalize:
            selected = [w.capitalize() for w in selected]

        phrase = separator.join(selected)

        if include_digit:
            phrase += separator + str(secrets.randbelow(9999)).zfill(4)

        passphrases.append(phrase)

    # ── Entropy calculation ────────────────────────
    # Each word adds log2(wordlist_size) bits
    word_entropy = words * math.log2(len(word_list))
    digit_entropy = math.log2(9999) if include_digit else 0
    total_entropy = word_entropy + digit_entropy

    return {
        "passphrases": passphrases,
        "count": len(passphrases),
        "words": words,
        "separator": separator,
        "entropy": round(total_entropy, 2),
        "strength": _calculate_strength(total_entropy),
        "crack_time": _estimate_crack_time(total_entropy),
    }


def generate_hash(
    source,
    algorithms_list: list = None,
    encoding: str = "utf-8",
    hmac_key: str = None,
    output_format: str = "hex",
) -> dict:
    """
    Generate hash(es) from text or file.

    Args:
        source          : text string | file object | bytes
        algorithms_list : list of algorithms to use
                          md5 | sha1 | sha224 | sha256 | sha384 | sha512
                          sha3_224 | sha3_256 | sha3_384 | sha3_512
                          blake2b | blake2s | shake_128 | shake_256
                          default: all
        encoding        : text encoding                 (default: utf-8)
        hmac_key        : HMAC secret key              (default: None)
        output_format   : hex | base64 | base64url | int (default: hex)

    Returns:
        {
            'hashes'      : dict,
            'input_type'  : str,
            'input_size'  : int,
            'encoding'    : str,
            'output_format': str,
            'is_hmac'     : bool,
        }
    """
    import hashlib
    import hmac as hmac_lib
    import base64

    # ── Read source ───────────────────────────────
    input_type = "text"

    if hasattr(source, "read"):
        data = source.read()
        input_type = "file"
        if isinstance(data, str):
            data = data.encode(encoding)
    elif isinstance(source, bytes):
        data = source
        input_type = "bytes"
    elif isinstance(source, str):
        data = source.encode(encoding)
        input_type = "text"
    else:
        raise ValueError("source must be a string, bytes, or file object.")

    if not data:
        raise ValueError("Empty input. Nothing to hash.")

    # ── Validate output format ─────────────────────
    valid_formats = ("hex", "base64", "base64url", "int")
    if output_format not in valid_formats:
        raise ValueError(
            f'Invalid output_format: "{output_format}". '
            f"Must be one of: {valid_formats}"
        )

    # ── All supported algorithms ───────────────────
    all_algorithms = {
        "md5": {"fn": hashlib.md5, "bits": 128, "secure": False},
        "sha1": {"fn": hashlib.sha1, "bits": 160, "secure": False},
        "sha224": {"fn": hashlib.sha224, "bits": 224, "secure": True},
        "sha256": {"fn": hashlib.sha256, "bits": 256, "secure": True},
        "sha384": {"fn": hashlib.sha384, "bits": 384, "secure": True},
        "sha512": {"fn": hashlib.sha512, "bits": 512, "secure": True},
        "sha3_224": {"fn": hashlib.sha3_224, "bits": 224, "secure": True},
        "sha3_256": {"fn": hashlib.sha3_256, "bits": 256, "secure": True},
        "sha3_384": {"fn": hashlib.sha3_384, "bits": 384, "secure": True},
        "sha3_512": {"fn": hashlib.sha3_512, "bits": 512, "secure": True},
        "blake2b": {"fn": hashlib.blake2b, "bits": 512, "secure": True},
        "blake2s": {"fn": hashlib.blake2s, "bits": 256, "secure": True},
        "shake_128": {"fn": hashlib.shake_128, "bits": 128, "secure": True},
        "shake_256": {"fn": hashlib.shake_256, "bits": 256, "secure": True},
    }

    # ── Select algorithms ─────────────────────────
    if algorithms_list is None:
        selected = list(all_algorithms.keys())
    else:
        selected = [a.lower().strip() for a in algorithms_list]
        invalid = [a for a in selected if a not in all_algorithms]
        if invalid:
            raise ValueError(
                f"Unsupported algorithm(s): {invalid}. "
                f"Supported: {list(all_algorithms.keys())}"
            )

    # ── Format digest ─────────────────────────────
    def format_digest(digest_bytes: bytes) -> str:
        if output_format == "hex":
            return digest_bytes.hex()
        elif output_format == "base64":
            return base64.b64encode(digest_bytes).decode("utf-8")
        elif output_format == "base64url":
            return base64.urlsafe_b64encode(digest_bytes).decode("utf-8")
        elif output_format == "int":
            return str(int.from_bytes(digest_bytes, "big"))
        return digest_bytes.hex()

    # ── Compute hashes ────────────────────────────
    hashes = {}
    is_hmac = bool(hmac_key)
    hmac_key_b = hmac_key.encode(encoding) if hmac_key else None

    for algo in selected:
        info = all_algorithms[algo]
        try:
            if is_hmac:
                # HMAC mode
                if algo in ("shake_128", "shake_256"):
                    hashes[algo] = {
                        "hash": "[HMAC not supported for SHAKE]",
                        "bits": info["bits"],
                        "secure": info["secure"],
                        "error": True,
                    }
                    continue

                h = hmac_lib.new(
                    hmac_key_b,
                    data,
                    info["fn"],
                )
                digest = h.digest()

            else:
                # Regular hash
                if algo == "shake_128":
                    h = hashlib.shake_128(data)
                    digest = h.digest(16)  # 128 bits = 16 bytes
                elif algo == "shake_256":
                    h = hashlib.shake_256(data)
                    digest = h.digest(32)  # 256 bits = 32 bytes
                else:
                    h = info["fn"](data)
                    digest = h.digest()

            hashes[algo] = {
                "hash": format_digest(digest),
                "bits": info["bits"],
                "length": (
                    len(digest) * 2
                    if output_format == "hex"
                    else len(format_digest(digest))
                ),
                "secure": info["secure"],
                "error": False,
            }

        except Exception as e:
            hashes[algo] = {
                "hash": None,
                "error": True,
                "message": str(e),
            }

    return {
        "hashes": hashes,
        "input_type": input_type,
        "input_size": len(data),
        "input_size_kb": round(len(data) / 1024, 2),
        "encoding": encoding,
        "output_format": output_format,
        "is_hmac": is_hmac,
        "algorithm_count": len(hashes),
    }


def compare_hashes(
    hash1: str,
    hash2: str,
) -> dict:
    """
    Securely compare two hash strings.
    Uses constant-time comparison to prevent timing attacks.

    Args:
        hash1 : first hash string
        hash2 : second hash string

    Returns:
        { 'match', 'hash1_length', 'hash2_length' }
    """
    import hmac

    h1 = hash1.strip().lower()
    h2 = hash2.strip().lower()

    # Constant-time comparison
    match = hmac.compare_digest(h1, h2)

    return {
        "match": match,
        "hash1_length": len(h1),
        "hash2_length": len(h2),
        "message": "Hashes match." if match else "Hashes do not match.",
    }


def generate_checksum(
    source,
    algorithm: str = "sha256",
) -> dict:
    """
    Generate a file checksum for integrity verification.

    Args:
        source    : file object OR bytes
        algorithm : hash algorithm to use    (default: sha256)

    Returns:
        { 'checksum', 'algorithm', 'filename', 'size' }
    """
    result = generate_hash(
        source,
        algorithms_list=[algorithm],
        output_format="hex",
    )

    return {
        "checksum": result["hashes"][algorithm]["hash"],
        "algorithm": algorithm,
        "size_kb": result["input_size_kb"],
        "size": result["input_size"],
    }


def word_to_jpg(
    source,
    filename: str = "document.docx",
    dpi: int = 200,
    quality: int = 85,
    pages: list = None,
) -> list[dict]:
    """
    Convert Word (.docx) → JPG images (one per page).

    Strategy:
        1. Convert Word → PDF (using word_to_pdf service)
        2. Convert PDF pages → JPG images (using pymupdf)

    Args:
        source   : uploaded file object OR raw bytes
        filename : original filename
        dpi      : image resolution 72-600    (default: 200)
        quality  : JPG quality 1-95           (default: 85)
        pages    : list of page numbers (1-based) None = all pages

    Returns:
        list of {
            'page'    : int,
            'bytes'   : bytes,
            'width'   : int,
            'height'  : int,
            'filename': str,
            'size_kb' : float,
        }
    """
    import fitz
    from PIL import Image

    # ── Read file bytes ────────────────────────────
    if hasattr(source, "read"):
        file_bytes = source.read()
    else:
        file_bytes = source

    if not file_bytes:
        raise ValueError("Empty file.")

    # ── Step 1: Word → PDF ─────────────────────────
    pdf_bytes = word_to_pdf(
        io.BytesIO(file_bytes),
        filename=filename,
    )

    if not pdf_bytes:
        raise RuntimeError("Failed to convert Word to PDF.")

    # ── Step 2: PDF → JPG ─────────────────────────
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    total_pages = len(doc)

    # ── Validate pages ────────────────────────────
    if pages is not None:
        invalid = [p for p in pages if not (1 <= p <= total_pages)]
        if invalid:
            raise ValueError(
                f"Invalid page numbers: {invalid}. "
                f"Document has {total_pages} pages (1-{total_pages})."
            )
        pages_to_convert = pages
    else:
        pages_to_convert = list(range(1, total_pages + 1))

    # ── Convert each page → JPG ───────────────────
    zoom = dpi / 72
    matrix = fitz.Matrix(zoom, zoom)
    results = []

    base_name = (
        filename.replace(".docx", "")
        .replace(".doc", "")
        .replace(".DOCX", "")
        .replace(".DOC", "")
    )

    for page_num in pages_to_convert:
        page = doc[page_num - 1]
        pixmap = page.get_pixmap(matrix=matrix, alpha=False)

        image = Image.frombytes(
            "RGB",
            [pixmap.width, pixmap.height],
            pixmap.samples,
        )

        buffer = io.BytesIO()
        image.save(
            buffer,
            format="JPEG",
            quality=quality,
            optimize=True,
        )
        buffer.seek(0)
        jpg_bytes = buffer.read()

        results.append(
            {
                "page": page_num,
                "bytes": jpg_bytes,
                "width": pixmap.width,
                "height": pixmap.height,
                "filename": f"{base_name}_page_{page_num}.jpg",
                "size_kb": round(len(jpg_bytes) / 1024, 2),
            }
        )

    doc.close()
    return results
